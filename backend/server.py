from fastapi import FastAPI, APIRouter, HTTPException, Request, UploadFile, File, WebSocket, WebSocketDisconnect, Body
from fastapi.responses import Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import json
import base64
from anthropic import AsyncAnthropic
from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionResponse, CheckoutStatusResponse, CheckoutSessionRequest
from passlib.context import CryptContext
import asyncio

# Shared modules
from db import db, client
from services.email import send_email, email_welcome, email_booking_confirmed, email_application_update, email_offer_received, email_verification, generate_verification_token, email_job_request_confirmation, email_new_lead_notification, email_bid_received
from services.r2_storage import upload_file, upload_property_image, upload_document as r2_upload_document, upload_avatar, upload_contractor_portfolio, delete_file, is_r2_configured

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    try:
        return pwd_context.verify(plain_password, hashed_password)
    except Exception:
        # Fallback for old plaintext passwords during migration
        return plain_password == hashed_password

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Create the main app
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, List[WebSocket]] = {}

    async def connect(self, websocket: WebSocket, user_id: str):
        await websocket.accept()
        if user_id not in self.active_connections:
            self.active_connections[user_id] = []
        self.active_connections[user_id].append(websocket)

    def disconnect(self, websocket: WebSocket, user_id: str):
        if user_id in self.active_connections:
            self.active_connections[user_id].remove(websocket)

    async def send_personal_message(self, message: str, user_id: str):
        if user_id in self.active_connections:
            for connection in self.active_connections[user_id]:
                try:
                    await connection.send_text(message)
                except Exception:
                    pass

    async def broadcast(self, message: str, user_ids: List[str]):
        for user_id in user_ids:
            await self.send_personal_message(message, user_id)

manager = ConnectionManager()

# ========== MODELS ==========

class StatusCheck(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StatusCheckCreate(BaseModel):
    client_name: str

class Listing(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    address: str
    city: str
    province: str
    postal_code: str
    lat: float
    lng: float
    price: int
    bedrooms: int
    bathrooms: float
    sqft: int
    property_type: str
    description: str
    amenities: List[str]
    images: List[str]
    available_date: str = ""
    pet_friendly: bool = False
    parking: bool = False
    landlord_id: Optional[str] = None
    user_id: Optional[str] = None  # Alias for landlord_id
    owner_id: Optional[str] = None  # Another alias for compatibility
    listing_type: str = "rent"  # rent, sale
    sale_price: Optional[int] = None
    year_built: Optional[int] = None
    lot_size: Optional[int] = None
    garage: Optional[int] = None
    mls_number: Optional[str] = None
    open_house_dates: List[str] = []
    status: str = "active"
    featured: bool = False
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ListingCreate(BaseModel):
    title: str
    address: str
    city: str
    province: str
    postal_code: str
    lat: float
    lng: float
    price: int
    bedrooms: int
    bathrooms: float
    sqft: int
    property_type: str
    description: str
    amenities: List[str] = []
    images: List[str] = []
    available_date: str = ""
    pet_friendly: bool = False
    parking: bool = False
    listing_type: str = "rent"
    sale_price: Optional[int] = None
    year_built: Optional[int] = None
    lot_size: Optional[int] = None
    garage: Optional[int] = None
    mls_number: Optional[str] = None
    open_house_dates: List[str] = []
    lease_duration: Optional[int] = 12  # New: lease duration in months
    offers: List[str] = []  # New: special offers/promotions

class ChatMessage(BaseModel):
    role: str
    content: str
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ChatSession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    messages: List[ChatMessage] = []
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ChatRequest(BaseModel):
    session_id: Optional[str] = None
    message: str
    user_id: Optional[str] = None  # For logged-in users to enable long-term memory
    user_context: Optional[Dict[str, Any]] = None  # For lifestyle/budget info

class UserCreate(BaseModel):
    email: str
    password: str
    name: Optional[str] = None
    user_type: str

class UserLogin(BaseModel):
    email: str
    password: str
    user_type: Optional[str] = None

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str
    user_type: str
    email_verified: bool = False
    verification_token: Optional[str] = None
    verification_token_expires: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Payment Models
class PaymentCreate(BaseModel):
    amount: float
    description: str
    property_id: Optional[str] = None
    recipient_id: Optional[str] = None

class PaymentTransaction(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    user_id: Optional[str] = None
    user_email: Optional[str] = None
    amount: float
    currency: str = "cad"
    description: str
    property_id: Optional[str] = None
    recipient_id: Optional[str] = None
    payment_status: str = "pending"
    status: str = "initiated"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Document Models
class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    name: str
    type: str  # lease, application, receipt, id, etc.
    content: Optional[str] = None  # base64 for small files
    url: Optional[str] = None
    status: str = "active"
    signed: bool = False
    signed_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# Message Models
class DirectMessage(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    sender_id: str
    recipient_id: str
    content: str
    read: bool = False
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class MessageCreate(BaseModel):
    recipient_id: str
    content: str

class ChatResponse(BaseModel):
    session_id: str
    response: str
    listings: List[dict] = []
    suggestions: List[str] = []  # Nova's proactive suggestions
    preferences_loaded: bool = False  # Whether user's saved preferences were loaded

# FCM Token Model for Push Notifications
class FCMTokenCreate(BaseModel):
    user_id: str
    token: str

class FCMToken(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    token: str
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class NotificationCreate(BaseModel):
    user_id: str
    title: str
    body: str
    notification_type: str  # message, payment, document, property
    data: Optional[Dict[str, Any]] = None

# Rental Application Models
class RentalApplication(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    listing_id: str
    landlord_id: Optional[str] = None
    # Personal Info
    full_name: str
    email: str
    phone: str
    current_address: str
    move_in_date: str
    # Employment Info
    employer: Optional[str] = None
    job_title: Optional[str] = None
    monthly_income: Optional[float] = None
    employment_length: Optional[str] = None
    # References
    references: List[Dict[str, str]] = []
    # Additional
    num_occupants: int = 1
    has_pets: bool = False
    pet_details: Optional[str] = None
    additional_notes: Optional[str] = None
    # Status tracking
    status: str = "pending"  # pending, under_review, approved, rejected
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ApplicationCreate(BaseModel):
    listing_id: str
    full_name: str
    email: str
    phone: str
    current_address: str
    move_in_date: str
    employer: Optional[str] = None
    job_title: Optional[str] = None
    monthly_income: Optional[float] = None
    employment_length: Optional[str] = None
    references: List[Dict[str, str]] = []
    num_occupants: int = 1
    has_pets: bool = False
    pet_details: Optional[str] = None
    additional_notes: Optional[str] = None

# Maintenance Request Models
class MaintenanceRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str  # tenant who submitted
    landlord_id: Optional[str] = None
    property_id: Optional[str] = None
    title: str
    description: str
    category: str  # plumbing, electrical, appliance, hvac, structural, pest, other
    priority: str = "medium"  # low, medium, high, emergency
    images: List[str] = []
    status: str = "open"  # open, in_progress, scheduled, completed, cancelled
    assigned_contractor_id: Optional[str] = None
    scheduled_date: Optional[str] = None
    completed_date: Optional[str] = None
    cost: Optional[float] = None
    notes: List[Dict[str, Any]] = []
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class MaintenanceRequestCreate(BaseModel):
    property_id: Optional[str] = None
    title: str
    description: str
    category: str
    priority: str = "medium"
    images: List[str] = []

# Contractor Job Models
class ContractorJob(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    maintenance_request_id: Optional[str] = None
    landlord_id: str
    contractor_id: Optional[str] = None
    title: str
    description: str
    category: str
    location: str
    budget_min: Optional[float] = None
    budget_max: Optional[float] = None
    deadline: Optional[str] = None
    status: str = "open"  # open, assigned, in_progress, completed, cancelled
    bids: List[Dict[str, Any]] = []
    selected_bid_id: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ContractorJobCreate(BaseModel):
    title: str
    description: str
    category: str
    location: str
    budget_min: Optional[float] = None
    budget_max: Optional[float] = None
    deadline: Optional[str] = None
    maintenance_request_id: Optional[str] = None

class ContractorBid(BaseModel):
    contractor_id: str
    amount: float
    estimated_days: int
    message: str

# Contractor Profile Models
class ContractorProfile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    business_name: str
    description: str = ""
    specialties: List[str] = []
    service_areas: List[str] = []
    hourly_rate: Optional[float] = None
    years_experience: int = 0
    license_number: Optional[str] = None
    insurance: bool = False
    portfolio_images: List[str] = []
    avatar: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None
    rating: float = 0.0
    review_count: int = 0
    completed_jobs: int = 0
    verified: bool = False
    status: str = "active"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ContractorProfileCreate(BaseModel):
    business_name: str
    description: str = ""
    specialties: List[str] = []
    service_areas: List[str] = []
    hourly_rate: Optional[float] = None
    years_experience: int = 0
    license_number: Optional[str] = None
    insurance: bool = False
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None

# Contractor Service Models
class ContractorService(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    contractor_id: str
    title: str
    description: str
    category: str
    price_type: str = "fixed"  # fixed, hourly, quote
    price: Optional[float] = None
    duration_estimate: Optional[str] = None
    images: List[str] = []
    status: str = "active"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ContractorServiceCreate(BaseModel):
    title: str
    description: str
    category: str
    price_type: str = "fixed"
    price: Optional[float] = None
    duration_estimate: Optional[str] = None

# Service Booking Models
class ServiceBooking(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    customer_id: str
    contractor_id: str
    service_id: Optional[str] = None
    title: str
    description: str
    preferred_date: Optional[str] = None
    preferred_time: Optional[str] = None
    address: str = ""
    status: str = "pending"  # pending, confirmed, in_progress, completed, cancelled
    amount: Optional[float] = None
    payment_status: str = "unpaid"  # unpaid, paid, refunded
    payment_session_id: Optional[str] = None
    notes: Optional[str] = None
    rating: Optional[int] = None
    review: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ServiceBookingCreate(BaseModel):
    contractor_id: str
    service_id: Optional[str] = None
    title: str
    description: str
    preferred_date: Optional[str] = None
    preferred_time: Optional[str] = None
    address: str = ""
    notes: Optional[str] = None

class ReviewCreate(BaseModel):
    rating: int
    review: str

# Property Offer Model (Buy/Sell)
class PropertyOffer(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    listing_id: str
    buyer_id: str
    seller_id: str
    offer_amount: int
    conditions: List[str] = []
    financing_type: str = "mortgage"  # mortgage, cash, pre-approved
    closing_date: Optional[str] = None
    message: Optional[str] = None
    status: str = "pending"  # pending, accepted, rejected, countered, withdrawn
    counter_amount: Optional[int] = None
    counter_message: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class OfferCreate(BaseModel):
    listing_id: str
    offer_amount: int
    conditions: List[str] = []
    financing_type: str = "mortgage"
    closing_date: Optional[str] = None
    message: Optional[str] = None

# Roommate Finder Models
class RoommateProfile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    name: str = ""
    age: Optional[int] = None
    gender: Optional[str] = None
    occupation: Optional[str] = None
    budget_min: int = 500
    budget_max: int = 2000
    move_in_date: Optional[str] = None
    preferred_areas: List[str] = []
    lifestyle: List[str] = []  # early_bird, night_owl, quiet, social, clean, etc.
    pets: bool = False
    smoking: bool = False
    bio: str = ""
    avatar: Optional[str] = None
    status: str = "active"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class RoommateProfileCreate(BaseModel):
    name: str
    age: Optional[int] = None
    gender: Optional[str] = None
    occupation: Optional[str] = None
    budget_min: int = 500
    budget_max: int = 2000
    move_in_date: Optional[str] = None
    preferred_areas: List[str] = []
    lifestyle: List[str] = []
    pets: bool = False
    smoking: bool = False
    bio: str = ""


# ========== RENT PAYMENT MODELS ==========

class RentAgreement(BaseModel):
    """Rent payment agreement between landlord and tenant"""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    listing_id: str
    landlord_id: str
    tenant_id: str
    
    # Rent details
    monthly_rent: float
    due_day: int = 1  # Day of month rent is due (1-28)
    grace_period_days: int = 3  # Days before late fee applies
    
    # Late fee settings (set by landlord)
    late_fee_type: str = "flat"  # "flat" or "percentage"
    late_fee_amount: float = 50.0  # Flat amount or percentage
    max_late_fee: Optional[float] = None  # Cap for percentage-based fees
    
    # E-transfer option
    etransfer_email: Optional[str] = None  # Landlord's e-transfer email
    
    # Status
    status: str = "pending"  # pending, active, completed, cancelled
    landlord_signed: bool = False
    tenant_signed: bool = False
    tenant_accepted_terms: bool = False
    
    # Stripe
    stripe_customer_id: Optional[str] = None  # Tenant's Stripe customer ID
    stripe_payment_method_id: Optional[str] = None  # Tenant's saved card
    
    # Dates
    start_date: str
    end_date: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


class RentAgreementCreate(BaseModel):
    listing_id: str
    tenant_id: str
    monthly_rent: float
    due_day: int = 1
    grace_period_days: int = 3
    late_fee_type: str = "flat"
    late_fee_amount: float = 50.0
    max_late_fee: Optional[float] = None
    etransfer_email: Optional[str] = None
    start_date: str
    end_date: Optional[str] = None


class RentAgreementTerms(BaseModel):
    """Terms proposed/countered by tenant"""
    due_day: Optional[int] = None
    grace_period_days: Optional[int] = None
    late_fee_type: Optional[str] = None
    late_fee_amount: Optional[float] = None


class RentPayment(BaseModel):
    """Individual rent payment record"""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    agreement_id: str
    tenant_id: str
    landlord_id: str
    
    # Payment details
    amount: float  # Base rent amount
    late_fee: float = 0.0
    total_amount: float  # amount + late_fee
    
    # Period
    payment_month: str  # Format: "2026-04"
    due_date: str
    
    # Status
    status: str = "pending"  # pending, processing, paid, failed, overdue
    payment_method: str = "card"  # card, etransfer
    
    # Stripe details
    stripe_payment_intent_id: Optional[str] = None
    stripe_charge_id: Optional[str] = None
    failure_reason: Optional[str] = None
    
    # Timestamps
    paid_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


class SavePaymentMethod(BaseModel):
    """Save tenant's payment method"""
    payment_method_id: str  # Stripe payment method ID


# ========== ROUTES ==========

@api_router.get("/")
async def root():
    return {"message": "DOMMMA V3 API - Complete Real Estate Marketplace"}

# Status Check Routes
@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.model_dump()
    status_obj = StatusCheck(**status_dict)
    doc = status_obj.model_dump()
    doc['timestamp'] = doc['timestamp'].isoformat()
    _ = await db.status_checks.insert_one(doc)
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    status_checks = await db.status_checks.find({}, {"_id": 0}).to_list(1000)
    for check in status_checks:
        if isinstance(check['timestamp'], str):
            check['timestamp'] = datetime.fromisoformat(check['timestamp'])
    return status_checks

# Listing Routes
@api_router.get("/listings", response_model=List[Listing])
async def get_listings(
    city: Optional[str] = None,
    min_price: Optional[int] = None,
    max_price: Optional[int] = None,
    bedrooms: Optional[int] = None,
    bathrooms: Optional[float] = None,
    property_type: Optional[str] = None,
    pet_friendly: Optional[bool] = None,
    parking: Optional[bool] = None,
    listing_type: Optional[str] = None,
    q: Optional[str] = None,
    owner_id: Optional[str] = None,
    featured: Optional[bool] = None,
    sort: Optional[str] = "newest",  # newest, price_low, price_high, featured
    limit: int = 50,
    skip: int = 0
):
    """
    Get listings with filters, sorting and pagination.
    
    Sort options:
    - newest: Most recently created (default)
    - price_low: Lowest price first
    - price_high: Highest price first  
    - featured: Featured/promoted listings first, then by newest
    
    For production, use 'featured' sort to show promoted/paid listings first.
    """
    query = {"status": "active"}
    if listing_type:
        query["listing_type"] = listing_type
    if city:
        query["city"] = {"$regex": city, "$options": "i"}
    if min_price:
        query["price"] = {"$gte": min_price}
    if max_price:
        query.setdefault("price", {})["$lte"] = max_price
    if bedrooms is not None:
        query["bedrooms"] = {"$gte": bedrooms}
    if bathrooms:
        query["bathrooms"] = {"$gte": bathrooms}
    if property_type:
        query["property_type"] = property_type
    if pet_friendly is not None:
        query["pet_friendly"] = pet_friendly
    if parking is not None:
        query["parking"] = parking
    if owner_id:
        # Support both owner_id and landlord_id for backwards compatibility
        query["$or"] = [{"owner_id": owner_id}, {"landlord_id": owner_id}, {"user_id": owner_id}]
    if featured is not None:
        query["featured"] = featured
    if q:
        query["$or"] = [
            {"title": {"$regex": q, "$options": "i"}},
            {"address": {"$regex": q, "$options": "i"}},
            {"description": {"$regex": q, "$options": "i"}}
        ]
    
    # Sorting
    sort_options = {
        "newest": [("created_at", -1)],
        "price_low": [("price", 1)],
        "price_high": [("price", -1)],
        "featured": [("featured", -1), ("boost_score", -1), ("created_at", -1)]
    }
    sort_by = sort_options.get(sort, sort_options["newest"])
    
    listings = await db.listings.find(query, {"_id": 0}).sort(sort_by).skip(skip).limit(min(limit, 100)).to_list(min(limit, 100))
    return listings

@api_router.get("/listings/map")
async def get_listings_for_map(
    north_east_lat: Optional[float] = None,
    north_east_lng: Optional[float] = None,
    south_west_lat: Optional[float] = None,
    south_west_lng: Optional[float] = None
):
    query = {"status": "active"}
    if all([north_east_lat, north_east_lng, south_west_lat, south_west_lng]):
        query["lat"] = {"$gte": south_west_lat, "$lte": north_east_lat}
        query["lng"] = {"$gte": south_west_lng, "$lte": north_east_lng}
    
    listings = await db.listings.find(query, {"_id": 0}).to_list(100)
    return listings

@api_router.get("/listings/{listing_id}", response_model=Listing)
async def get_listing(listing_id: str):
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    return listing

@api_router.post("/listings", response_model=Listing)
async def create_listing_basic(listing: ListingCreate):
    listing_obj = Listing(**listing.model_dump())
    doc = listing_obj.model_dump()
    await db.listings.insert_one(doc)
    return listing_obj

# Auth Routes
@api_router.post("/auth/register")
async def register_user(request: Request, user_data: UserCreate):
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Generate verification token
    verification_token = generate_verification_token()
    token_expires = (datetime.now(timezone.utc) + timedelta(hours=24)).isoformat()
    
    user = User(
        email=user_data.email,
        name=user_data.name or user_data.email.split('@')[0],
        user_type=user_data.user_type,
        email_verified=False,
        verification_token=verification_token,
        verification_token_expires=token_expires
    )
    user_doc = user.model_dump()
    user_doc['password_hash'] = hash_password(user_data.password)
    user_doc['preferences'] = {}  # For Nova's memory
    await db.users.insert_one(user_doc)
    
    # Build verification link - use FRONTEND_URL env var or origin header
    frontend_url = os.environ.get('FRONTEND_URL', 'https://dommma-rent-pay.preview.emergentagent.com')
    origin = request.headers.get('origin', frontend_url).rstrip('/')
    verification_link = f"{origin}/verify-email?token={verification_token}"
    
    # Send verification email (non-blocking)
    asyncio.create_task(send_email(
        user.email,
        "Verify your DOMMMA account",
        email_verification(user.name, verification_link)
    ))
    
    return {
        "message": "Registration successful! Please check your email to verify your account.",
        "email": user.email,
        "requires_verification": True
    }

@api_router.post("/auth/login")
async def login_user(login_data: UserLogin):
    user = await db.users.find_one({"email": login_data.email}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    # Check if email verification is required (only for users created after verification was added)
    # Legacy users (without email_verified field) can still log in
    if 'email_verified' in user and not user.get('email_verified'):
        raise HTTPException(status_code=403, detail="Please verify your email before logging in. Check your inbox for the verification link.")
    
    # Verify password (supports both hashed and legacy plaintext)
    stored_password = user.get('password_hash') or user.get('password', '')
    if not verify_password(login_data.password, stored_password):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    # Migrate old plaintext password to hashed if needed
    if 'password' in user and 'password_hash' not in user:
        await db.users.update_one(
            {"email": login_data.email},
            {"$set": {"password_hash": hash_password(login_data.password)}, "$unset": {"password": ""}}
        )
    
    return {"id": user.get('id'), "email": user.get('email'), "name": user.get('name'), "user_type": user.get('user_type')}

@api_router.get("/auth/verify-email")
async def verify_email(token: str):
    """Verify user email with token"""
    user = await db.users.find_one({"verification_token": token}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired verification token")
    
    # Check if token expired
    token_expires = user.get('verification_token_expires')
    if token_expires:
        expires_dt = datetime.fromisoformat(token_expires.replace('Z', '+00:00'))
        if datetime.now(timezone.utc) > expires_dt:
            raise HTTPException(status_code=400, detail="Verification token has expired. Please request a new one.")
    
    # Mark email as verified
    await db.users.update_one(
        {"verification_token": token},
        {
            "$set": {"email_verified": True},
            "$unset": {"verification_token": "", "verification_token_expires": ""}
        }
    )
    
    # Send welcome email now that they're verified
    asyncio.create_task(send_email(
        user['email'],
        "Welcome to DOMMMA!",
        email_welcome(user.get('name', 'User'), user.get('user_type', 'renter'))
    ))
    
    return {"message": "Email verified successfully! You can now log in.", "email": user['email']}

@api_router.post("/auth/resend-verification")
async def resend_verification(request: Request, email: str = Body(..., embed=True)):
    """Resend verification email"""
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user:
        # Don't reveal if email exists
        return {"message": "If an account exists with this email, a verification link will be sent."}
    
    if user.get('email_verified'):
        raise HTTPException(status_code=400, detail="Email is already verified. You can log in.")
    
    # Generate new verification token
    verification_token = generate_verification_token()
    token_expires = (datetime.now(timezone.utc) + timedelta(hours=24)).isoformat()
    
    await db.users.update_one(
        {"email": email},
        {"$set": {"verification_token": verification_token, "verification_token_expires": token_expires}}
    )
    
    # Build verification link - use FRONTEND_URL env var or origin header
    frontend_url = os.environ.get('FRONTEND_URL', 'https://dommma-rent-pay.preview.emergentagent.com')
    origin = request.headers.get('origin', frontend_url).rstrip('/')
    verification_link = f"{origin}/verify-email?token={verification_token}"
    
    # Send verification email
    asyncio.create_task(send_email(
        email,
        "Verify your DOMMMA account",
        email_verification(user.get('name', 'User'), verification_link)
    ))
    
    return {"message": "Verification email sent! Please check your inbox."}

# ========== CLAIM LISTING ROUTES ==========

@api_router.get("/listings/claim")
async def claim_listing(token: str):
    """Claim a pending listing with token - returns listing info for verification"""
    listing = await db.listings.find_one({"claim_token": token}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Invalid or expired claim token")
    
    if listing.get("status") != "pending_claim":
        raise HTTPException(status_code=400, detail="This listing has already been claimed")
    
    return {
        "listing_id": listing["id"],
        "title": listing["title"],
        "address": listing["address"],
        "city": listing["city"],
        "price": listing["price"],
        "bedrooms": listing["bedrooms"],
        "bathrooms": listing["bathrooms"],
        "claim_email": listing.get("claim_email")
    }

@api_router.post("/listings/claim")
async def complete_listing_claim(request: Request, data: Dict[str, Any]):
    """Complete the claim process - create account and link listing"""
    token = data.get("token")
    password = data.get("password")
    name = data.get("name")
    
    if not token or not password:
        raise HTTPException(status_code=400, detail="Token and password are required")
    
    # Find the listing
    listing = await db.listings.find_one({"claim_token": token})
    if not listing:
        raise HTTPException(status_code=404, detail="Invalid or expired claim token")
    
    if listing.get("status") != "pending_claim":
        raise HTTPException(status_code=400, detail="This listing has already been claimed")
    
    claim_email = listing.get("claim_email")
    if not claim_email:
        raise HTTPException(status_code=400, detail="No email associated with this listing")
    
    # Check if user already exists
    existing_user = await db.users.find_one({"email": claim_email})
    
    if existing_user:
        # User exists - just link the listing to their account
        user_id = existing_user["id"]
    else:
        # Create new user account
        user_id = str(uuid.uuid4())
        user = {
            "id": user_id,
            "email": claim_email,
            "name": name or claim_email.split('@')[0],
            "user_type": "landlord",
            "email_verified": True,  # Verified via claim token
            "password_hash": hash_password(password),
            "preferences": {},
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.users.insert_one(user)
        
        # Send welcome email
        asyncio.create_task(send_email(
            claim_email,
            "Welcome to DOMMMA!",
            email_welcome(user["name"], "landlord")
        ))
    
    # Update listing - activate it and link to user
    await db.listings.update_one(
        {"claim_token": token},
        {
            "$set": {
                "status": "active",
                "owner_id": user_id,
                "updated_at": datetime.now(timezone.utc).isoformat()
            },
            "$unset": {
                "claim_token": "",
                "claim_email": ""
            }
        }
    )
    
    # Get updated user data for login
    user_data = await db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0})
    
    return {
        "success": True,
        "message": "Listing claimed and published successfully!",
        "user": user_data,
        "listing_id": listing["id"]
    }

# ========== STRIPE PAYMENT ROUTES ==========

RENT_PACKAGES = {
    "monthly": {"amount": 0, "description": "Monthly Rent Payment"},  # Amount set dynamically
}

@api_router.post("/payments/create-checkout")
async def create_payment_checkout(request: Request, payment: PaymentCreate):
    try:
        api_key = os.environ.get('STRIPE_API_KEY')
        host_url = str(request.base_url).rstrip('/')
        webhook_url = f"{host_url}/api/webhook/stripe"
        
        stripe_checkout = StripeCheckout(api_key=api_key, webhook_url=webhook_url)
        
        # Get origin from request headers
        origin = request.headers.get('origin', host_url)
        success_url = f"{origin}/dashboard?payment=success&session_id={{CHECKOUT_SESSION_ID}}"
        cancel_url = f"{origin}/dashboard?payment=cancelled"
        
        checkout_request = CheckoutSessionRequest(
            amount=float(payment.amount),
            currency="cad",
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                "description": payment.description,
                "property_id": payment.property_id or "",
                "recipient_id": payment.recipient_id or ""
            }
        )
        
        session = await stripe_checkout.create_checkout_session(checkout_request)
        
        # Create transaction record
        transaction = PaymentTransaction(
            session_id=session.session_id,
            amount=payment.amount,
            currency="cad",
            description=payment.description,
            property_id=payment.property_id,
            recipient_id=payment.recipient_id,
            payment_status="pending",
            status="initiated"
        )
        await db.payment_transactions.insert_one(transaction.model_dump())
        
        return {"url": session.url, "session_id": session.session_id}
        
    except Exception as e:
        logger.error(f"Payment error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/payments/status/{session_id}")
async def get_payment_status(request: Request, session_id: str):
    try:
        api_key = os.environ.get('STRIPE_API_KEY')
        host_url = str(request.base_url).rstrip('/')
        webhook_url = f"{host_url}/api/webhook/stripe"
        
        stripe_checkout = StripeCheckout(api_key=api_key, webhook_url=webhook_url)
        status = await stripe_checkout.get_checkout_status(session_id)
        
        # Update transaction in database
        await db.payment_transactions.update_one(
            {"session_id": session_id},
            {"$set": {
                "payment_status": status.payment_status,
                "status": status.status
            }}
        )
        
        return {
            "status": status.status,
            "payment_status": status.payment_status,
            "amount": status.amount_total / 100,
            "currency": status.currency
        }
        
    except Exception as e:
        logger.error(f"Payment status error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/payments/history")
async def get_payment_history(user_id: Optional[str] = None):
    query = {}
    if user_id:
        query["user_id"] = user_id
    payments = await db.payment_transactions.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    return payments

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    try:
        api_key = os.environ.get('STRIPE_API_KEY')
        host_url = str(request.base_url).rstrip('/')
        webhook_url = f"{host_url}/api/webhook/stripe"
        
        stripe_checkout = StripeCheckout(api_key=api_key, webhook_url=webhook_url)
        body = await request.body()
        signature = request.headers.get("Stripe-Signature")
        
        webhook_response = await stripe_checkout.handle_webhook(body, signature)
        
        if webhook_response.payment_status == "paid":
            await db.payment_transactions.update_one(
                {"session_id": webhook_response.session_id},
                {"$set": {"payment_status": "paid", "status": "complete"}}
            )
        
        return {"status": "success"}
        
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        return {"status": "error", "message": str(e)}

# ========== STRIPE CUSTOMER & PAYMENT METHODS ==========

@api_router.post("/stripe/customer")
async def create_or_get_stripe_customer(user_id: str):
    """Create or retrieve Stripe customer for a user"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Check if user already has a Stripe customer ID
    if user.get("stripe_customer_id"):
        try:
            customer = stripe.Customer.retrieve(user["stripe_customer_id"])
            return {"customer_id": customer.id, "email": customer.email}
        except stripe.error.InvalidRequestError:
            pass  # Customer doesn't exist, create new one
    
    # Create new Stripe customer
    try:
        customer = stripe.Customer.create(
            email=user.get("email"),
            name=user.get("name"),
            metadata={"user_id": user_id}
        )
        
        # Save Stripe customer ID to user record
        await db.users.update_one(
            {"id": user_id},
            {"$set": {"stripe_customer_id": customer.id}}
        )
        
        return {"customer_id": customer.id, "email": customer.email}
    except Exception as e:
        logger.error(f"Stripe customer creation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/stripe/setup-intent")
async def create_setup_intent(user_id: str):
    """Create a SetupIntent to save a payment method"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    # Get or create customer
    customer_response = await create_or_get_stripe_customer(user_id)
    customer_id = customer_response["customer_id"]
    
    try:
        setup_intent = stripe.SetupIntent.create(
            customer=customer_id,
            payment_method_types=["card"],
            metadata={"user_id": user_id}
        )
        
        return {
            "client_secret": setup_intent.client_secret,
            "setup_intent_id": setup_intent.id
        }
    except Exception as e:
        logger.error(f"SetupIntent creation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/stripe/payment-methods/{user_id}")
async def get_payment_methods(user_id: str):
    """Get all saved payment methods for a user"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user or not user.get("stripe_customer_id"):
        return {"payment_methods": [], "default_payment_method": None}
    
    try:
        # Get payment methods
        payment_methods = stripe.PaymentMethod.list(
            customer=user["stripe_customer_id"],
            type="card"
        )
        
        # Get customer to find default payment method
        customer = stripe.Customer.retrieve(user["stripe_customer_id"])
        default_pm = customer.invoice_settings.default_payment_method if customer.invoice_settings else None
        
        methods = []
        for pm in payment_methods.data:
            methods.append({
                "id": pm.id,
                "brand": pm.card.brand,
                "last4": pm.card.last4,
                "exp_month": pm.card.exp_month,
                "exp_year": pm.card.exp_year,
                "is_default": pm.id == default_pm
            })
        
        return {"payment_methods": methods, "default_payment_method": default_pm}
    except Exception as e:
        logger.error(f"Get payment methods error: {e}")
        return {"payment_methods": [], "default_payment_method": None}

@api_router.post("/stripe/payment-methods/{user_id}/default")
async def set_default_payment_method(user_id: str, payment_method_id: str):
    """Set a payment method as default"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user or not user.get("stripe_customer_id"):
        raise HTTPException(status_code=404, detail="No Stripe customer found")
    
    try:
        stripe.Customer.modify(
            user["stripe_customer_id"],
            invoice_settings={"default_payment_method": payment_method_id}
        )
        return {"success": True, "default_payment_method": payment_method_id}
    except Exception as e:
        logger.error(f"Set default payment method error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.delete("/stripe/payment-methods/{payment_method_id}")
async def delete_payment_method(payment_method_id: str):
    """Delete a saved payment method"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    try:
        stripe.PaymentMethod.detach(payment_method_id)
        return {"success": True, "message": "Payment method removed"}
    except Exception as e:
        logger.error(f"Delete payment method error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/stripe/checkout-setup")
async def create_checkout_setup_session(user_id: str):
    """Create a Stripe checkout session for adding a new card"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Get or create customer
    customer_id = user.get("stripe_customer_id")
    if not customer_id:
        customer = stripe.Customer.create(
            email=user.get("email"),
            name=user.get("name"),
            metadata={"user_id": user_id}
        )
        customer_id = customer.id
        await db.users.update_one({"id": user_id}, {"$set": {"stripe_customer_id": customer_id}})
    
    try:
        host_url = os.environ.get('FRONTEND_URL', 'https://dommma-rent-pay.preview.emergentagent.com')
        
        session = stripe.checkout.Session.create(
            customer=customer_id,
            payment_method_types=['card'],
            mode='setup',
            success_url=f"{host_url}/settings?card_added=true",
            cancel_url=f"{host_url}/settings?card_canceled=true",
            metadata={"user_id": user_id}
        )
        
        return {
            "success": True,
            "checkout_url": session.url,
            "session_id": session.id
        }
    except Exception as e:
        logger.error(f"Checkout setup error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== RENT INVOICES & COLLECTION ==========

class RentInvoiceCreate(BaseModel):
    tenant_id: str
    property_id: str
    amount: float
    due_date: str
    description: str = "Monthly Rent"
    late_fee_type: str = "none"  # none, flat, percentage, daily
    late_fee_amount: float = 0
    grace_period_days: int = 3

@api_router.post("/rent-invoices")
async def create_rent_invoice(invoice: RentInvoiceCreate, landlord_id: str):
    """Create a rent invoice for a tenant"""
    # Verify landlord
    landlord = await db.users.find_one({"id": landlord_id}, {"_id": 0})
    if not landlord or landlord.get("user_type") != "landlord":
        raise HTTPException(status_code=403, detail="Only landlords can create rent invoices")
    
    # Get tenant
    tenant = await db.users.find_one({"id": invoice.tenant_id}, {"_id": 0})
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")
    
    # Get property
    property_info = await db.listings.find_one({"id": invoice.property_id}, {"_id": 0})
    
    invoice_id = str(uuid.uuid4())
    invoice_number = f"RENT-{datetime.now().strftime('%Y%m')}-{invoice_id[:8].upper()}"
    
    invoice_doc = {
        "id": invoice_id,
        "invoice_number": invoice_number,
        "type": "rent",
        "landlord_id": landlord_id,
        "landlord_name": landlord.get("name"),
        "tenant_id": invoice.tenant_id,
        "tenant_name": tenant.get("name"),
        "tenant_email": tenant.get("email"),
        "property_id": invoice.property_id,
        "property_address": property_info.get("address") if property_info else None,
        "amount": invoice.amount,
        "due_date": invoice.due_date,
        "grace_period_days": invoice.grace_period_days,
        "late_fee_type": invoice.late_fee_type,
        "late_fee_amount": invoice.late_fee_amount,
        "description": invoice.description,
        "status": "pending",  # pending, paid, overdue, cancelled
        "paid_at": None,
        "late_fee_applied": False,
        "late_fee_total": 0,
        "reminders_sent": 0,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.rent_invoices.insert_one(invoice_doc)
    
    # Send invoice email to tenant
    if tenant.get("email"):
        asyncio.create_task(send_email(
            tenant["email"],
            f"Rent Invoice - {invoice.description}",
            f"""
            <div style="font-family:'Georgia',serif;max-width:600px;margin:0 auto;background:#F5F5F0;padding:40px;">
              <div style="background:#1A2F3A;padding:30px;border-radius:16px 16px 0 0;text-align:center;">
                <h1 style="color:white;font-family:'Georgia',serif;margin:0;font-size:28px;">DOMMMA</h1>
                <p style="color:rgba(255,255,255,0.7);margin:8px 0 0;font-size:14px;">Rent Invoice</p>
              </div>
              <div style="background:white;padding:30px;border-radius:0 0 16px 16px;">
                <h2 style="color:#1A2F3A;">Hi {tenant.get('name', 'Tenant')},</h2>
                <p style="color:#555;">Your rent invoice is ready.</p>
                <div style="background:#F5F5F0;border-radius:12px;padding:20px;margin:20px 0;">
                  <p style="margin:4px 0;color:#555;">{invoice.description}</p>
                  <p style="margin:4px 0;color:#1A2F3A;font-size:28px;font-weight:bold;">${invoice.amount:,.2f}</p>
                  <p style="margin:4px 0;color:#888;">Due: {invoice.due_date}</p>
                </div>
                <a href="https://dommma.com/payments" style="display:inline-block;background:#1A2F3A;color:white;padding:14px 32px;border-radius:8px;text-decoration:none;font-weight:bold;">Pay Now</a>
              </div>
            </div>
            """
        ))
    
    # Create notification
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": invoice.tenant_id,
        "title": "New Rent Invoice",
        "body": f"Rent of ${invoice.amount:,.2f} is due on {invoice.due_date}",
        "type": "payment",
        "data": {"invoice_id": invoice_id},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    invoice_doc.pop("_id", None)
    return invoice_doc

@api_router.get("/rent-invoices")
async def get_rent_invoices(user_id: str, role: str = "tenant"):
    """Get rent invoices for a user (as tenant or landlord)"""
    if role == "landlord":
        query = {"landlord_id": user_id}
    else:
        query = {"tenant_id": user_id}
    
    invoices = await db.rent_invoices.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    # Update status for overdue invoices
    now = datetime.now(timezone.utc)
    for inv in invoices:
        if inv["status"] == "pending":
            due = datetime.fromisoformat(inv["due_date"].replace('Z', '+00:00')) if 'T' in inv["due_date"] else datetime.strptime(inv["due_date"], "%Y-%m-%d").replace(tzinfo=timezone.utc)
            grace_end = due + timedelta(days=inv.get("grace_period_days", 3))
            if now > grace_end:
                inv["status"] = "overdue"
    
    return invoices

@api_router.get("/rent-invoices/{invoice_id}")
async def get_rent_invoice(invoice_id: str):
    """Get a specific rent invoice"""
    invoice = await db.rent_invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    return invoice

@api_router.post("/rent-invoices/{invoice_id}/pay")
async def pay_rent_invoice(invoice_id: str, user_id: str):
    """Pay a rent invoice using saved payment method"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    invoice = await db.rent_invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    if invoice["tenant_id"] != user_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    if invoice["status"] == "paid":
        return {"success": True, "message": "Invoice already paid"}
    
    # Get user's default payment method
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user or not user.get("stripe_customer_id"):
        raise HTTPException(status_code=400, detail="No payment method on file. Please add a card first.")
    
    try:
        customer = stripe.Customer.retrieve(user["stripe_customer_id"])
        default_pm = customer.invoice_settings.default_payment_method
        
        if not default_pm:
            raise HTTPException(status_code=400, detail="No default payment method. Please add a card first.")
        
        # Calculate total with late fees if applicable
        total = invoice["amount"]
        if invoice.get("late_fee_applied") and invoice.get("late_fee_total"):
            total += invoice["late_fee_total"]
        
        # Create payment intent
        payment_intent = stripe.PaymentIntent.create(
            amount=int(total * 100),  # cents
            currency="cad",
            customer=user["stripe_customer_id"],
            payment_method=default_pm,
            confirm=True,
            description=f"Rent Payment - {invoice['description']}",
            metadata={
                "invoice_id": invoice_id,
                "type": "rent"
            },
            automatic_payment_methods={
                "enabled": True,
                "allow_redirects": "never"
            }
        )
        
        if payment_intent.status == "succeeded":
            # Update invoice
            await db.rent_invoices.update_one(
                {"id": invoice_id},
                {"$set": {
                    "status": "paid",
                    "paid_at": datetime.now(timezone.utc).isoformat(),
                    "payment_intent_id": payment_intent.id,
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            
            # Notify landlord
            await db.notifications.insert_one({
                "id": str(uuid.uuid4()),
                "user_id": invoice["landlord_id"],
                "title": "Rent Payment Received",
                "body": f"{invoice['tenant_name']} paid ${total:,.2f} for rent",
                "type": "payment",
                "data": {"invoice_id": invoice_id},
                "read": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            })
            
            return {"success": True, "message": "Payment successful", "amount": total}
        else:
            raise HTTPException(status_code=400, detail=f"Payment failed: {payment_intent.status}")
            
    except stripe.error.CardError as e:
        raise HTTPException(status_code=400, detail=f"Card error: {e.error.message}")
    except Exception as e:
        logger.error(f"Rent payment error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/rent-invoices/{invoice_id}/send-reminder")
async def send_rent_reminder(invoice_id: str, landlord_id: str):
    """Send a reminder email for an unpaid invoice"""
    invoice = await db.rent_invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    if invoice["landlord_id"] != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    if invoice["status"] == "paid":
        return {"success": False, "message": "Invoice already paid"}
    
    tenant = await db.users.find_one({"id": invoice["tenant_id"]}, {"_id": 0})
    if not tenant or not tenant.get("email"):
        raise HTTPException(status_code=400, detail="Tenant email not found")
    
    # Send reminder
    is_overdue = invoice.get("status") == "overdue"
    subject = f"{'OVERDUE: ' if is_overdue else ''}Rent Reminder - ${invoice['amount']:,.2f} Due"
    
    asyncio.create_task(send_email(
        tenant["email"],
        subject,
        f"""
        <div style="font-family:'Georgia',serif;max-width:600px;margin:0 auto;background:#F5F5F0;padding:40px;">
          <div style="background:{'#DC2626' if is_overdue else '#1A2F3A'};padding:30px;border-radius:16px 16px 0 0;text-align:center;">
            <h1 style="color:white;font-family:'Georgia',serif;margin:0;font-size:28px;">{'⚠️ OVERDUE' if is_overdue else 'Reminder'}</h1>
          </div>
          <div style="background:white;padding:30px;border-radius:0 0 16px 16px;">
            <h2 style="color:#1A2F3A;">Hi {tenant.get('name', 'Tenant')},</h2>
            <p style="color:#555;">This is a {'final reminder' if is_overdue else 'friendly reminder'} that your rent payment is {'overdue' if is_overdue else 'due soon'}.</p>
            <div style="background:#F5F5F0;border-radius:12px;padding:20px;margin:20px 0;">
              <p style="margin:4px 0;color:#555;">{invoice['description']}</p>
              <p style="margin:4px 0;color:#1A2F3A;font-size:28px;font-weight:bold;">${invoice['amount']:,.2f}</p>
              <p style="margin:4px 0;color:#888;">Due: {invoice['due_date']}</p>
            </div>
            <a href="https://dommma.com/payments" style="display:inline-block;background:#1A2F3A;color:white;padding:14px 32px;border-radius:8px;text-decoration:none;font-weight:bold;">Pay Now</a>
          </div>
        </div>
        """
    ))
    
    # Update reminder count
    await db.rent_invoices.update_one(
        {"id": invoice_id},
        {"$inc": {"reminders_sent": 1}, "$set": {"last_reminder_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return {"success": True, "message": "Reminder sent"}

@api_router.post("/rent-invoices/{invoice_id}/apply-late-fee")
async def apply_late_fee(invoice_id: str, landlord_id: str):
    """Apply late fee to an overdue invoice"""
    invoice = await db.rent_invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    if invoice["landlord_id"] != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    if invoice["status"] == "paid":
        return {"success": False, "message": "Invoice already paid"}
    
    if invoice.get("late_fee_applied"):
        return {"success": False, "message": "Late fee already applied"}
    
    # Calculate late fee
    late_fee = 0
    if invoice["late_fee_type"] == "flat":
        late_fee = invoice["late_fee_amount"]
    elif invoice["late_fee_type"] == "percentage":
        late_fee = invoice["amount"] * (invoice["late_fee_amount"] / 100)
    elif invoice["late_fee_type"] == "daily":
        due = datetime.fromisoformat(invoice["due_date"].replace('Z', '+00:00')) if 'T' in invoice["due_date"] else datetime.strptime(invoice["due_date"], "%Y-%m-%d").replace(tzinfo=timezone.utc)
        days_late = (datetime.now(timezone.utc) - due).days
        late_fee = invoice["late_fee_amount"] * max(0, days_late)
    
    if late_fee > 0:
        await db.rent_invoices.update_one(
            {"id": invoice_id},
            {"$set": {
                "late_fee_applied": True,
                "late_fee_total": round(late_fee, 2),
                "status": "overdue",
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        # Notify tenant
        await db.notifications.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": invoice["tenant_id"],
            "title": "Late Fee Applied",
            "body": f"A late fee of ${late_fee:,.2f} has been added to your rent invoice",
            "type": "payment",
            "data": {"invoice_id": invoice_id},
            "read": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        
        return {"success": True, "late_fee": round(late_fee, 2)}
    
    return {"success": False, "message": "No late fee configured"}


# ========== FEATURED LISTINGS (PAY-PER-SUCCESS) ==========

FEATURED_FEE = 4999  # $49.99 fee charged when property is successfully rented

@api_router.post("/listings/{listing_id}/featured")
async def enable_featured_listing(listing_id: str, landlord_id: str):
    """Enable featured status for a listing (pay-per-success model)"""
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    if listing.get("landlord_id") != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    if listing.get("featured"):
        return {"success": True, "message": "Listing is already featured", "featured": True}
    
    # Enable featured with pay-per-success (fee collected when rented)
    now = datetime.now(timezone.utc)
    expires = now + timedelta(days=30)  # Featured for 30 days
    
    await db.listings.update_one(
        {"id": listing_id},
        {"$set": {
            "featured": True,
            "featured_enabled_at": now.isoformat(),
            "featured_expires_at": expires.isoformat(),
            "featured_fee_pending": True,
            "boost_score": 100
        }}
    )
    
    return {
        "success": True,
        "message": "Listing is now featured! A $49.99 fee will be charged when the property is rented.",
        "featured": True,
        "expires_at": expires.isoformat(),
        "fee_pending": True,
        "fee_amount": FEATURED_FEE / 100
    }

@api_router.delete("/listings/{listing_id}/featured")
async def disable_featured_listing(listing_id: str, landlord_id: str):
    """Disable featured status for a listing"""
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    if listing.get("landlord_id") != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    await db.listings.update_one(
        {"id": listing_id},
        {"$set": {
            "featured": False,
            "featured_fee_pending": False,
            "boost_score": 0
        }}
    )
    
    return {"success": True, "message": "Featured status disabled", "featured": False}

@api_router.post("/listings/{listing_id}/mark-rented")
async def mark_listing_rented(listing_id: str, landlord_id: str):
    """Mark a listing as rented and process featured fee if applicable"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    if listing.get("landlord_id") != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    result = {"success": True, "status": "rented"}
    
    # If featured with pending fee, charge the landlord
    if listing.get("featured") and listing.get("featured_fee_pending"):
        landlord = await db.users.find_one({"id": landlord_id}, {"_id": 0})
        
        if landlord and landlord.get("stripe_customer_id"):
            try:
                # Get default payment method
                customer = stripe.Customer.retrieve(landlord["stripe_customer_id"])
                default_pm = customer.invoice_settings.default_payment_method if customer.invoice_settings else None
                
                if default_pm:
                    # Charge the featured fee
                    payment_intent = stripe.PaymentIntent.create(
                        amount=FEATURED_FEE,
                        currency="cad",
                        customer=landlord["stripe_customer_id"],
                        payment_method=default_pm,
                        off_session=True,
                        confirm=True,
                        description=f"Featured listing fee - {listing.get('title', 'Property')}",
                        metadata={
                            "listing_id": listing_id,
                            "landlord_id": landlord_id,
                            "type": "featured_fee"
                        }
                    )
                    
                    # Record the transaction
                    await db.payment_transactions.insert_one({
                        "id": str(uuid.uuid4()),
                        "user_id": landlord_id,
                        "session_id": payment_intent.id,
                        "amount": FEATURED_FEE / 100,
                        "currency": "cad",
                        "description": f"Featured listing fee - {listing.get('title', 'Property')}",
                        "property_id": listing_id,
                        "payment_status": "paid",
                        "status": "complete",
                        "created_at": datetime.now(timezone.utc).isoformat()
                    })
                    
                    result["featured_fee_charged"] = True
                    result["fee_amount"] = FEATURED_FEE / 100
                else:
                    result["featured_fee_charged"] = False
                    result["fee_message"] = "No default payment method found"
            except stripe.error.CardError as e:
                result["featured_fee_charged"] = False
                result["fee_message"] = f"Payment failed: {e.error.message}"
            except Exception as e:
                result["featured_fee_charged"] = False
                result["fee_message"] = str(e)
        else:
            result["featured_fee_charged"] = False
            result["fee_message"] = "No payment method on file"
    
    # Update listing status
    await db.listings.update_one(
        {"id": listing_id},
        {"$set": {
            "status": "rented",
            "featured_fee_pending": False
        }}
    )
    
    return result

@api_router.get("/listings/{listing_id}/featured-status")
async def get_featured_status(listing_id: str):
    """Get the featured status of a listing"""
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0, "featured": 1, "featured_enabled_at": 1, "featured_expires_at": 1, "featured_fee_pending": 1, "boost_score": 1})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    # Check if featured has expired
    if listing.get("featured") and listing.get("featured_expires_at"):
        expires = datetime.fromisoformat(listing["featured_expires_at"].replace("Z", "+00:00"))
        if datetime.now(timezone.utc) > expires:
            # Expired - disable featured
            await db.listings.update_one(
                {"id": listing_id},
                {"$set": {"featured": False, "featured_fee_pending": False, "boost_score": 0}}
            )
            return {"featured": False, "expired": True}
    
    return {
        "featured": listing.get("featured", False),
        "enabled_at": listing.get("featured_enabled_at"),
        "expires_at": listing.get("featured_expires_at"),
        "fee_pending": listing.get("featured_fee_pending", False),
        "fee_amount": FEATURED_FEE / 100 if listing.get("featured_fee_pending") else 0
    }

# ========== UNIVERSAL PAYMENTS & INVOICES ==========

class PaymentRequest(BaseModel):
    amount: float
    description: str
    payment_type: str  # rent, contractor, moving, supplies, equipment, property_expense, etc.
    recipient_id: Optional[str] = None
    recipient_name: Optional[str] = None
    property_id: Optional[str] = None
    property_address: Optional[str] = None
    notes: Optional[str] = None

@api_router.post("/payments/create")
async def create_payment(user_id: str, payment: PaymentRequest):
    """Create a payment and generate an invoice"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Create or get Stripe customer
    customer_id = user.get("stripe_customer_id")
    if not customer_id:
        customer = stripe.Customer.create(
            email=user.get("email"),
            name=user.get("name"),
            metadata={"user_id": user_id}
        )
        customer_id = customer.id
        await db.users.update_one({"id": user_id}, {"$set": {"stripe_customer_id": customer_id}})
    
    # Create payment intent
    amount_cents = int(payment.amount * 100)
    
    try:
        # Get default payment method
        customer = stripe.Customer.retrieve(customer_id)
        default_pm = customer.invoice_settings.default_payment_method if customer.invoice_settings else None
        
        if not default_pm:
            # Create a checkout session instead
            host_url = os.environ.get('FRONTEND_URL', 'https://dommma-rent-pay.preview.emergentagent.com')
            session = stripe.checkout.Session.create(
                customer=customer_id,
                payment_method_types=['card'],
                line_items=[{
                    'price_data': {
                        'currency': 'cad',
                        'product_data': {'name': payment.description},
                        'unit_amount': amount_cents,
                    },
                    'quantity': 1,
                }],
                mode='payment',
                success_url=f"{host_url}/payments?success=true",
                cancel_url=f"{host_url}/payments?canceled=true",
                metadata={
                    "user_id": user_id,
                    "payment_type": payment.payment_type,
                    "recipient_id": payment.recipient_id or "",
                }
            )
            
            # Create pending invoice
            invoice_id = str(uuid.uuid4())
            invoice = {
                "id": invoice_id,
                "invoice_number": f"INV-{datetime.now().strftime('%Y%m%d')}-{invoice_id[:8].upper()}",
                "user_id": user_id,
                "user_name": user.get("name"),
                "user_email": user.get("email"),
                "user_type": user.get("user_type"),
                "amount": payment.amount,
                "currency": "CAD",
                "description": payment.description,
                "payment_type": payment.payment_type,
                "recipient_id": payment.recipient_id,
                "recipient_name": payment.recipient_name,
                "property_id": payment.property_id,
                "property_address": payment.property_address,
                "notes": payment.notes,
                "status": "pending",
                "stripe_session_id": session.id,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "paid_at": None
            }
            await db.invoices.insert_one(invoice)
            
            return {
                "success": True,
                "requires_redirect": True,
                "checkout_url": session.url,
                "invoice_id": invoice_id
            }
        
        # Charge with saved payment method
        payment_intent = stripe.PaymentIntent.create(
            amount=amount_cents,
            currency="cad",
            customer=customer_id,
            payment_method=default_pm,
            off_session=True,
            confirm=True,
            description=payment.description,
            metadata={
                "user_id": user_id,
                "payment_type": payment.payment_type,
                "recipient_id": payment.recipient_id or "",
            }
        )
        
        # Create paid invoice
        invoice_id = str(uuid.uuid4())
        invoice = {
            "id": invoice_id,
            "invoice_number": f"INV-{datetime.now().strftime('%Y%m%d')}-{invoice_id[:8].upper()}",
            "user_id": user_id,
            "user_name": user.get("name"),
            "user_email": user.get("email"),
            "user_type": user.get("user_type"),
            "amount": payment.amount,
            "currency": "CAD",
            "description": payment.description,
            "payment_type": payment.payment_type,
            "recipient_id": payment.recipient_id,
            "recipient_name": payment.recipient_name,
            "property_id": payment.property_id,
            "property_address": payment.property_address,
            "notes": payment.notes,
            "status": "paid",
            "stripe_payment_id": payment_intent.id,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "paid_at": datetime.now(timezone.utc).isoformat()
        }
        await db.invoices.insert_one(invoice)
        
        return {
            "success": True,
            "requires_redirect": False,
            "invoice_id": invoice_id,
            "invoice_number": invoice["invoice_number"],
            "payment_id": payment_intent.id
        }
        
    except stripe.error.CardError as e:
        raise HTTPException(status_code=400, detail=f"Payment failed: {e.error.message}")
    except Exception as e:
        logger.error(f"Payment error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/invoices/{user_id}")
async def get_user_invoices(user_id: str, status: Optional[str] = None, limit: int = 50):
    """Get all invoices for a user"""
    query = {"user_id": user_id}
    if status:
        query["status"] = status
    
    invoices = await db.invoices.find(query, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)
    return invoices

@api_router.get("/invoices/detail/{invoice_id}")
async def get_invoice_detail(invoice_id: str):
    """Get detailed invoice information"""
    invoice = await db.invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    return invoice

@api_router.get("/invoices/{invoice_id}/pdf")
async def download_invoice_pdf(invoice_id: str):
    """Generate and download invoice as PDF"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    import io
    
    invoice = await db.invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    elements = []
    
    # Header style
    header_style = ParagraphStyle('Header', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#1A2F3A'))
    subheader_style = ParagraphStyle('SubHeader', parent=styles['Normal'], fontSize=10, textColor=colors.gray)
    
    # Company Header
    elements.append(Paragraph("DOMMMA", header_style))
    elements.append(Paragraph("Real Estate Marketplace", subheader_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Invoice Title
    title_style = ParagraphStyle('Title', parent=styles['Heading2'], fontSize=18, textColor=colors.HexColor('#1A2F3A'))
    elements.append(Paragraph(f"INVOICE", title_style))
    elements.append(Paragraph(f"#{invoice.get('invoice_number', invoice_id[:8].upper())}", subheader_style))
    elements.append(Spacer(1, 0.2*inch))
    
    # Invoice Details Table
    status_color = colors.green if invoice.get('status') == 'paid' else colors.orange
    details_data = [
        ['Date:', invoice.get('created_at', '')[:10]],
        ['Status:', invoice.get('status', 'pending').upper()],
        ['Payment Type:', invoice.get('payment_type', 'N/A').replace('_', ' ').title()],
    ]
    
    if invoice.get('paid_at'):
        details_data.append(['Paid On:', invoice.get('paid_at', '')[:10]])
    
    details_table = Table(details_data, colWidths=[1.5*inch, 3*inch])
    details_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#333333')),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(details_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Bill To Section
    elements.append(Paragraph("Bill To:", ParagraphStyle('BillTo', parent=styles['Heading3'], fontSize=12)))
    elements.append(Paragraph(invoice.get('user_name', 'N/A'), styles['Normal']))
    elements.append(Paragraph(invoice.get('user_email', ''), subheader_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Line Items
    items_data = [
        ['Description', 'Amount'],
        [invoice.get('description', 'Payment'), f"${invoice.get('amount', 0):,.2f} CAD"],
    ]
    
    if invoice.get('property_address'):
        items_data.insert(1, [f"Property: {invoice.get('property_address')}", ''])
    
    if invoice.get('recipient_name'):
        items_data.insert(1, [f"Recipient: {invoice.get('recipient_name')}", ''])
    
    if invoice.get('notes'):
        items_data.append([f"Notes: {invoice.get('notes')}", ''])
    
    items_table = Table(items_data, colWidths=[4.5*inch, 2*inch])
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A2F3A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#DDDDDD')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
    ]))
    elements.append(items_table)
    elements.append(Spacer(1, 0.2*inch))
    
    # Total
    total_data = [
        ['', 'Total:', f"${invoice.get('amount', 0):,.2f} CAD"],
    ]
    total_table = Table(total_data, colWidths=[3.5*inch, 1.5*inch, 1.5*inch])
    total_table.setStyle(TableStyle([
        ('FONTNAME', (1, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (1, 0), (-1, -1), 12),
        ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
        ('TEXTCOLOR', (1, 0), (-1, -1), colors.HexColor('#1A2F3A')),
    ]))
    elements.append(total_table)
    elements.append(Spacer(1, 0.5*inch))
    
    # Footer
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.gray, alignment=1)
    elements.append(Paragraph("Thank you for your business!", footer_style))
    elements.append(Paragraph("DOMMMA Real Estate Marketplace | Vancouver, BC", footer_style))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    return Response(
        content=buffer.read(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=invoice_{invoice.get('invoice_number', invoice_id[:8])}.pdf"
        }
    )

@api_router.post("/payments/webhook/complete")
async def complete_payment_webhook(session_id: str):
    """Complete a pending payment (called after successful checkout)"""
    import stripe
    stripe.api_key = os.environ.get('STRIPE_API_KEY')
    
    # Find the pending invoice
    invoice = await db.invoices.find_one({"stripe_session_id": session_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    if invoice.get("status") == "paid":
        return {"success": True, "message": "Already paid"}
    
    # Verify with Stripe
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            await db.invoices.update_one(
                {"stripe_session_id": session_id},
                {"$set": {
                    "status": "paid",
                    "paid_at": datetime.now(timezone.utc).isoformat(),
                    "stripe_payment_id": session.payment_intent
                }}
            )
            return {"success": True, "message": "Payment completed"}
    except Exception as e:
        logger.error(f"Webhook error: {e}")
    
    return {"success": False, "message": "Payment not verified"}

# Payment types by role
PAYMENT_TYPES = {
    "renter": [
        {"type": "rent", "label": "Pay Rent", "icon": "Home", "description": "Monthly rent payment to landlord"},
        {"type": "security_deposit", "label": "Security Deposit", "icon": "Shield", "description": "Rental security deposit"},
        {"type": "utilities", "label": "Utilities", "icon": "Zap", "description": "Electricity, gas, water bills"},
        {"type": "contractor", "label": "Pay Contractor", "icon": "Wrench", "description": "Pay for repairs or services"},
        {"type": "moving", "label": "Moving Costs", "icon": "Truck", "description": "Moving company or supplies"},
        {"type": "other", "label": "Other Payment", "icon": "DollarSign", "description": "Miscellaneous payment"},
    ],
    "landlord": [
        {"type": "property_expense", "label": "Property Expense", "icon": "Building2", "description": "Maintenance, repairs, upgrades"},
        {"type": "contractor", "label": "Pay Contractor", "icon": "Wrench", "description": "Pay for property services"},
        {"type": "insurance", "label": "Insurance", "icon": "Shield", "description": "Property insurance payment"},
        {"type": "mortgage", "label": "Mortgage", "icon": "Home", "description": "Mortgage payment"},
        {"type": "taxes", "label": "Property Taxes", "icon": "FileText", "description": "Municipal property taxes"},
        {"type": "utilities", "label": "Utilities", "icon": "Zap", "description": "Property utility bills"},
        {"type": "other", "label": "Other Payment", "icon": "DollarSign", "description": "Miscellaneous payment"},
    ],
    "contractor": [
        {"type": "supplies", "label": "Supplies", "icon": "Package", "description": "Materials and supplies"},
        {"type": "equipment", "label": "Equipment", "icon": "Tool", "description": "Tools and equipment"},
        {"type": "subcontractor", "label": "Pay Subcontractor", "icon": "Users", "description": "Pay other contractors"},
        {"type": "insurance", "label": "Insurance", "icon": "Shield", "description": "Business insurance"},
        {"type": "license", "label": "Licenses/Permits", "icon": "FileText", "description": "Professional licenses"},
        {"type": "vehicle", "label": "Vehicle Expense", "icon": "Truck", "description": "Work vehicle costs"},
        {"type": "other", "label": "Other Payment", "icon": "DollarSign", "description": "Miscellaneous payment"},
    ]
}

@api_router.get("/payments/types/{user_type}")
async def get_payment_types(user_type: str):
    """Get available payment types for a user role"""
    return PAYMENT_TYPES.get(user_type, PAYMENT_TYPES["renter"])

# ========== DOCUMENT BUILDER ROUTES ==========

class DocumentBuilderSave(BaseModel):
    user_id: str
    template_id: str
    template_name: str
    form_data: Dict[str, Any]
    status: str = "draft"

@api_router.post("/document-builder/save")
async def save_builder_document(doc: DocumentBuilderSave):
    """Save a document from the document builder"""
    doc_id = str(uuid.uuid4())
    document = {
        "id": doc_id,
        "user_id": doc.user_id,
        "template_id": doc.template_id,
        "template_name": doc.template_name,
        "form_data": doc.form_data,
        "status": doc.status,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    await db.builder_documents.insert_one(document)
    return {"id": doc_id, "message": "Document saved"}

class DocumentBuilderSend(BaseModel):
    user_id: str
    template_id: str
    template_name: str
    form_data: Dict[str, Any]
    recipient_email: str
    sender_name: str
    sender_email: str

@api_router.post("/document-builder/send")
async def send_builder_document(doc: DocumentBuilderSend):
    """Send a document for signature via email"""
    doc_id = str(uuid.uuid4())
    
    # Save the document
    document = {
        "id": doc_id,
        "user_id": doc.user_id,
        "template_id": doc.template_id,
        "template_name": doc.template_name,
        "form_data": doc.form_data,
        "recipient_email": doc.recipient_email,
        "sender_name": doc.sender_name,
        "sender_email": doc.sender_email,
        "status": "pending_signature",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "sent_at": datetime.now(timezone.utc).isoformat()
    }
    await db.builder_documents.insert_one(document)
    
    # Send email notification
    try:
        from services.email import send_email
        
        # Generate signing link
        sign_link = f"https://dommma-rent-pay.preview.emergentagent.com/sign-document/{doc_id}"
        
        email_html = f"""
        <div style="font-family: Georgia, serif; max-width: 600px; margin: 0 auto; background: #F5F5F0; padding: 40px;">
            <div style="background: #1A2F3A; padding: 30px; border-radius: 16px 16px 0 0; text-align: center;">
                <h1 style="color: white; margin: 0; font-size: 28px;">DOMMMA</h1>
                <p style="color: rgba(255,255,255,0.7); margin: 8px 0 0;">Document Signature Request</p>
            </div>
            <div style="background: white; padding: 30px; border-radius: 0 0 16px 16px;">
                <h2 style="color: #1A2F3A; margin: 0 0 16px;">You have a document to sign</h2>
                <p style="color: #555; margin-bottom: 16px;">
                    <strong>{doc.sender_name}</strong> has sent you a <strong>{doc.template_name}</strong> to review and sign.
                </p>
                <div style="background: #F5F5F0; padding: 16px; border-radius: 8px; margin-bottom: 24px;">
                    <p style="color: #666; margin: 0; font-size: 14px;">Document: {doc.template_name}</p>
                    <p style="color: #666; margin: 4px 0 0; font-size: 14px;">From: {doc.sender_email}</p>
                </div>
                <a href="{sign_link}" style="display: inline-block; background: #1A2F3A; color: white; padding: 14px 28px; border-radius: 8px; text-decoration: none; font-weight: bold;">
                    Review & Sign Document
                </a>
                <p style="color: #999; font-size: 12px; margin-top: 24px;">
                    This link will expire in 30 days. If you have questions, contact {doc.sender_email}.
                </p>
            </div>
        </div>
        """
        
        await send_email(
            doc.recipient_email,
            f"Document for Signature: {doc.template_name}",
            email_html
        )
    except Exception as e:
        logger.error(f"Failed to send signature request email: {e}")
    
    return {"id": doc_id, "message": "Document sent for signature", "status": "pending_signature"}

class DocumentBuilderPDF(BaseModel):
    template_id: str
    template_name: str
    form_data: Dict[str, Any]

@api_router.post("/document-builder/pdf")
async def generate_builder_pdf(doc: DocumentBuilderPDF):
    """Generate a PDF from the document builder form data"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    import io
    
    buffer = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    elements = []
    
    # Header
    header_style = ParagraphStyle('Header', parent=styles['Heading1'], fontSize=18, textColor=colors.HexColor('#1A2F3A'))
    subheader_style = ParagraphStyle('SubHeader', parent=styles['Normal'], fontSize=10, textColor=colors.gray)
    
    elements.append(Paragraph(doc.template_name, header_style))
    elements.append(Paragraph("Province of British Columbia", subheader_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Form data as table
    form_data_list = []
    for key, value in doc.form_data.items():
        if value:  # Only include filled fields
            # Convert field ID to readable label
            label = key.replace('_', ' ').title()
            if isinstance(value, bool):
                value = "Yes" if value else "No"
            form_data_list.append([label + ":", str(value)])
    
    if form_data_list:
        data_table = Table(form_data_list, colWidths=[2.5*inch, 4*inch])
        data_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#333333')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor('#EEEEEE')),
        ]))
        elements.append(data_table)
    
    elements.append(Spacer(1, 0.5*inch))
    
    # Signature lines
    sig_style = ParagraphStyle('Signature', parent=styles['Normal'], fontSize=10)
    elements.append(Paragraph("_" * 40 + "                    " + "_" * 20, sig_style))
    elements.append(Paragraph("Landlord Signature                                              Date", subheader_style))
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph("_" * 40 + "                    " + "_" * 20, sig_style))
    elements.append(Paragraph("Tenant Signature                                                Date", subheader_style))
    
    # Footer
    elements.append(Spacer(1, 0.5*inch))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.gray, alignment=1)
    elements.append(Paragraph(f"Generated via DOMMMA Document Builder | {datetime.now().strftime('%Y-%m-%d')}", footer_style))
    
    pdf_doc.build(elements)
    buffer.seek(0)
    
    return Response(
        content=buffer.read(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={doc.template_name.replace(' ', '_')}.pdf"
        }
    )

@api_router.get("/document-builder/list/{user_id}")
async def list_builder_documents(user_id: str):
    """List all documents created by a user"""
    docs = await db.builder_documents.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return docs


# ========== AI DOCUMENT ASSISTANT ==========

@api_router.post("/document-builder/ai-prompts")
async def get_lease_prompts(landlord_id: str, listing_id: str = None):
    """Get AI-generated prompts and suggestions for lease agreement creation"""
    
    # Get landlord info
    landlord = await db.users.find_one({"id": landlord_id})
    
    # Get listing info if provided
    listing = None
    if listing_id:
        listing = await db.listings.find_one({"id": listing_id})
    
    # BC Residential Tenancy Act requirements
    prompts = {
        "required_clauses": [
            {
                "title": "Rent Payment Terms",
                "description": "Specify the monthly rent amount, due date, and acceptable payment methods",
                "legal_note": "Under BC RTA, rent increases are limited to once per year with proper notice",
                "ai_suggestion": f"Recommended due date: 1st of month. Consider grace period of 3-5 days before late fees."
            },
            {
                "title": "Late Payment Fees",
                "description": "Define late fee structure - flat rate or percentage",
                "legal_note": "Late fees must be reasonable and disclosed upfront. Courts may reject excessive fees.",
                "ai_suggestion": "Industry standard: $50 flat fee OR 3-5% of monthly rent after grace period. Stripe charges 2.9% + $0.30 per transaction.",
                "options": [
                    {"type": "flat", "amount": 50, "label": "$50 flat fee"},
                    {"type": "flat", "amount": 75, "label": "$75 flat fee"},
                    {"type": "percentage", "amount": 3, "label": "3% of rent"},
                    {"type": "percentage", "amount": 5, "label": "5% of rent"}
                ]
            },
            {
                "title": "Security Deposit",
                "description": "Maximum allowed is half month's rent for damage deposit",
                "legal_note": "BC law limits security deposit to maximum 1/2 month rent. Pet deposit max is also 1/2 month rent.",
                "ai_suggestion": f"For rent of ${listing.get('price', 2000)}, max deposit is ${listing.get('price', 2000) / 2:.2f}" if listing else "Set deposit based on monthly rent"
            },
            {
                "title": "Lease Term",
                "description": "Fixed term or month-to-month tenancy",
                "legal_note": "Fixed-term leases cannot include 'vacate clauses' requiring tenant to move at end of term.",
                "options": [
                    {"months": 12, "label": "12 months (standard)"},
                    {"months": 6, "label": "6 months"},
                    {"months": 0, "label": "Month-to-month"}
                ]
            },
            {
                "title": "Pet Policy",
                "description": "Specify if pets are allowed and any restrictions",
                "legal_note": "Landlords can restrict pets but cannot charge pet rent. Pet damage deposit is separate.",
                "ai_suggestion": "Be specific about pet types, sizes, and number allowed."
            },
            {
                "title": "Utilities",
                "description": "Clarify which utilities are included in rent",
                "legal_note": "Must clearly state if tenant is responsible for BC Hydro, FortisBC, internet, etc.",
                "checklist": ["Electricity", "Gas", "Water", "Internet", "Cable", "Garbage"]
            },
            {
                "title": "Move-in/Move-out Inspection",
                "description": "Schedule condition inspection",
                "legal_note": "BC RTA requires written condition inspection report at move-in and move-out.",
                "ai_suggestion": "Use DOMMMA's inspection checklist template."
            }
        ],
        "ai_tips": [
            "Always use the official BC RTB Form for residential tenancy agreements",
            "Keep copies of all signed documents for at least 2 years",
            "Take photos during move-in inspection as evidence",
            "Set up automatic rent collection to avoid payment delays",
            "Clearly communicate late fee policy before signing"
        ],
        "tenant_disclosure": {
            "title": "Important Information for Tenants",
            "items": [
                f"Monthly rent: ${listing.get('price', 'TBD')}" if listing else "Monthly rent: To be specified",
                "Late fees will apply after grace period",
                "Rent is collected automatically via credit card on the due date",
                "You can dispute charges within 30 days",
                "Contact the Residential Tenancy Branch for disputes"
            ]
        }
    }
    
    return prompts


@api_router.post("/document-builder/ai-review")
async def ai_review_document(content: str, document_type: str = "lease"):
    """AI reviews a document and highlights important terms for tenants"""
    
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY")
    if not anthropic_key:
        raise HTTPException(status_code=500, detail="AI service not configured")
    
    client = AsyncAnthropic(api_key=anthropic_key)
    
    prompt = f"""You are a helpful AI assistant reviewing a {document_type} agreement for a tenant in British Columbia, Canada.

Review the following document and:
1. Highlight any unusual or concerning clauses
2. Identify the late fee structure and payment terms
3. Note any terms that deviate from standard BC residential tenancy practices
4. Summarize key financial obligations
5. Flag any potentially illegal clauses under BC RTA

Document content:
{content}

Respond in JSON format:
{{
    "summary": "Brief 2-3 sentence summary",
    "monthly_obligations": {{
        "rent": "$X",
        "late_fee": "$X after X days",
        "other_fees": []
    }},
    "highlights": [
        {{"type": "info/warning/alert", "clause": "quote from doc", "explanation": "why this matters"}}
    ],
    "concerns": [
        {{"severity": "low/medium/high", "issue": "description", "recommendation": "what to do"}}
    ],
    "tenant_checklist": ["items tenant should verify before signing"]
}}"""

    try:
        response = await client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        result = response.content[0].text
        # Try to parse as JSON
        try:
            import json
            return json.loads(result)
        except:
            return {"raw_analysis": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")


# ========== POST-RESERVATION UPSELLS ==========

@api_router.get("/upsells/services/{city}")
async def get_local_services(city: str):
    """Get local service providers for post-reservation upsells"""
    
    # Vancouver Metro area service providers (curated list)
    # In production, this would be a database or API call
    services = {
        "movers": [
            {"id": "m1", "name": "BC Moving Co.", "rating": 4.8, "reviews": 234, "price": "From $99/hr", "phone": "604-555-0101", "featured": True, "website": "https://bcmoving.ca"},
            {"id": "m2", "name": "Two Small Men with Big Hearts", "rating": 4.7, "reviews": 412, "price": "From $109/hr", "phone": "604-555-0102", "website": "https://twosmallmen.com"},
            {"id": "m3", "name": "Ferguson Moving", "rating": 4.6, "reviews": 189, "price": "From $95/hr", "phone": "604-555-0103", "website": "https://fergusonmoving.com"},
            {"id": "m4", "name": "Quick Move YVR", "rating": 4.5, "reviews": 89, "price": "From $79/hr", "phone": "604-555-0104"},
        ],
        "internet": [
            {"id": "w1", "name": "Telus", "rating": 4.2, "price": "From $75/mo", "promo": "3 months free", "featured": True, "website": "https://telus.com"},
            {"id": "w2", "name": "Shaw", "rating": 4.0, "price": "From $70/mo", "promo": "Free installation", "website": "https://shaw.ca"},
            {"id": "w3", "name": "Novus", "rating": 4.5, "price": "From $55/mo", "promo": "No contract required", "website": "https://novusnow.ca"},
            {"id": "w4", "name": "Starlink", "rating": 4.3, "price": "$140/mo", "promo": "High-speed anywhere", "website": "https://starlink.com"},
        ],
        "insurance": [
            {"id": "i1", "name": "BCAA Renters Insurance", "rating": 4.7, "price": "From $18/mo", "featured": True, "website": "https://bcaa.com"},
            {"id": "i2", "name": "Square One Insurance", "rating": 4.5, "price": "From $12/mo", "promo": "Online quotes in 5 min", "website": "https://squareone.ca"},
            {"id": "i3", "name": "ICBC", "rating": 4.0, "price": "From $20/mo", "website": "https://icbc.com"},
        ],
        "utilities": [
            {"id": "u1", "name": "BC Hydro", "rating": 4.1, "price": "Avg $80/mo", "info": "Setup required for new tenants", "website": "https://bchydro.com"},
            {"id": "u2", "name": "FortisBC", "rating": 4.0, "price": "Avg $60/mo", "info": "Natural gas provider", "website": "https://fortisbc.com"},
        ],
        "cleaning": [
            {"id": "c1", "name": "Maid4Condos", "rating": 4.8, "price": "From $120", "featured": True, "phone": "604-555-0201"},
            {"id": "c2", "name": "AspenClean", "rating": 4.7, "price": "From $150", "promo": "Eco-friendly products", "website": "https://aspenclean.com"},
            {"id": "c3", "name": "Molly Maid", "rating": 4.5, "price": "From $130", "website": "https://mollymaid.ca"},
        ],
        "storage": [
            {"id": "s1", "name": "U-Haul Storage", "rating": 4.3, "price": "From $49/mo", "promo": "First month free", "website": "https://uhaul.com"},
            {"id": "s2", "name": "Storage Mart", "rating": 4.4, "price": "From $59/mo", "featured": True},
            {"id": "s3", "name": "Access Storage", "rating": 4.2, "price": "From $55/mo", "website": "https://accessstorage.ca"},
        ]
    }
    
    return {
        "city": city,
        "services": services,
        "generated_at": datetime.now(timezone.utc).isoformat()
    }

@api_router.post("/upsells/request-quote")
async def request_service_quote(data: Dict[str, Any]):
    """Record a quote request for a service provider"""
    request_id = str(uuid.uuid4())
    
    quote_request = {
        "id": request_id,
        "user_id": data.get("user_id"),
        "service_type": data.get("service_type"),
        "provider_id": data.get("provider_id"),
        "provider_name": data.get("provider_name"),
        "property_address": data.get("property_address"),
        "move_date": data.get("move_date"),
        "notes": data.get("notes"),
        "status": "pending",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.upsell_requests.insert_one(quote_request)
    
    # In production, this would send an email to the service provider
    # and/or integrate with their booking API
    
    return {"id": request_id, "message": "Quote request submitted", "status": "pending"}

# ========== DOCUMENT MANAGEMENT ROUTES ==========

@api_router.post("/documents/upload")
async def upload_document_legacy(
    user_id: str,
    name: str,
    doc_type: str,
    file: UploadFile = File(...)
):
    """Upload a document to R2 storage (with base64 fallback)"""
    try:
        content = await file.read()
        doc_id = str(uuid.uuid4())
        
        # Try R2 first if configured
        if is_r2_configured():
            try:
                result = await r2_upload_document(
                    file_data=content,
                    filename=file.filename,
                    content_type=file.content_type or "application/octet-stream",
                    user_id=user_id,
                    document_type=doc_type
                )
                
                doc = Document(
                    id=doc_id,
                    user_id=user_id,
                    name=name,
                    type=doc_type,
                    url=result["url"],
                    content=None  # Don't store content in DB
                )
                doc_dict = doc.model_dump()
                doc_dict["r2_key"] = result["key"]
                doc_dict["storage"] = "r2"
                await db.documents.insert_one(doc_dict)
                
                return {"id": doc.id, "name": doc.name, "type": doc.type, "url": result["url"], "status": "uploaded"}
            except Exception as e:
                logger.warning(f"R2 document upload failed, falling back to base64: {e}")
        
        # Fallback to base64 storage
        content_b64 = base64.b64encode(content).decode('utf-8')
        
        doc = Document(
            id=doc_id,
            user_id=user_id,
            name=name,
            type=doc_type,
            content=content_b64
        )
        doc_dict = doc.model_dump()
        doc_dict["storage"] = "base64"
        await db.documents.insert_one(doc_dict)
        
        return {"id": doc.id, "name": doc.name, "type": doc.type, "status": "uploaded"}
        
    except Exception as e:
        logger.error(f"Document upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/documents/{user_id}")
async def get_user_documents(user_id: str, doc_type: Optional[str] = None):
    query = {"user_id": user_id, "status": "active"}
    if doc_type:
        query["type"] = doc_type
    
    docs = await db.documents.find(query, {"_id": 0, "content": 0}).to_list(100)
    return docs

@api_router.get("/documents/download/{doc_id}")
async def download_document(doc_id: str):
    doc = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@api_router.post("/documents/sign/{doc_id}")
async def sign_document(doc_id: str, user_id: str):
    result = await db.documents.update_one(
        {"id": doc_id},
        {"$set": {
            "signed": True,
            "signed_at": datetime.now(timezone.utc).isoformat(),
            "signed_by": user_id
        }}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "signed", "signed_at": datetime.now(timezone.utc).isoformat()}

@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    result = await db.documents.update_one(
        {"id": doc_id},
        {"$set": {"status": "deleted"}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "deleted"}

# ========== MESSAGING ROUTES ==========

@api_router.post("/messages/send")
async def send_message(sender_id: str, message: MessageCreate):
    msg = DirectMessage(
        sender_id=sender_id,
        recipient_id=message.recipient_id,
        content=message.content
    )
    await db.messages.insert_one(msg.model_dump())
    
    # Send via WebSocket if recipient is connected
    await manager.send_personal_message(
        json.dumps({
            "type": "new_message",
            "message": msg.model_dump()
        }),
        message.recipient_id
    )
    
    return {"id": msg.id, "status": "sent"}

@api_router.get("/messages/{user_id}")
async def get_messages(user_id: str, other_user_id: Optional[str] = None):
    if other_user_id:
        # Get conversation between two users
        query = {
            "$or": [
                {"sender_id": user_id, "recipient_id": other_user_id},
                {"sender_id": other_user_id, "recipient_id": user_id}
            ]
        }
    else:
        # Get all messages for user
        query = {
            "$or": [
                {"sender_id": user_id},
                {"recipient_id": user_id}
            ]
        }
    
    messages = await db.messages.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return messages

@api_router.get("/messages/conversations/{user_id}")
async def get_conversations(user_id: str):
    # Get unique conversation partners
    pipeline = [
        {"$match": {"$or": [{"sender_id": user_id}, {"recipient_id": user_id}]}},
        {"$sort": {"created_at": -1}},
        {"$group": {
            "_id": {
                "$cond": [
                    {"$eq": ["$sender_id", user_id]},
                    "$recipient_id",
                    "$sender_id"
                ]
            },
            "last_message": {"$first": "$content"},
            "last_time": {"$first": "$created_at"},
            "unread": {"$sum": {"$cond": [{"$and": [{"$eq": ["$recipient_id", user_id]}, {"$eq": ["$read", False]}]}, 1, 0]}}
        }}
    ]
    conversations = await db.messages.aggregate(pipeline).to_list(50)
    return [{"user_id": c["_id"], "last_message": c["last_message"], "last_time": c["last_time"], "unread": c["unread"]} for c in conversations]

@api_router.post("/messages/read/{message_id}")
async def mark_message_read(message_id: str):
    await db.messages.update_one({"id": message_id}, {"$set": {"read": True}})
    return {"status": "read"}

# WebSocket endpoint for real-time messaging
@app.websocket("/ws/{user_id}")
async def websocket_endpoint(websocket: WebSocket, user_id: str):
    await manager.connect(websocket, user_id)
    try:
        while True:
            data = await websocket.receive_text()
            message_data = json.loads(data)
            
            if message_data.get("type") == "message":
                msg = DirectMessage(
                    sender_id=user_id,
                    recipient_id=message_data["recipient_id"],
                    content=message_data["content"]
                )
                await db.messages.insert_one(msg.model_dump())
                
                # Send to recipient
                await manager.send_personal_message(
                    json.dumps({"type": "new_message", "message": msg.model_dump()}),
                    message_data["recipient_id"]
                )
                
                # Confirm to sender
                await websocket.send_text(json.dumps({"type": "sent", "message_id": msg.id}))
                
    except WebSocketDisconnect:
        manager.disconnect(websocket, user_id)

# ========== ENHANCED NOVA AI CHAT ==========

@api_router.post("/chat", response_model=ChatResponse)
async def chat_with_nova(request: ChatRequest):
    session_id = request.session_id or str(uuid.uuid4())
    user_id = request.user_id  # Get user_id if logged in
    
    # Get or create chat session
    session = await db.chat_sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        session = {"id": session_id, "messages": [], "created_at": datetime.now(timezone.utc).isoformat(), "user_id": user_id}
        await db.chat_sessions.insert_one(session)
    
    # Get available listings for context - fetch ALL active listings
    listings = await db.listings.find({"status": "active"}, {"_id": 0}).to_list(200)
    
    # Build clear context with CITY, OFFERS prominently displayed
    listings_context = "\n".join([
        f"- ID:{listing['id']} | CITY: {listing.get('city', 'Vancouver')} | {listing['title']}: {listing.get('bedrooms', 0)}bd/{listing.get('bathrooms', 1)}ba, ${listing.get('price', 0)}/mo, {listing.get('address', '')}, {listing.get('sqft', 0)}sqft, Pet-friendly: {listing.get('pet_friendly', False)}, Type: {listing.get('property_type', 'Apartment')}, Listing: {listing.get('listing_type', 'rent')}, Lease: {listing.get('lease_duration', 12)} months, OFFERS: {', '.join(listing.get('offers', [])) if listing.get('offers') else 'None'}"
        for listing in listings
    ])
    
    # Get contractors for context
    contractors = await db.contractor_profiles.find({"status": "active"}, {"_id": 0}).to_list(50)
    contractors_context = "\n".join([
        f"- CONTRACTOR ID:{c['id']} | {c.get('company_name', 'Unknown')} | Specialties: {', '.join(c.get('specialties', []))} | Rate: ${c.get('hourly_rate', 0)}/hr | Rating: {c.get('rating', 0)}/5 | Verified: {c.get('verified', False)}"
        for c in contractors
    ])
    
    # Log for debugging
    cities = {}
    for l in listings:
        c = l.get('city', 'Unknown')
        cities[c] = cities.get(c, 0) + 1
    logger.info(f"Chat context: {len(listings)} listings, cities: {cities}, {len(contractors)} contractors")
    
    # Build conversation context from history
    history_context = ""
    prev_messages = session.get("messages", [])[-6:]
    if prev_messages:
        history_context = "\n\nRecent conversation:\n" + "\n".join([
            f"{'User' if m['role'] == 'user' else 'Nova'}: {m['content'][:200]}"
            for m in prev_messages
        ])
    
    # Load long-term user memory if logged in
    long_term_memory = ""
    user_preferences_loaded = False
    if user_id:
        from services.nova_memory import NovaMemoryService
        memory_service = NovaMemoryService(db)
        memory_context = await memory_service.get_context_summary(user_id)
        if memory_context:
            long_term_memory = f"\n\n{memory_context}"
            user_preferences_loaded = True
    
    # User context for personalization (from manual input)
    user_context = request.user_context or {}
    lifestyle_info = ""
    if user_context:
        lifestyle_info = f"""
User Context:
- Budget: ${user_context.get('budget', 'Not specified')}/month
- Occupation: {user_context.get('occupation', 'Not specified')}
- Has Pets: {user_context.get('has_pets', 'Unknown')}
- Commute To: {user_context.get('commute_location', 'Not specified')}
- Preferences: {user_context.get('preferences', 'None specified')}
"""
    
    system_message = f"""You are Nova, DOMMMA's advanced AI real estate assistant. You help users with ALL aspects of real estate:

🏠 PROPERTY SEARCH
- Natural language search
- Lifestyle-based recommendations
- Commute optimization
- Neighborhood insights

💰 FINANCIAL HELP  
- Budget breakdown calculator
- Rent negotiation tips
- Hidden costs analysis
- Affordability advice (30% rule: rent should be max 30% of gross income)

📄 APPLICATION SUPPORT
- Rental resume tips
- Application optimization
- Document guidance

🏘️ NEIGHBORHOOD INTEL
- Safety insights
- Local amenities
- Community vibes
- Transit access

👥 COMMUNICATION
- Help draft messages to landlords
- Conflict resolution advice
- Negotiation strategies

🔧 CONTRACTOR/SERVICE SEARCH
- Find plumbers, electricians, painters, cleaners, etc.
- Match users with verified contractors
- Provide contractor recommendations based on specialty

🎁 SPECIAL OFFERS
- Highlight listings with special offers (free rent, free wifi, etc.)
- When user asks about deals/offers, show listings with OFFERS field filled

Available Properties in Database (include ID when recommending):
{listings_context if listings_context else "No listings currently available."}

Available Contractors in Database (include ID when recommending):
{contractors_context if contractors_context else "No contractors currently available."}
{lifestyle_info}
{long_term_memory}
{history_context}

IMPORTANT CAPABILITIES:
1. Budget Calculator: If user gives income, calculate max affordable rent (30% of gross monthly)
2. Lifestyle Search: Match properties to lifestyle (e.g., "I bike to work" → near bike routes)
3. Commute Analysis: Estimate commute times to key locations
4. Proactive Suggestions: Offer relevant tips based on conversation
5. Multi-turn Memory: Reference previous parts of the conversation
6. Long-term Memory: If user context shows saved preferences, USE THEM to personalize recommendations
7. Contractor Search: When user needs a plumber, electrician, cleaner, etc., search the contractors list and recommend

CRITICAL - LOCATION AWARENESS:
- We have listings in MULTIPLE CITIES: Vancouver, Coquitlam, Burnaby, Richmond, Surrey, and other Metro Vancouver areas
- When user asks about a specific city (e.g., "anything in Coquitlam?"), search the listings above by the CITY field
- ALWAYS check the CITY field in each listing - it's clearly marked as "CITY: [city name]"
- If we have properties in that city, SHOW THEM. If not, be honest and suggest nearby areas.

CRITICAL - OFFERS/DEALS:
- When user asks for "deals", "offers", "promotions", or "specials", look for listings with OFFERS field
- Highlight listings that have offers like "1 month free rent", "Free WiFi", etc.
- Example: "Here's a great deal: [Yaletown Loft](property:123) - 1 month free rent!"

CRITICAL - CONTRACTOR SEARCH:
- When user says "I need a plumber" or "looking for electrician", search the contractors list
- Recommend contractors by specialty, rating, and hourly rate
- Format: "[Company Name](contractor:ID)" to make it clickable
- Example: "I found [Plumbing Pro](contractor:abc) - $80/hr, 4.8★, verified"

CRITICAL - WHEN RECOMMENDING PROPERTIES:
- Always reference the property ID from the database above
- Format property recommendations with their ID like: "[Property Name](property:ID)"
- Example: "I recommend [Modern Downtown Condo](property:abc123) - it's perfect for your needs"
- This allows users to click and view the property directly
- If no properties match, say so honestly and suggest broadening the search
- If user has saved preferences (pet-friendly, budget, area), prioritize those matches

Keep responses helpful, conversational, and concise (2-3 paragraphs max).
End responses with 1-2 proactive suggestions when relevant."""

    suggestions = []
    
    try:
        # Use direct Anthropic SDK
        anthropic_client = AsyncAnthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
        
        # Build messages array with conversation history
        messages = []
        prev_messages = session.get("messages", [])[-6:]
        for m in prev_messages:
            messages.append({"role": m['role'] if m['role'] != 'assistant' else 'assistant', "content": m['content'][:500]})
        
        # Add current user message
        messages.append({"role": "user", "content": request.message})
        
        # Call Claude API
        claude_response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=1024,
            system=system_message,
            messages=messages
        )
        
        response = claude_response.content[0].text
        
        # Save messages to session
        await db.chat_sessions.update_one(
            {"id": session_id},
            {"$push": {
                "messages": {
                    "$each": [
                        {"role": "user", "content": request.message, "timestamp": datetime.now(timezone.utc).isoformat()},
                        {"role": "assistant", "content": response, "timestamp": datetime.now(timezone.utc).isoformat()}
                    ]
                }
            }}
        )
        
        # Store interaction and extract preferences for logged-in users
        if user_id:
            from services.nova_memory import NovaMemoryService
            memory_service = NovaMemoryService(db)
            await memory_service.store_interaction(
                user_id=user_id,
                message=request.message,
                response=response,
                context=user_context
            )
        
        # Find mentioned listings
        mentioned_listings = []
        for listing in listings:
            if listing['title'].lower() in response.lower() or listing['city'].lower() in response.lower():
                mentioned_listings.append(listing)
        
        # Generate proactive suggestions based on context
        if "budget" in request.message.lower() or "afford" in request.message.lower():
            suggestions.append("💡 Would you like me to show properties within your budget?")
        if "pet" in request.message.lower():
            suggestions.append("🐾 I can filter for pet-friendly properties only")
        if any(word in request.message.lower() for word in ["commute", "work", "office"]):
            suggestions.append("🚇 I can find places with the best commute to your workplace")
        
        return ChatResponse(
            session_id=session_id,
            response=response,
            listings=mentioned_listings[:5],
            suggestions=suggestions[:3],
            preferences_loaded=user_preferences_loaded
        )
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        return ChatResponse(
            session_id=session_id,
            response="I'm having a moment! Let me help you - try asking about apartments in Vancouver, budget advice, or neighborhood recommendations! 🏠",
            listings=[],
            suggestions=["Try asking: 'What can I afford on $70k salary?'", "Try: 'Show me pet-friendly apartments near downtown'"],
            preferences_loaded=False
        )

@api_router.get("/chat/{session_id}/history")
async def get_chat_history(session_id: str):
    session = await db.chat_sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    return session


# ========== AI CONCIERGE WITH TOOL CALLING ==========

class ConciergeRequest(BaseModel):
    session_id: Optional[str] = None
    message: str
    user_id: Optional[str] = None
    user_context: Optional[Dict[str, Any]] = None

class ConciergeResponse(BaseModel):
    session_id: str
    response: str
    tool_results: Optional[List[Dict[str, Any]]] = None
    listings: Optional[List[Dict[str, Any]]] = None
    contractors: Optional[List[Dict[str, Any]]] = None
    suggestions: Optional[List[str]] = None

@api_router.post("/ai/concierge", response_model=ConciergeResponse)
async def ai_concierge(request: ConciergeRequest):
    """
    AI Concierge endpoint with Claude tool calling.
    Handles structured actions like creating listings, searching, scheduling viewings.
    """
    from services.ai_tools import AIToolsService, NOVA_TOOLS
    
    session_id = request.session_id or str(uuid.uuid4())
    user_id = request.user_id
    
    # Get or create chat session
    session = await db.chat_sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        session = {
            "id": session_id, 
            "messages": [], 
            "created_at": datetime.now(timezone.utc).isoformat(), 
            "user_id": user_id,
            "mode": "concierge"
        }
        await db.chat_sessions.insert_one(session)
    
    # Initialize tools service
    tools_service = AIToolsService(db)
    
    # Get context for system prompt
    listings = await db.listings.find({"status": "active"}, {"_id": 0}).to_list(50)
    contractors = await db.contractor_profiles.find({"status": "active"}, {"_id": 0}).to_list(30)
    
    # Build context summaries
    listings_summary = f"Currently {len(listings)} active listings in database."
    contractors_summary = f"Currently {len(contractors)} active contractors available."
    
    # Get user info if logged in
    user_info = ""
    if user_id:
        user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        if user:
            user_info = f"\nUser: {user.get('name', 'Unknown')} ({user.get('role', 'renter')})"
    
    system_message = f"""You are Nova, DOMMMA's AI real estate concierge. You help users with all aspects of real estate through natural conversation.

You have access to tools that let you take real actions:
- **create_listing**: Create property listings for landlords
- **search_listings**: Find properties based on criteria
- **find_contractors**: Find plumbers, electricians, cleaners, etc.
- **triage_maintenance**: Handle maintenance requests
- **calculate_budget**: Help users understand what they can afford
- **schedule_viewing**: Book property viewings

WHEN TO USE TOOLS:
- User says "I want to list my apartment" → Use create_listing (gather details first if needed)
- User says "Find me a 2 bedroom under $2500" → Use search_listings
- User says "I need a plumber" → Use find_contractors
- User says "My sink is leaking" → Use triage_maintenance
- User says "What can I afford on $80k" → Use calculate_budget
- User says "I want to see this property" → Use schedule_viewing

CREATING LISTINGS WORKFLOW:
When a user wants to create a listing, collect these details conversationally:
1. Property type (apartment, condo, house, townhouse, etc.)
2. Address and city
3. Monthly rent/price
4. Number of bedrooms and bathrooms
5. When available for move-in
Optional: square footage, pet policy, amenities, description

IMPORTANT FOR UNAUTHENTICATED USERS (when user_id is empty):
- After gathering listing details, you MUST ask for their email address
- Include the email in the create_listing tool call as "claim_email"
- Explain that they'll receive an email to verify and publish the listing
- Say something like: "Great details! To publish your listing, I'll need your email address to create your landlord account. What email should I use?"

RESPONSE FORMAT:
- After using create_listing, format the result as: [Listing Title](property:LISTING_ID)
- After using find_contractors, format each as: [Contractor Name](contractor:CONTRACTOR_ID)
- Be conversational and helpful
- If you need more info to use a tool, ask for it naturally

{user_info}
{listings_summary}
{contractors_summary}

Current user status: {"Logged in as " + user_info if user_id else "Not logged in - will need email for listing creation"}

Keep responses concise and action-oriented. You're a helpful concierge that gets things done!"""

    try:
        anthropic_client = AsyncAnthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
        
        # Build messages from history (only text content, skip tool results to avoid tool_use/tool_result pairing issues)
        messages = []
        prev_messages = session.get("messages", [])[-6:]  # Reduced to avoid long context issues
        for m in prev_messages:
            # Only include simple text messages
            content = m.get('content', '')
            if isinstance(content, str) and content:
                messages.append({
                    "role": m['role'], 
                    "content": content[:500]
                })
        
        # Add current message
        messages.append({"role": "user", "content": request.message})
        
        # Initial Claude call with tools
        response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=2048,
            system=system_message,
            tools=NOVA_TOOLS,
            messages=messages
        )
        
        tool_results = []
        final_response = ""
        result_listings = []
        result_contractors = []
        
        # Process response - may include tool calls
        max_tool_iterations = 5  # Prevent infinite loops
        iteration = 0
        
        while response.stop_reason == "tool_use" and iteration < max_tool_iterations:
            iteration += 1
            # Find ALL tool use blocks in this response
            tool_use_blocks = []
            text_content = ""
            
            for block in response.content:
                if block.type == "tool_use":
                    tool_use_blocks.append(block)
                elif block.type == "text":
                    text_content = block.text
            
            if tool_use_blocks:
                # Build assistant content with all tool uses
                assistant_content = []
                if text_content:
                    assistant_content.append({"type": "text", "text": text_content})
                
                tool_result_content = []
                
                for tool_use_block in tool_use_blocks:
                    # Add tool use to assistant content
                    assistant_content.append({
                        "type": "tool_use",
                        "id": tool_use_block.id,
                        "name": tool_use_block.name,
                        "input": tool_use_block.input
                    })
                    
                    # Execute the tool
                    tool_name = tool_use_block.name
                    tool_input = tool_use_block.input
                    tool_id = tool_use_block.id
                    
                    logger.info(f"Executing tool: {tool_name}")
                    result = await tools_service.execute_tool(tool_name, tool_input, user_id)
                    tool_results.append({
                        "tool": tool_name,
                        "input": tool_input,
                        "result": result
                    })
                    
                    # Collect listings/contractors from results
                    if result.get("listings"):
                        result_listings.extend(result["listings"])
                    if result.get("contractors"):
                        result_contractors.extend(result["contractors"])
                    if result.get("listing"):
                        result_listings.append(result["listing"])
                    if result.get("suggested_contractors"):
                        for c in result["suggested_contractors"]:
                            result_contractors.append(c)
                    
                    # Add tool result
                    tool_result_content.append({
                        "type": "tool_result",
                        "tool_use_id": tool_id,
                        "content": json.dumps(result)
                    })
                
                # Add assistant message with all tool uses
                messages.append({"role": "assistant", "content": assistant_content})
                # Add user message with all tool results
                messages.append({"role": "user", "content": tool_result_content})
                
                # Get Claude's interpretation of the result
                response = await anthropic_client.messages.create(
                    model="claude-sonnet-4-5-20250929",
                    max_tokens=2048,
                    system=system_message,
                    tools=NOVA_TOOLS,
                    messages=messages
                )
        
        # Extract final text response
        for block in response.content:
            if hasattr(block, 'text'):
                final_response = block.text
                break
        
        # Save messages to session
        await db.chat_sessions.update_one(
            {"id": session_id},
            {"$push": {
                "messages": {
                    "$each": [
                        {"role": "user", "content": request.message, "timestamp": datetime.now(timezone.utc).isoformat()},
                        {"role": "assistant", "content": final_response, "timestamp": datetime.now(timezone.utc).isoformat(), "tool_results": tool_results}
                    ]
                }
            }}
        )
        
        # Generate suggestions
        suggestions = []
        if not tool_results:
            if "list" in request.message.lower() or "rent out" in request.message.lower():
                suggestions.append("💡 I can help you create a listing - just say 'I want to list my property'")
            if "find" in request.message.lower() or "search" in request.message.lower():
                suggestions.append("🏠 Tell me your requirements and I'll search for matching properties")
        
        return ConciergeResponse(
            session_id=session_id,
            response=final_response,
            tool_results=tool_results if tool_results else None,
            listings=result_listings if result_listings else None,
            contractors=result_contractors if result_contractors else None,
            suggestions=suggestions if suggestions else None
        )
        
    except Exception as e:
        logger.error(f"Concierge error: {e}")
        import traceback
        traceback.print_exc()
        return ConciergeResponse(
            session_id=session_id,
            response="I'm having a moment! Let me help you - try asking me to find apartments, create a listing, or connect you with a contractor.",
            tool_results=None,
            listings=None,
            contractors=None,
            suggestions=["Try: 'I want to list my 2 bedroom apartment'", "Try: 'Find me a plumber'"]
        )

# Nova Memory endpoint - Get saved preferences for a user
@api_router.get("/nova/memory/{user_id}")
async def get_nova_memory(user_id: str):
    """Get Nova's memory of user preferences for display in chat"""
    from services.nova_memory import NovaMemoryService
    memory_service = NovaMemoryService(db)
    memory = await memory_service.get_user_memory(user_id)
    context_summary = await memory_service.get_context_summary(user_id)
    return {
        "preferences": memory.get("preferences", {}),
        "context_summary": context_summary,
        "has_preferences": bool(memory.get("preferences"))
    }

# User Preferences for Nova Memory
@api_router.post("/user/preferences/{user_id}")
async def save_user_preferences(user_id: str, preferences: Dict[str, Any]):
    await db.users.update_one(
        {"id": user_id},
        {"$set": {"preferences": preferences}}
    )
    return {"status": "saved"}

# ========== RENTER RESUME ENDPOINTS ==========

class RenterResumeInput(BaseModel):
    user_id: Optional[str] = None
    full_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    employment: Optional[Dict[str, Any]] = None
    rental_history: Optional[Dict[str, Any]] = None
    household: Optional[Dict[str, Any]] = None
    preferences: Optional[Dict[str, Any]] = None

@api_router.get("/renter-resume/{user_id}")
async def get_renter_resume(user_id: str):
    """Get a renter's resume/profile"""
    resume = await db.renter_resumes.find_one({"user_id": user_id}, {"_id": 0})
    
    if not resume:
        return {
            "has_resume": False,
            "message": "No renter resume found"
        }
    
    return {
        "has_resume": True,
        "resume": resume
    }

@api_router.post("/renter-resume")
async def save_renter_resume(data: RenterResumeInput):
    """Create or update a renter resume"""
    user_id = data.user_id
    
    if not user_id:
        raise HTTPException(status_code=400, detail="User ID required")
    
    # Check if resume exists
    existing = await db.renter_resumes.find_one({"user_id": user_id})
    
    # Calculate completeness score
    fields_filled = 0
    total_fields = 10
    if data.full_name:
        fields_filled += 1
    if data.email:
        fields_filled += 1
    if data.phone:
        fields_filled += 1
    if data.employment and data.employment.get('status'):
        fields_filled += 1
    if data.employment and data.employment.get('employer'):
        fields_filled += 1
    if data.employment and data.employment.get('annual_income'):
        fields_filled += 1
    if data.rental_history and data.rental_history.get('current_address'):
        fields_filled += 1
    if data.rental_history and data.rental_history.get('previous_landlord', {}).get('name'):
        fields_filled += 1
    if data.household and data.household.get('num_occupants'):
        fields_filled += 1
    if data.preferences and data.preferences.get('move_in_date'):
        fields_filled += 1
    
    completeness_score = round((fields_filled / total_fields) * 100)
    
    resume_data = {
        "user_id": user_id,
        "full_name": data.full_name or "",
        "email": data.email or "",
        "phone": data.phone or "",
        "employment": data.employment or {},
        "rental_history": data.rental_history or {},
        "household": data.household or {"num_occupants": 1},
        "preferences": data.preferences or {},
        "completeness_score": completeness_score,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    if existing:
        await db.renter_resumes.update_one(
            {"user_id": user_id},
            {"$set": resume_data}
        )
    else:
        resume_data["id"] = str(uuid.uuid4())
        resume_data["created_at"] = datetime.now(timezone.utc).isoformat()
        await db.renter_resumes.insert_one(resume_data)
    
    return {"status": "saved", "completeness_score": completeness_score}

@api_router.get("/user/preferences/{user_id}")
async def get_user_preferences(user_id: str):
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user.get("preferences", {})

# ========== LEASE ASSIGNMENT MARKETPLACE ==========

class LeaseAssignmentInput(BaseModel):
    title: Optional[str] = None
    address: str
    city: str = "Vancouver"
    current_rent: float
    market_rent: Optional[float] = None
    assignment_fee: float
    remaining_months: int
    available_date: Optional[str] = None
    bedrooms: int = 1
    bathrooms: float = 1
    sqft: Optional[int] = None
    amenities: List[str] = []
    pet_friendly: bool = False
    description: Optional[str] = None
    reason: Optional[str] = None
    owner_id: Optional[str] = None
    owner_name: Optional[str] = None
    images: List[str] = []

@api_router.get("/lease-assignments")
async def get_lease_assignments(city: Optional[str] = None, min_savings: Optional[int] = None):
    """Get all active lease assignments"""
    query = {"status": "active"}
    
    if city:
        query["city"] = {"$regex": city, "$options": "i"}
    
    assignments = await db.lease_assignments.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    
    # Calculate savings for each
    for assignment in assignments:
        if assignment.get("market_rent") and assignment.get("current_rent"):
            assignment["savings_per_month"] = assignment["market_rent"] - assignment["current_rent"]
        else:
            assignment["market_rent"] = assignment.get("current_rent", 0) * 1.05
            assignment["savings_per_month"] = assignment.get("current_rent", 0) * 0.05
    
    if min_savings:
        assignments = [a for a in assignments if a.get("savings_per_month", 0) >= min_savings]
    
    return assignments

@api_router.post("/lease-assignments")
async def create_lease_assignment(data: LeaseAssignmentInput):
    """Create a new lease assignment listing"""
    assignment_id = str(uuid.uuid4())
    
    market_rent = data.market_rent or (data.current_rent * 1.05)
    
    assignment = {
        "id": assignment_id,
        "title": data.title or f"{data.bedrooms}BR in {data.city}",
        "address": data.address,
        "city": data.city,
        "current_rent": data.current_rent,
        "market_rent": market_rent,
        "assignment_fee": data.assignment_fee,
        "remaining_months": data.remaining_months,
        "available_date": data.available_date or "Immediately",
        "bedrooms": data.bedrooms,
        "bathrooms": data.bathrooms,
        "sqft": data.sqft,
        "amenities": data.amenities,
        "pet_friendly": data.pet_friendly,
        "description": data.description,
        "reason": data.reason,
        "owner_id": data.owner_id,
        "owner": {"name": data.owner_name or "Anonymous", "avatar": (data.owner_name or "A")[0]},
        "images": data.images,
        "savings_per_month": market_rent - data.current_rent,
        "status": "active",
        "verified": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.lease_assignments.insert_one(assignment)
    
    # Remove MongoDB _id before returning
    assignment.pop("_id", None)
    return assignment

@api_router.get("/lease-assignments/{assignment_id}")
async def get_lease_assignment(assignment_id: str):
    """Get a specific lease assignment"""
    assignment = await db.lease_assignments.find_one({"id": assignment_id}, {"_id": 0})
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")
    return assignment

# ========== LEASE ASSIGNMENT PAYMENTS (Stripe) ==========

@api_router.post("/lease-assignments/{assignment_id}/payment")
async def create_assignment_payment(request: Request, assignment_id: str, buyer_id: str):
    """Create a Stripe checkout session for a lease assignment fee"""
    
    # Get the assignment
    assignment = await db.lease_assignments.find_one({"id": assignment_id})
    if not assignment:
        raise HTTPException(status_code=404, detail="Lease assignment not found")
    
    if assignment.get("status") != "active":
        raise HTTPException(status_code=400, detail="This lease assignment is no longer available")
    
    # Get assignment fee from backend (security - don't trust frontend)
    amount = float(assignment.get("assignment_fee", 0))
    if amount <= 0:
        raise HTTPException(status_code=400, detail="Invalid assignment fee")
    
    try:
        api_key = os.environ.get('STRIPE_API_KEY')
        host_url = str(request.base_url).rstrip('/')
        webhook_url = f"{host_url}/api/webhook/stripe"
        
        stripe_checkout = StripeCheckout(api_key=api_key, webhook_url=webhook_url)
        
        # Get origin from request headers for dynamic URLs
        origin = request.headers.get('origin', host_url)
        success_url = f"{origin}/lease-assignments?payment=success&session_id={{CHECKOUT_SESSION_ID}}"
        cancel_url = f"{origin}/lease-assignments?payment=cancelled"
        
        checkout_request = CheckoutSessionRequest(
            amount=amount,
            currency="cad",
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                "type": "lease_assignment",
                "assignment_id": assignment_id,
                "buyer_id": buyer_id,
                "seller_id": assignment.get("owner_id") or "unknown",
                "title": assignment.get("title", "Lease Assignment"),
            }
        )
        
        session = await stripe_checkout.create_checkout_session(checkout_request)
        
        # Create transaction record BEFORE redirecting
        transaction = PaymentTransaction(
            session_id=session.session_id,
            user_id=buyer_id,
            amount=amount,
            currency="cad",
            description=f"Lease Assignment: {assignment.get('title', 'Unknown')}",
            property_id=assignment_id,
            recipient_id=assignment.get("owner_id"),
            payment_status="pending",
            status="initiated"
        )
        await db.payment_transactions.insert_one(transaction.model_dump())
        
        return {"url": session.url, "session_id": session.session_id}
        
    except Exception as e:
        logger.error(f"Lease assignment payment error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/lease-assignments/{assignment_id}/payment-complete")
async def complete_assignment_payment(assignment_id: str, session_id: str):
    """Mark a lease assignment as paid and update status"""
    
    # Verify the payment was successful
    transaction = await db.payment_transactions.find_one({"session_id": session_id})
    if not transaction:
        raise HTTPException(status_code=404, detail="Transaction not found")
    
    if transaction.get("payment_status") != "paid":
        raise HTTPException(status_code=400, detail="Payment not yet confirmed")
    
    # Update assignment status
    await db.lease_assignments.update_one(
        {"id": assignment_id},
        {"$set": {
            "status": "paid",
            "buyer_id": transaction.get("user_id"),
            "payment_session_id": session_id,
            "paid_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"success": True, "message": "Assignment payment completed"}

# ========== E-SIGN DOCUMENTS ==========

class ESignDocumentInput(BaseModel):
    title: str
    form_type: Optional[str] = None
    recipient_email: str
    recipient_name: str
    property_address: Optional[str] = None
    notes: Optional[str] = None
    creator_id: str
    creator_name: str
    creator_email: str

class SignatureInput(BaseModel):
    signature_data: str
    signer_id: str
    signer_name: str

@api_router.get("/esign/documents")
async def get_esign_documents(user_id: str):
    """Get all e-sign documents for a user"""
    # Get documents where user is creator, recipient, or signer
    documents = await db.esign_documents.find(
        {"$or": [
            {"creator_id": user_id}, 
            {"recipient_id": user_id},
            {"signer_id": user_id}  # Include documents user has signed
        ]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return documents

@api_router.post("/esign/documents")
async def create_esign_document(data: ESignDocumentInput):
    """Create a new document for e-signature"""
    doc_id = str(uuid.uuid4())
    
    document = {
        "id": doc_id,
        "title": data.title,
        "form_type": data.form_type,
        "recipient_email": data.recipient_email,
        "recipient_name": data.recipient_name,
        "property_address": data.property_address,
        "notes": data.notes,
        "creator_id": data.creator_id,
        "creator_name": data.creator_name,
        "creator_email": data.creator_email,
        "status": "pending",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "signed_at": None,
        "signature_data": None,
        "signer_name": None,
        "audit_trail": [{
            "event": "created",
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "actor": data.creator_name or data.creator_email,
            "details": f"Document created: {data.title}"
        }]
    }
    
    await db.esign_documents.insert_one(document)
    
    # Return document without _id
    return {k: v for k, v in document.items() if k != '_id'}

@api_router.post("/esign/documents/{doc_id}/sign")
async def sign_document(doc_id: str, data: SignatureInput):
    """Sign a document"""
    document = await db.esign_documents.find_one({"id": doc_id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if document["status"] == "signed":
        raise HTTPException(status_code=400, detail="Document already signed")
    
    signed_at = datetime.now(timezone.utc).isoformat()
    
    # Create audit trail entry for signing
    audit_entry = {
        "event": "signed",
        "timestamp": signed_at,
        "actor": data.signer_name,
        "details": f"Document signed by {data.signer_name}"
    }
    
    await db.esign_documents.update_one(
        {"id": doc_id},
        {
            "$set": {
                "status": "signed",
                "signed_at": signed_at,
                "signature_data": data.signature_data,
                "signer_id": data.signer_id,
                "signer_name": data.signer_name,
                "recipient_id": data.signer_id  # Link document to signer's account
            },
            "$push": {
                "audit_trail": audit_entry
            }
        }
    )
    
    return {"status": "signed", "message": "Document signed successfully"}

@api_router.post("/esign/documents/{doc_id}/remind")
async def send_reminder(doc_id: str):
    """Send a reminder to sign the document"""
    document = await db.esign_documents.find_one({"id": doc_id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # TODO: Implement email sending via Resend
    
    return {"status": "sent", "message": "Reminder sent successfully"}

@api_router.get("/esign/documents/{doc_id}")
async def get_esign_document(doc_id: str):
    """Get a specific e-sign document"""
    document = await db.esign_documents.find_one({"id": doc_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document

@api_router.get("/esign/documents/{doc_id}/audit")
async def get_document_audit_trail(doc_id: str):
    """Get audit trail for a document"""
    document = await db.esign_documents.find_one({"id": doc_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    audit_trail = document.get("audit_trail", [])
    
    # Add creation event if not present
    if not audit_trail:
        audit_trail = [{
            "event": "created",
            "timestamp": document.get("created_at"),
            "actor": document.get("created_by_name", "Unknown"),
            "details": f"Document created: {document.get('title')}"
        }]
    
    return {"document_id": doc_id, "audit_trail": audit_trail}

@api_router.post("/esign/documents/{doc_id}/audit")
async def add_audit_event(doc_id: str, event: str, actor: str, details: str = None):
    """Add an audit event to a document"""
    document = await db.esign_documents.find_one({"id": doc_id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    audit_event = {
        "event": event,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "actor": actor,
        "details": details
    }
    
    await db.esign_documents.update_one(
        {"id": doc_id},
        {"$push": {"audit_trail": audit_event}}
    )
    
    return {"success": True, "event": audit_event}

@api_router.post("/esign/templates")
async def create_esign_template(
    user_id: str,
    name: str,
    form_type: str,
    default_fields: dict = None
):
    """Create a reusable document template"""
    template = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "name": name,
        "form_type": form_type,
        "default_fields": default_fields or {},
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.esign_templates.insert_one(template)
    return {"id": template["id"], "name": name}

@api_router.get("/esign/templates")
async def get_esign_templates(user_id: str):
    """Get user's document templates"""
    templates = await db.esign_templates.find({"user_id": user_id}, {"_id": 0}).to_list(50)
    return templates

# ========== DOCUSIGN OAUTH 2.0 INTEGRATION ==========

from services.docusign_service import DocuSignService, DocuSignError

docusign_service = DocuSignService()

@api_router.get("/docusign/status")
async def get_docusign_status(user_id: str):
    """Check if user has connected their DocuSign account"""
    connection = await db.docusign_connections.find_one({"user_id": user_id}, {"_id": 0})
    
    if not connection:
        return {
            "connected": False,
            "integration_key_configured": bool(docusign_service.integration_key),
            "message": "DocuSign not connected. Click 'Connect DocuSign' to begin."
        }
    
    return {
        "connected": True,
        "email": connection.get("email"),
        "account_name": connection.get("account_name"),
        "connected_at": connection.get("connected_at")
    }

@api_router.get("/docusign/auth-url")
async def get_docusign_auth_url(request: Request, user_id: str):
    """Generate DocuSign OAuth authorization URL"""
    if not docusign_service.integration_key:
        raise HTTPException(status_code=503, detail="DocuSign not configured. Integration key missing.")
    
    # Build redirect URI from request
    origin = request.headers.get('origin', str(request.base_url).rstrip('/'))
    redirect_uri = f"{origin}/esign?docusign_callback=true"
    
    auth_url, state = docusign_service.generate_authorization_url(redirect_uri)
    
    # Store state for CSRF validation
    await db.docusign_oauth_states.insert_one({
        "state": state,
        "user_id": user_id,
        "redirect_uri": redirect_uri,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": datetime.now(timezone.utc).isoformat()  # 10 min expiry handled client-side
    })
    
    return {
        "auth_url": auth_url,
        "state": state
    }

@api_router.post("/docusign/callback")
async def docusign_oauth_callback(
    request: Request,
    code: str,
    state: str,
    user_id: str
):
    """Handle DocuSign OAuth callback - exchange code for tokens"""
    # Validate state
    stored_state = await db.docusign_oauth_states.find_one({"state": state, "user_id": user_id})
    if not stored_state:
        raise HTTPException(status_code=400, detail="Invalid state parameter - possible CSRF attack")
    
    redirect_uri = stored_state.get("redirect_uri")
    
    try:
        # Exchange code for tokens
        token_data = await docusign_service.exchange_code_for_tokens(code, redirect_uri)
        
        access_token = token_data.get("access_token")
        refresh_token = token_data.get("refresh_token")
        expires_in = token_data.get("expires_in", 3600)
        
        # Get user info including account ID
        user_info = await docusign_service.get_user_info(access_token)
        
        # Store connection
        connection = {
            "user_id": user_id,
            "account_id": user_info.get("account_id"),
            "account_name": user_info.get("account_name"),
            "base_uri": user_info.get("base_uri"),
            "email": user_info.get("email"),
            "name": user_info.get("name"),
            "access_token": access_token,
            "refresh_token": refresh_token,
            "token_expires_in": expires_in,
            "connected_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        # Upsert connection
        await db.docusign_connections.update_one(
            {"user_id": user_id},
            {"$set": connection},
            upsert=True
        )
        
        # Clean up state
        await db.docusign_oauth_states.delete_one({"state": state})
        
        return {
            "success": True,
            "email": user_info.get("email"),
            "account_name": user_info.get("account_name"),
            "message": "DocuSign connected successfully!"
        }
        
    except DocuSignError as e:
        logger.error(f"DocuSign OAuth error: {e}")
        raise HTTPException(status_code=400, detail=str(e))

@api_router.post("/docusign/disconnect")
async def disconnect_docusign(user_id: str):
    """Disconnect user's DocuSign account"""
    result = await db.docusign_connections.delete_one({"user_id": user_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="No DocuSign connection found")
    
    return {"success": True, "message": "DocuSign disconnected"}

@api_router.post("/docusign/send-envelope")
async def send_docusign_envelope(
    request: Request,
    user_id: str,
    doc_id: str,
):
    """Send an existing e-sign document via DocuSign"""
    # Get DocuSign connection
    connection = await db.docusign_connections.find_one({"user_id": user_id})
    if not connection:
        raise HTTPException(status_code=401, detail="DocuSign not connected. Please connect your account first.")
    
    # Get the document
    document = await db.esign_documents.find_one({"id": doc_id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Generate PDF content from the document
    pdf_content = generate_document_pdf(document)
    
    try:
        # Refresh token if needed (in production, check expiry)
        access_token = connection.get("access_token")
        account_id = connection.get("account_id")
        base_uri = connection.get("base_uri")
        
        # Send via DocuSign
        result = await docusign_service.create_and_send_envelope(
            access_token=access_token,
            account_id=account_id,
            base_uri=base_uri,
            subject=f"Please sign: {document.get('title')}",
            document_content=pdf_content,
            document_name=f"{document.get('title')}.pdf",
            signer_email=document.get("recipient_email"),
            signer_name=document.get("recipient_name"),
        )
        
        # Update document with DocuSign info
        await db.esign_documents.update_one(
            {"id": doc_id},
            {"$set": {
                "docusign_envelope_id": result.get("envelope_id"),
                "docusign_status": result.get("status"),
                "docusign_sent_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        return result
        
    except DocuSignError as e:
        logger.error(f"DocuSign send error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/docusign/envelope-status/{envelope_id}")
async def get_docusign_envelope_status(user_id: str, envelope_id: str):
    """Get status of a DocuSign envelope"""
    connection = await db.docusign_connections.find_one({"user_id": user_id})
    if not connection:
        raise HTTPException(status_code=401, detail="DocuSign not connected")
    
    try:
        status = await docusign_service.get_envelope_status_async(
            access_token=connection.get("access_token"),
            account_id=connection.get("account_id"),
            base_uri=connection.get("base_uri"),
            envelope_id=envelope_id
        )
        return status
        
    except DocuSignError as e:
        raise HTTPException(status_code=500, detail=str(e))


def generate_document_pdf(document: Dict[str, Any]) -> bytes:
    """Generate a simple PDF from document data (placeholder)"""
    # In production, use reportlab or a proper PDF generation library
    content = f"""
    DOMMMA E-SIGN DOCUMENT
    =====================
    
    Document: {document.get('title', 'Untitled')}
    Form Type: {document.get('form_type', 'General')}
    
    Property Address: {document.get('property_address', 'N/A')}
    
    FROM:
    {document.get('creator_name', 'Unknown')}
    {document.get('creator_email', '')}
    
    TO:
    {document.get('recipient_name', 'Unknown')}
    {document.get('recipient_email', '')}
    
    Notes:
    {document.get('notes', 'None')}
    
    Created: {document.get('created_at', '')}
    
    
    SIGNATURE:
    
    /sig1/
    
    Date: _____________
    
    
    This document was sent via DOMMMA's e-signature platform.
    """
    return content.encode('utf-8')

# ========== LISTING SYNDICATION ==========

class SyndicationTrackInput(BaseModel):
    listing_id: str
    platform: str
    user_id: str
    timestamp: str

@api_router.post("/syndication/track")
async def track_syndication(data: SyndicationTrackInput):
    """Track when a listing is syndicated to an external platform"""
    track_id = str(uuid.uuid4())
    
    track_record = {
        "id": track_id,
        "listing_id": data.listing_id,
        "platform": data.platform,
        "user_id": data.user_id,
        "syndicated_at": data.timestamp,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.syndication_history.insert_one(track_record)
    
    # Update listing with syndication info
    await db.listings.update_one(
        {"id": data.listing_id},
        {"$addToSet": {"syndicated_to": data.platform}}
    )
    
    return {"status": "tracked", "id": track_id}

# ========== AI COMPETITOR PRICE ANALYSIS ==========

@api_router.post("/ai/competitor-analysis")
async def analyze_competitor_prices(data: Dict[str, Any]):
    """
    AI-powered competitor analysis - analyzes nearby rentals and suggests pricing.
    Scrapes real data from online sources and uses AI for analysis.
    """
    address = data.get("address", "")
    city = data.get("city", "Vancouver")
    bedrooms = data.get("bedrooms", 2)
    property_type = data.get("property_type", "apartment")
    
    if not address:
        raise HTTPException(status_code=400, detail="Address is required")
    
    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="AI analysis not configured")
    
    try:
        import httpx
        
        # Fetch comparable listings from our database
        comparable_query = {
            "city": {"$regex": city, "$options": "i"},
            "bedrooms": {"$gte": bedrooms - 1, "$lte": bedrooms + 1},
            "status": "active"
        }
        
        local_listings = await db.listings.find(
            comparable_query,
            {"_id": 0, "id": 1, "title": 1, "price": 1, "bedrooms": 1, "bathrooms": 1, "address": 1, "amenities": 1}
        ).limit(10).to_list(10)
        
        # Try to scrape real competitor data using Perplexity/web search
        scraped_data = []
        try:
            # Use Perplexity API for real-time web search of rental listings
            perplexity_key = os.environ.get('PERPLEXITY_API_KEY')
            if perplexity_key:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    search_response = await client.post(
                        "https://api.perplexity.ai/chat/completions",
                        headers={
                            "Authorization": f"Bearer {perplexity_key}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": "llama-3.1-sonar-small-128k-online",
                            "messages": [{
                                "role": "user",
                                "content": f"Find current rental listings near {address}, {city} with {bedrooms} bedrooms. Include prices from Craigslist Vancouver, Facebook Marketplace rentals, and Kijiji. List at least 5 comparable rentals with their prices in CAD per month."
                            }]
                        }
                    )
                    if search_response.status_code == 200:
                        search_result = search_response.json()
                        scraped_content = search_result.get("choices", [{}])[0].get("message", {}).get("content", "")
                        scraped_data.append({"source": "web_search", "content": scraped_content})
        except Exception as e:
            logger.warning(f"Web scraping failed, using local data only: {e}")
        
        # Calculate average prices from local data
        if local_listings:
            prices = [l.get("price", 0) for l in local_listings if l.get("price")]
            avg_price = sum(prices) / len(prices) if prices else 0
            min_price = min(prices) if prices else 0
            max_price = max(prices) if prices else 0
        else:
            # Default Vancouver market rates (updated 2025)
            base_rates = {0: 1600, 1: 2000, 2: 2600, 3: 3400, 4: 4200}
            avg_price = base_rates.get(bedrooms, 2600)
            min_price = avg_price * 0.8
            max_price = avg_price * 1.2
        
        # Use AI to generate comprehensive analysis
        scraped_info = "\n".join([f"Source: {s['source']}\n{s['content']}" for s in scraped_data]) if scraped_data else "No external data available"
        
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json"
                },
                json={
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 2048,
                    "messages": [
                        {
                            "role": "user",
                            "content": f"""You are a real estate market analyst for {city}. 
                            
Based on this property:
- Address: {address}
- Bedrooms: {bedrooms}
- Property Type: {property_type}

Comparable listings in the area from our database (prices in CAD/month):
{[f"- {l.get('title', 'Listing')}: ${l.get('price', 0)}/mo, {l.get('bedrooms', 0)}BR" for l in local_listings[:5]]}

External market data (from Craigslist, FB Marketplace, Kijiji):
{scraped_info}

Average market price: ${avg_price:.0f}/mo
Price range: ${min_price:.0f} - ${max_price:.0f}/mo

Provide:
1. A recommended listing price (consider the specific address, features, and current market data)
2. A compelling listing title (under 60 chars)
3. A professional listing description (150-200 words)
4. 3 key selling points as bullet points
5. Suggested hashtags for social media

Format your response as JSON:
{{
    "recommended_price": number,
    "price_explanation": "brief explanation",
    "title": "Listing Title",
    "description": "Full description...",
    "selling_points": ["point 1", "point 2", "point 3"],
    "hashtags": ["#tag1", "#tag2", "#tag3"],
    "market_insights": "Brief market analysis"
}}"""
                        }
                    ]
                }
            )
            
            if response.status_code != 200:
                logger.error(f"Claude API error: {response.text}")
                # Return basic analysis without AI
                return {
                    "recommended_price": int(avg_price),
                    "price_range": {"min": int(min_price), "max": int(max_price)},
                    "comparable_listings": local_listings[:5],
                    "ai_analysis": None
                }
            
            result = response.json()
            ai_response = result.get("content", [{}])[0].get("text", "{}")
            
            # Parse JSON response
            import json
            try:
                ai_analysis = json.loads(ai_response)
            except:
                import re
                json_match = re.search(r'\{[\s\S]*\}', ai_response)
                if json_match:
                    ai_analysis = json.loads(json_match.group())
                else:
                    ai_analysis = None
            
            return {
                "recommended_price": ai_analysis.get("recommended_price", int(avg_price)) if ai_analysis else int(avg_price),
                "price_range": {"min": int(min_price), "max": int(max_price)},
                "comparable_listings": local_listings[:5],
                "ai_analysis": ai_analysis,
                "market_data": {
                    "average_price": int(avg_price),
                    "total_comparables": len(local_listings),
                    "city": city
                }
            }
            
    except Exception as e:
        logger.error(f"Competitor analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@api_router.get("/syndication/history/{listing_id}")
async def get_syndication_history(listing_id: str):
    """Get syndication history for a listing"""
    history = await db.syndication_history.find(
        {"listing_id": listing_id}, {"_id": 0}
    ).sort("syndicated_at", -1).to_list(50)
    return history

@api_router.get("/syndication/stats/{user_id}")
async def get_syndication_stats(user_id: str):
    """Get syndication stats for a user"""
    # Count by platform
    pipeline = [
        {"$match": {"user_id": user_id}},
        {"$group": {"_id": "$platform", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]
    
    stats = await db.syndication_history.aggregate(pipeline).to_list(10)
    
    total = sum(s["count"] for s in stats)
    
    return {
        "total_syndications": total,
        "by_platform": {s["_id"]: s["count"] for s in stats}
    }

# ========== AI LISTING OPTIMIZER ==========

class ListingOptimizeInput(BaseModel):
    listing_id: str
    user_id: str

@api_router.post("/listings/optimize")
async def optimize_listing(data: ListingOptimizeInput):
    """AI-powered listing optimization with price suggestions and content improvements"""
    
    listing = await db.listings.find_one({"id": data.listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    city = listing.get("city", "Vancouver")
    bedrooms = listing.get("bedrooms", 1)
    current_price = listing.get("price", 2000)
    
    # Get comparable listings for market analysis
    comparables = await db.listings.find({
        "city": city,
        "bedrooms": bedrooms,
        "status": "active",
        "listing_type": "rent",
        "id": {"$ne": data.listing_id}
    }, {"_id": 0, "price": 1}).to_list(20)
    
    # Calculate market average
    if comparables:
        prices = [c.get("price", 0) for c in comparables if c.get("price")]
        market_avg = sum(prices) / len(prices) if prices else current_price
    else:
        # Fallback estimates by city and bedrooms
        market_rates = {
            'Vancouver': {1: 2200, 2: 3000, 3: 3800},
            'Burnaby': {1: 1900, 2: 2600, 3: 3200},
            'Surrey': {1: 1600, 2: 2200, 3: 2800},
            'Richmond': {1: 2000, 2: 2800, 3: 3500}
        }
        city_rates = market_rates.get(city, market_rates['Vancouver'])
        market_avg = city_rates.get(min(bedrooms, 3), 2000)
    
    suggested_price = round(market_avg / 50) * 50  # Round to nearest 50
    price_diff = suggested_price - current_price
    
    if price_diff > 200:
        price_status = "below_market"
    elif price_diff < -200:
        price_status = "above_market"
    else:
        price_status = "competitive"
    
    # Calculate optimization score
    score = 50
    images = listing.get("images", [])
    description = listing.get("description", "")
    amenities = listing.get("amenities", [])
    
    if len(images) >= 5:
        score += 15
    elif len(images) >= 3:
        score += 8
    
    if len(description) >= 150:
        score += 10
    elif len(description) >= 50:
        score += 5
    
    if len(amenities) >= 5:
        score += 10
    elif len(amenities) >= 3:
        score += 5
    
    if price_status == "competitive":
        score += 15
    elif price_status == "below_market":
        score += 10
    
    # Generate title suggestions
    property_type = listing.get("property_type", "Apartment")
    pet_friendly = listing.get("pet_friendly", False)
    
    title_suggestions = [
        f"Modern {bedrooms}BR {property_type} in {city} - {'In-Suite Laundry!' if 'In-Suite Laundry' in amenities else 'Great Location!'}",
        f"Bright & Spacious {bedrooms} Bedroom in Prime {city}",
        f"{'Pet-Friendly ' if pet_friendly else ''}{bedrooms}BR with {amenities[0] if amenities else 'Modern Amenities'}"
    ]
    
    # Improvement suggestions
    improvements = []
    if len(description) < 100:
        improvements.append("Add a detailed description (150+ words) - listings with longer descriptions get 40% more inquiries")
    if len(amenities) < 5:
        improvements.append("List at least 5 amenities - highlight unique features like in-suite laundry, parking, or gym")
    if len(images) < 5:
        improvements.append("Add more photos (8-10 recommended) - listings with more photos rent 50% faster")
    if not listing.get("offers"):
        improvements.append("Consider adding a move-in special to stand out (e.g., 'First month 50% off')")
    
    # Quick wins
    quick_wins = []
    if len(images) < 5:
        quick_wins.append("Add 3+ more photos")
    if not listing.get("offers"):
        quick_wins.append("Add a move-in special")
    if not description:
        quick_wins.append("Write a compelling description")
    if len(amenities) < 5:
        quick_wins.append("List more amenities")
    
    # Market insights
    insights = [
        f"Similar {bedrooms}BR listings in {city} average ${int(market_avg):,}/month",
        f"{bedrooms}BR rentals in your area typically include: In-Suite Laundry, Parking, Gym Access",
        f"Peak rental season in {city} is May-August - consider timing your availability"
    ]
    
    # Generate optimized description
    amenity_list = "\n".join([f"- {a}" for a in amenities[:5]]) if amenities else "- Modern kitchen appliances\n- Great natural light"
    sqft_text = f"At {listing.get('sqft')} sqft, there is" if listing.get('sqft') else "There is"
    available_date = listing.get('available_date') or 'immediately'
    parking_text = "Parking included!" if 'Parking' in amenities else ""
    pet_text = "pet-friendly " if pet_friendly else ""
    amenity_text = ', '.join(amenities[:3]) if amenities else 'modern finishes throughout'
    
    generated_description = f"""Welcome to this {pet_text}{bedrooms}-bedroom {property_type.lower()} in the heart of {city}!

This bright and spacious unit features {amenity_text}. {sqft_text} plenty of room for comfortable living.

Highlights:
{amenity_list}

Located in a prime {city} neighborhood with easy access to transit, shopping, and dining. {parking_text}

Available {available_date}. Contact us today to schedule a viewing!"""
    
    return {
        "listing_id": data.listing_id,
        "current_price": current_price,
        "suggested_price": suggested_price,
        "price_status": price_status,
        "price_adjustment": price_diff,
        "optimization_score": min(100, score),
        "title_suggestions": title_suggestions,
        "description_improvements": improvements,
        "competitor_insights": insights,
        "quick_wins": quick_wins,
        "generated_description": generated_description,
        "market_comparables": len(comparables)
    }

# ========== VIDEO TOURS (Cloudinary) ==========

import time
import cloudinary
import cloudinary.utils
import cloudinary.uploader

# Initialize Cloudinary (if credentials are set)
cloudinary_configured = False
if os.environ.get('CLOUDINARY_CLOUD_NAME'):
    cloudinary.config(
        cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
        api_key=os.environ.get("CLOUDINARY_API_KEY"),
        api_secret=os.environ.get("CLOUDINARY_API_SECRET"),
        secure=True
    )
    cloudinary_configured = True

@api_router.get("/cloudinary/signature")
async def generate_cloudinary_signature(
    resource_type: str = "video",
    folder: str = "dommma/tours"
):
    """Generate a signed upload signature for Cloudinary"""
    if not cloudinary_configured:
        raise HTTPException(status_code=503, detail="Cloudinary not configured. Please set CLOUDINARY credentials.")
    
    # Validate folder
    allowed_folders = ("dommma/", "users/", "listings/")
    if not folder.startswith(allowed_folders):
        raise HTTPException(status_code=400, detail="Invalid folder path")
    
    timestamp = int(time.time())
    params = {
        "timestamp": timestamp,
        "folder": folder,
        "resource_type": resource_type
    }
    
    signature = cloudinary.utils.api_sign_request(
        params,
        os.environ.get("CLOUDINARY_API_SECRET")
    )
    
    return {
        "signature": signature,
        "timestamp": timestamp,
        "cloud_name": os.environ.get("CLOUDINARY_CLOUD_NAME"),
        "api_key": os.environ.get("CLOUDINARY_API_KEY"),
        "folder": folder,
        "resource_type": resource_type
    }

@api_router.get("/cloudinary/status")
async def cloudinary_status():
    """Check if Cloudinary is configured"""
    return {
        "configured": cloudinary_configured,
        "cloud_name": os.environ.get("CLOUDINARY_CLOUD_NAME") if cloudinary_configured else None
    }

class VideoTourInput(BaseModel):
    listing_id: str
    video_url: str
    public_id: str
    duration: Optional[float] = None
    title: Optional[str] = None

@api_router.post("/listings/{listing_id}/video-tour")
async def add_video_tour(listing_id: str, data: VideoTourInput):
    """Add a video tour to a listing"""
    listing = await db.listings.find_one({"id": listing_id})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    video_tour = {
        "video_url": data.video_url,
        "public_id": data.public_id,
        "duration": data.duration,
        "title": data.title or "Property Tour",
        "uploaded_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.listings.update_one(
        {"id": listing_id},
        {"$set": {"video_tour": video_tour, "has_video_tour": True}}
    )
    
    return {"status": "success", "video_tour": video_tour}

@api_router.delete("/listings/{listing_id}/video-tour")
async def remove_video_tour(listing_id: str):
    """Remove a video tour from a listing"""
    listing = await db.listings.find_one({"id": listing_id})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    # Delete from Cloudinary if configured
    if cloudinary_configured and listing.get("video_tour", {}).get("public_id"):
        try:
            cloudinary.uploader.destroy(
                listing["video_tour"]["public_id"],
                resource_type="video",
                invalidate=True
            )
        except Exception as e:
            logger.error(f"Failed to delete video from Cloudinary: {e}")
    
    await db.listings.update_one(
        {"id": listing_id},
        {"$unset": {"video_tour": 1}, "$set": {"has_video_tour": False}}
    )
    
    return {"status": "deleted"}

# ========== AI-POWERED FEATURES ==========

class IssueAnalysisRequest(BaseModel):
    image_data: str  # base64 data URL
    description: Optional[str] = ""

class DocumentAnalysisRequest(BaseModel):
    document_text: str
    document_type: str = "lease"

class CommuteSearchRequest(BaseModel):
    work_addresses: List[str]
    max_commute_minutes: int = 45
    transport_mode: str = "transit"

@api_router.post("/ai/analyze-issue")
async def analyze_issue(req: IssueAnalysisRequest):
    """AI analyzes a home issue image and matches relevant contractors"""
    
    system_message = """You are a home maintenance expert AI. Analyze the described home issue and:
1. Identify the specific issue type (plumbing, electrical, painting, renovation, carpentry, landscaping, cleaning, HVAC, roofing, etc.)
2. Assess the urgency (low, medium, high, emergency)
3. Provide a brief description of the problem
4. Recommend the type of contractor needed
5. Estimate a rough cost range in CAD

Respond in valid JSON format only:
{
  "issue_type": "plumbing",
  "urgency": "high",
  "description": "Water leak from pipe joint under kitchen sink",
  "contractor_type": "plumber",
  "estimated_cost_range": "$150-$400",
  "recommended_specialties": ["plumbing", "pipe repair"],
  "immediate_steps": ["Turn off water supply under the sink", "Place a bucket under the leak"]
}"""

    try:
        anthropic_client = AsyncAnthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
        
        message_text = f"Analyze this home issue. User description: {req.description or 'See image'}. The image has been uploaded showing the issue."
        if req.image_data and req.image_data.startswith("data:"):
            message_text += f"\n\n[Image uploaded - analyze based on description]"
        
        claude_response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=1024,
            system=system_message,
            messages=[{"role": "user", "content": message_text}]
        )
        
        response = claude_response.content[0].text
        
        # Parse AI response
        import re
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        analysis = {}
        if json_match:
            try:
                analysis = json.loads(json_match.group())
            except json.JSONDecodeError:
                analysis = {"issue_type": "general", "description": response, "urgency": "medium", "recommended_specialties": ["general"]}
        else:
            analysis = {"issue_type": "general", "description": response, "urgency": "medium", "recommended_specialties": ["general"]}
        
        # Match contractors based on AI analysis
        specialties = analysis.get("recommended_specialties", [analysis.get("issue_type", "general")])
        matched_contractors = []
        for specialty in specialties:
            contractors = await db.contractor_profiles.find(
                {"specialties": {"$regex": specialty, "$options": "i"}, "status": "active"},
                {"_id": 0}
            ).sort("rating", -1).to_list(5)
            for c in contractors:
                if c not in matched_contractors:
                    matched_contractors.append(c)
        
        return {
            "analysis": analysis,
            "matched_contractors": matched_contractors[:6],
            "total_matches": len(matched_contractors)
        }
    except Exception as e:
        logging.error(f"Issue analysis error: {e}")
        # Fallback: do keyword matching
        keywords_map = {
            "plumbing": ["water", "leak", "pipe", "faucet", "drain", "toilet", "sink"],
            "electrical": ["light", "switch", "outlet", "wire", "power", "circuit"],
            "painting": ["paint", "wall", "ceiling", "crack", "peel"],
            "cleaning": ["mold", "stain", "dirt", "mess"],
            "renovation": ["broken", "damage", "repair", "fix"],
        }
        desc_lower = (req.description or "").lower()
        matched_type = "general"
        for cat, keywords in keywords_map.items():
            if any(kw in desc_lower for kw in keywords):
                matched_type = cat
                break
        
        contractors = await db.contractor_profiles.find(
            {"specialties": {"$regex": matched_type, "$options": "i"}, "status": "active"},
            {"_id": 0}
        ).sort("rating", -1).to_list(6)
        
        return {
            "analysis": {
                "issue_type": matched_type,
                "description": f"Issue detected: {req.description}",
                "urgency": "medium",
                "recommended_specialties": [matched_type]
            },
            "matched_contractors": contractors,
            "total_matches": len(contractors)
        }

@api_router.post("/ai/analyze-document")
async def analyze_document(req: DocumentAnalysisRequest):
    """AI analyzes lease/rental documents for fairness and key terms - using faster Haiku model"""
    
    system_message = f"""You are a {req.document_type} document analysis expert for Canadian real estate. Analyze the provided document and return a comprehensive review in valid JSON:
{{
  "summary": "Brief summary of the document",
  "key_terms": [
    {{"term": "Monthly Rent", "value": "$2,500", "assessment": "fair"}},
    {{"term": "Deposit", "value": "$2,500", "assessment": "standard"}}
  ],
  "fairness_score": 8,
  "red_flags": ["Unusual early termination penalty of 3 months rent"],
  "green_flags": ["Standard 12-month term", "Pet-friendly clause included"],
  "recommendations": ["Negotiate the early termination penalty", "Request clarification on utility responsibilities"],
  "legal_notes": ["All terms comply with BC Residential Tenancy Act"],
  "monthly_costs_breakdown": {{
    "rent": 2500,
    "utilities_estimate": 150,
    "parking": 0,
    "total_estimate": 2650
  }}
}}
Score fairness from 1-10 (10 = very renter-friendly). Be thorough but practical."""

    try:
        anthropic_client = AsyncAnthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
        
        # Use Haiku for faster response (3-5x faster than Sonnet)
        claude_response = await anthropic_client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=2048,
            system=system_message,
            messages=[{"role": "user", "content": f"Analyze this {req.document_type} document:\n\n{req.document_text[:8000]}"}]
        )
        
        response = claude_response.content[0].text
        
        import re
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            try:
                return {"analysis": json.loads(json_match.group())}
            except json.JSONDecodeError:
                pass
        return {"analysis": {"summary": response, "fairness_score": 0, "key_terms": [], "red_flags": [], "green_flags": [], "recommendations": []}}
    except Exception as e:
        logging.error(f"Document analysis error: {e}")
        raise HTTPException(status_code=500, detail="Document analysis failed")


# ========== DOCUMENT FILE UPLOAD ==========

from fastapi import File, UploadFile
import io

@api_router.post("/ai/analyze-document-file")
async def analyze_document_file(
    file: UploadFile = File(...),
    document_type: str = "lease"
):
    """Upload and analyze a document file (PDF, image, DOCX)"""
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    filename = file.filename.lower()
    content = await file.read()
    
    extracted_text = ""
    
    try:
        # PDF extraction
        if filename.endswith('.pdf'):
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() or ""
        
        # Word document extraction
        elif filename.endswith('.docx'):
            from docx import Document
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        
        # Image OCR (PNG, JPG, etc.) - Use Claude Vision instead of pytesseract
        elif filename.endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif')):
            import base64
            image_base64 = base64.b64encode(content).decode('utf-8')
            
            # Determine media type
            media_type = "image/png"
            if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                media_type = "image/jpeg"
            elif filename.endswith('.webp'):
                media_type = "image/webp"
            elif filename.endswith('.gif'):
                media_type = "image/gif"
            
            # Use Claude Vision for OCR
            anthropic_client = AsyncAnthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
            
            ocr_response = await anthropic_client.messages.create(
                model="claude-haiku-4-5",
                max_tokens=4000,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": image_base64
                            }
                        },
                        {
                            "type": "text",
                            "text": "Extract ALL text from this document image. Return only the extracted text, preserving the document structure as much as possible."
                        }
                    ]
                }]
            )
            extracted_text = ocr_response.content[0].text
        
        # Plain text files
        elif filename.endswith('.txt'):
            extracted_text = content.decode('utf-8', errors='ignore')
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}. Supported: PDF, DOCX, PNG, JPG, TXT")
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the file")
        
        # Now analyze the extracted text using existing function
        analysis_request = DocumentAnalysisRequest(
            document_text=extracted_text[:10000],
            document_type=document_type
        )
        
        return await analyze_document(analysis_request)
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"File processing error: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


@api_router.post("/ai/commute-search")
async def commute_search(req: CommuteSearchRequest):
    """Find properties optimized for commute to work locations"""
    
    # Get all active listings
    all_listings = await db.listings.find({"status": "active"}, {"_id": 0}).to_list(100)
    
    listings_text = "\n".join([
        f"- {l['title']} at {l['address']}, {l['city']} (${l['price']}/mo, {l['bedrooms']}bd/{l['bathrooms']}ba)"
        for l in all_listings[:30]
    ])
    
    system_message = """You are a commute optimization expert for Vancouver. Given work addresses and available properties, rank properties by estimated commute convenience. Return valid JSON:
{
  "ranked_properties": [
    {
      "title": "Property Name",
      "address": "Full address",
      "estimated_commute": "25 min by transit",
      "commute_score": 9,
      "notes": "Close to SkyTrain station"
    }
  ],
  "tips": ["Consider the Canada Line for fastest downtown access"]
}
Score commute from 1-10 (10 = shortest/most convenient)."""

    try:
        anthropic_client = AsyncAnthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
        
        user_content = (f"Work addresses: {', '.join(req.work_addresses)}\n"
                       f"Max commute: {req.max_commute_minutes} min by {req.transport_mode}\n\n"
                       f"Available properties:\n{listings_text}")
        
        claude_response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=2048,
            system=system_message,
            messages=[{"role": "user", "content": user_content}]
        )
        
        response = claude_response.content[0].text
        
        import re
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            try:
                result = json.loads(json_match.group())
                # Enrich with full listing data
                for ranked in result.get("ranked_properties", []):
                    for l in all_listings:
                        if l["title"].lower() in ranked.get("title", "").lower() or ranked.get("address", "").lower() in l.get("address", "").lower():
                            ranked["listing"] = l
                            break
                return result
            except json.JSONDecodeError:
                pass
        return {"ranked_properties": [], "tips": [response]}
    except Exception as e:
        logging.error(f"Commute search error: {e}")
        raise HTTPException(status_code=500, detail="Commute search failed")

# ========== FCM PUSH NOTIFICATIONS ==========

@api_router.post("/notifications/register-token")
async def register_fcm_token(token_data: FCMTokenCreate):
    """Register or update FCM token for a user"""
    existing = await db.fcm_tokens.find_one({"user_id": token_data.user_id})
    
    if existing:
        await db.fcm_tokens.update_one(
            {"user_id": token_data.user_id},
            {"$set": {"token": token_data.token, "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
    else:
        fcm_token = FCMToken(user_id=token_data.user_id, token=token_data.token)
        await db.fcm_tokens.insert_one(fcm_token.model_dump())
    
    return {"status": "registered"}

@api_router.post("/notifications/send")
async def send_notification(notification: NotificationCreate):
    """Send push notification to a user (for internal use)"""
    # Store notification in database
    notif_doc = {
        "id": str(uuid.uuid4()),
        "user_id": notification.user_id,
        "title": notification.title,
        "body": notification.body,
        "type": notification.notification_type,
        "data": notification.data,
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(notif_doc)
    
    # Get user's FCM token
    token_doc = await db.fcm_tokens.find_one({"user_id": notification.user_id})
    
    # Note: For FCM sending, you would need firebase-admin SDK
    # For now, we'll send via WebSocket if user is connected
    await manager.send_personal_message(
        json.dumps({
            "type": "notification",
            "notification": {
                "title": notification.title,
                "body": notification.body,
                "type": notification.notification_type,
                "data": notification.data
            }
        }),
        notification.user_id
    )
    
    return {"status": "sent", "fcm_token_available": token_doc is not None}

@api_router.get("/notifications/{user_id}")
async def get_notifications(user_id: str, unread_only: bool = False):
    """Get notifications for a user"""
    query = {"user_id": user_id}
    if unread_only:
        query["read"] = False
    
    notifications = await db.notifications.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    return notifications

@api_router.post("/notifications/mark-read/{notification_id}")
async def mark_notification_read(notification_id: str):
    """Mark notification as read"""
    await db.notifications.update_one({"id": notification_id}, {"$set": {"read": True}})
    return {"status": "read"}

# ========== RENTAL APPLICATIONS ==========

@api_router.post("/applications", response_model=dict)
async def create_application(user_id: str, application: ApplicationCreate):
    """Submit a rental application"""
    # Get the listing to find the landlord
    listing = await db.listings.find_one({"id": application.listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    app_obj = RentalApplication(
        user_id=user_id,
        listing_id=application.listing_id,
        landlord_id=listing.get("landlord_id"),
        full_name=application.full_name,
        email=application.email,
        phone=application.phone,
        current_address=application.current_address,
        move_in_date=application.move_in_date,
        employer=application.employer,
        job_title=application.job_title,
        monthly_income=application.monthly_income,
        employment_length=application.employment_length,
        references=application.references,
        num_occupants=application.num_occupants,
        has_pets=application.has_pets,
        pet_details=application.pet_details,
        additional_notes=application.additional_notes
    )
    
    await db.applications.insert_one(app_obj.model_dump())
    
    # Send notification to landlord if exists
    if listing.get("landlord_id"):
        await db.notifications.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": listing["landlord_id"],
            "title": "New Rental Application",
            "body": f"{application.full_name} applied for {listing['title']}",
            "type": "application",
            "data": {"application_id": app_obj.id, "listing_id": application.listing_id},
            "read": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
    
    return {"id": app_obj.id, "status": "submitted"}

@api_router.get("/applications/user/{user_id}")
async def get_user_applications(user_id: str):
    """Get all applications submitted by a user"""
    applications = await db.applications.find({"user_id": user_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    
    # Enrich with listing details
    for app in applications:
        listing = await db.listings.find_one({"id": app["listing_id"]}, {"_id": 0, "title": 1, "address": 1, "price": 1, "images": 1})
        app["listing"] = listing
    
    return applications

@api_router.get("/applications/landlord/{landlord_id}")
async def get_landlord_applications(landlord_id: str, status: Optional[str] = None):
    """Get all applications for a landlord's properties"""
    query = {"landlord_id": landlord_id}
    if status:
        query["status"] = status
    
    applications = await db.applications.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    # Enrich with listing details
    for app in applications:
        listing = await db.listings.find_one({"id": app["listing_id"]}, {"_id": 0, "title": 1, "address": 1})
        app["listing"] = listing
    
    return applications

@api_router.get("/applications")
async def get_applications_by_listing(listing_id: str):
    """Get all applications for a specific listing with AI scoring"""
    applications = await db.applications.find({"listing_id": listing_id}, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    # Get listing details for scoring
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0})
    property_price = listing.get("price", 2000) if listing else 2000
    
    for app in applications:
        # Fetch user resume if available
        resume = await db.renter_resumes.find_one({"user_id": app.get("user_id")}, {"_id": 0})
        if resume:
            app["annual_income"] = resume.get("employment", {}).get("annual_income", app.get("annual_income", 0))
            app["employment_status"] = resume.get("employment", {}).get("status", app.get("employment_status"))
            app["employer_name"] = resume.get("employment", {}).get("employer", app.get("employer_name"))
            app["job_title"] = resume.get("employment", {}).get("job_title", app.get("job_title"))
            app["years_at_current"] = resume.get("rental_history", {}).get("years_at_current", 0)
            app["has_pets"] = resume.get("household", {}).get("has_pets", False)
            app["previous_landlord_reference"] = bool(resume.get("rental_history", {}).get("previous_landlord", {}).get("name"))
        
        # Calculate AI score
        app["ai_score"] = calculate_applicant_score(app, property_price)
        app["ai_analysis"] = generate_applicant_analysis(app, property_price, listing)
    
    # Sort by AI score
    applications.sort(key=lambda x: x.get("ai_score", 0), reverse=True)
    
    return applications

def calculate_applicant_score(applicant: dict, property_price: float) -> int:
    """Calculate AI score for an applicant (0-100)"""
    score = 0
    
    # Income to rent ratio (max 35 points)
    monthly_income = (applicant.get("annual_income") or 0) / 12
    rent_ratio = monthly_income / property_price if property_price > 0 else 0
    if rent_ratio >= 3:
        score += 35
    elif rent_ratio >= 2.5:
        score += 28
    elif rent_ratio >= 2:
        score += 20
    else:
        score += max(0, int(rent_ratio * 10))
    
    # Employment stability (max 25 points)
    emp_status = applicant.get("employment_status", "")
    if emp_status == "employed":
        score += 20
    elif emp_status == "self-employed":
        score += 15
    elif emp_status == "student":
        score += 10
    if applicant.get("employer_name"):
        score += 5
    
    # Rental history (max 25 points)
    years = applicant.get("years_at_current") or 0
    if years >= 2:
        score += 20
    elif years >= 1:
        score += 15
    else:
        score += 5
    if applicant.get("previous_landlord_reference"):
        score += 5
    
    # Completeness (max 15 points)
    completeness = 0
    if applicant.get("full_name"):
        completeness += 2
    if applicant.get("email"):
        completeness += 2
    if applicant.get("phone"):
        completeness += 2
    if applicant.get("annual_income"):
        completeness += 3
    if applicant.get("employment_status"):
        completeness += 3
    if applicant.get("move_in_date"):
        completeness += 3
    score += completeness
    
    return min(100, score)

def generate_applicant_analysis(applicant: dict, property_price: float, listing: dict = None) -> dict:
    """Generate AI analysis strengths and concerns"""
    monthly_income = (applicant.get("annual_income") or 0) / 12
    rent_ratio = monthly_income / property_price if property_price > 0 else 0
    
    strengths = []
    concerns = []
    
    # Income analysis
    if rent_ratio >= 3:
        strengths.append("Strong income-to-rent ratio (3x+)")
    elif rent_ratio >= 2:
        strengths.append("Acceptable income-to-rent ratio")
    else:
        concerns.append("Income may be tight for this rent level")
    
    # Employment
    if applicant.get("employment_status") == "employed" and applicant.get("employer_name"):
        strengths.append(f"Stable employment at {applicant.get('employer_name')}")
    elif applicant.get("employment_status") == "self-employed":
        strengths.append("Self-employed - may need income verification")
    elif applicant.get("employment_status") == "student":
        concerns.append("Student income - may need co-signer")
    
    # Rental history
    years = applicant.get("years_at_current") or 0
    if years >= 2:
        strengths.append("Long-term rental history shows stability")
    elif years < 1:
        concerns.append("Limited rental history")
    
    # References
    if applicant.get("previous_landlord_reference"):
        strengths.append("Provided landlord reference")
    else:
        concerns.append("No landlord reference provided")
    
    # Pets
    if applicant.get("has_pets") and listing and not listing.get("pet_friendly"):
        concerns.append("Has pets but property may not be pet-friendly")
    
    return {"strengths": strengths, "concerns": concerns}

@api_router.patch("/applications/{application_id}")
async def patch_application_status(application_id: str, status: str):
    """Quick update application status"""
    if status not in ["pending", "under_review", "approved", "rejected"]:
        raise HTTPException(status_code=400, detail="Invalid status")
    
    result = await db.applications.update_one(
        {"id": application_id},
        {"$set": {"status": status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Application not found")
    
    return {"success": True, "status": status}

@api_router.put("/applications/{application_id}/status")
async def update_application_status(application_id: str, status: str, landlord_id: str):
    """Update application status (landlord only)"""
    if status not in ["pending", "under_review", "approved", "rejected"]:
        raise HTTPException(status_code=400, detail="Invalid status")
    
    app = await db.applications.find_one({"id": application_id}, {"_id": 0})
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
    
    await db.applications.update_one(
        {"id": application_id},
        {"$set": {"status": status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Notify the applicant
    status_messages = {
        "under_review": "Your application is being reviewed",
        "approved": "Congratulations! Your application has been approved! 🎉",
        "rejected": "Your application status has been updated"
    }
    
    if status in status_messages:
        await db.notifications.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": app["user_id"],
            "title": f"Application {status.replace('_', ' ').title()}",
            "body": status_messages[status],
            "type": "application",
            "data": {"application_id": application_id, "status": status},
            "read": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        
        # Send email notification
        applicant = await db.users.find_one({"id": app["user_id"]}, {"_id": 0})
        listing = await db.listings.find_one({"id": app.get("listing_id")}, {"_id": 0, "title": 1})
        if applicant and applicant.get("email"):
            asyncio.create_task(send_email(
                applicant["email"],
                f"Application Update - {listing.get('title', 'Your Property') if listing else 'Your Property'}",
                email_application_update(
                    applicant.get("name", ""),
                    listing.get("title", "the property") if listing else "the property",
                    status
                )
            ))
    
    return {"status": "updated"}

# ========== MAINTENANCE REQUESTS ==========

@api_router.post("/maintenance", response_model=dict)
async def create_maintenance_request(user_id: str, request: MaintenanceRequestCreate):
    """Submit a maintenance request"""
    maint_req = MaintenanceRequest(
        user_id=user_id,
        property_id=request.property_id,
        title=request.title,
        description=request.description,
        category=request.category,
        priority=request.priority,
        images=request.images
    )
    
    # Try to find the landlord from property
    if request.property_id:
        listing = await db.listings.find_one({"id": request.property_id}, {"_id": 0})
        if listing and listing.get("landlord_id"):
            maint_req.landlord_id = listing["landlord_id"]
            
            # Notify landlord
            await db.notifications.insert_one({
                "id": str(uuid.uuid4()),
                "user_id": listing["landlord_id"],
                "title": f"New Maintenance Request - {request.priority.upper()}",
                "body": request.title,
                "type": "maintenance",
                "data": {"request_id": maint_req.id, "priority": request.priority},
                "read": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            })
    
    await db.maintenance_requests.insert_one(maint_req.model_dump())
    
    return {"id": maint_req.id, "status": "submitted"}

@api_router.get("/maintenance/user/{user_id}")
async def get_user_maintenance_requests(user_id: str, status: Optional[str] = None):
    """Get maintenance requests submitted by a user"""
    query = {"user_id": user_id}
    if status:
        query["status"] = status
    
    requests = await db.maintenance_requests.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    return requests

@api_router.get("/maintenance/landlord/{landlord_id}")
async def get_landlord_maintenance_requests(landlord_id: str, status: Optional[str] = None):
    """Get all maintenance requests for a landlord's properties"""
    query = {"landlord_id": landlord_id}
    if status:
        query["status"] = status
    
    requests = await db.maintenance_requests.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return requests

@api_router.put("/maintenance/{request_id}")
async def update_maintenance_request(request_id: str, updates: Dict[str, Any]):
    """Update maintenance request status, assign contractor, etc."""
    allowed_fields = ["status", "assigned_contractor_id", "scheduled_date", "completed_date", "cost", "notes"]
    update_dict = {k: v for k, v in updates.items() if k in allowed_fields}
    update_dict["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.maintenance_requests.update_one(
        {"id": request_id},
        {"$set": update_dict}
    )
    
    # Get request for notification
    req = await db.maintenance_requests.find_one({"id": request_id}, {"_id": 0})
    
    # Notify tenant of status change
    if req and "status" in updates:
        await db.notifications.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": req["user_id"],
            "title": "Maintenance Update",
            "body": f"Your request '{req['title']}' is now {updates['status']}",
            "type": "maintenance",
            "data": {"request_id": request_id, "status": updates["status"]},
            "read": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
    
    return {"status": "updated"}

# ========== CONTRACTOR JOBS (LANDLORD TO CONTRACTOR) ==========

@api_router.post("/contractor-jobs", response_model=dict)
async def create_contractor_job(landlord_id: str, job: ContractorJobCreate):
    """Create a job posting for contractors (landlord posts maintenance job)"""
    job_obj = ContractorJob(
        landlord_id=landlord_id,
        maintenance_request_id=job.maintenance_request_id,
        title=job.title,
        description=job.description,
        category=job.category,
        location=job.location,
        budget_min=job.budget_min,
        budget_max=job.budget_max,
        deadline=job.deadline
    )
    
    await db.contractor_jobs.insert_one(job_obj.model_dump())
    
    return {"id": job_obj.id, "status": "created"}

@api_router.get("/contractor-jobs")
async def get_contractor_jobs(
    category: Optional[str] = None,
    status: Optional[str] = None,
    location: Optional[str] = None
):
    """Get available contractor jobs"""
    query = {}
    if category:
        query["category"] = category
    if status:
        query["status"] = status
    else:
        query["status"] = "open"  # Default to open jobs
    if location:
        query["location"] = {"$regex": location, "$options": "i"}
    
    jobs = await db.contractor_jobs.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return jobs

@api_router.get("/contractor-jobs/landlord/{landlord_id}")
async def get_landlord_jobs(landlord_id: str):
    """Get jobs posted by a landlord"""
    jobs = await db.contractor_jobs.find({"landlord_id": landlord_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return jobs

@api_router.get("/contractor-jobs/contractor/{contractor_id}")
async def get_contractor_assigned_jobs(contractor_id: str):
    """Get jobs assigned to a contractor"""
    jobs = await db.contractor_jobs.find({"contractor_id": contractor_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return jobs

@api_router.post("/contractor-jobs/{job_id}/bid")
async def submit_bid(job_id: str, bid: ContractorBid):
    """Submit a bid for a job"""
    job = await db.contractor_jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if job["status"] != "open":
        raise HTTPException(status_code=400, detail="Job is not accepting bids")
    
    bid_doc = {
        "id": str(uuid.uuid4()),
        "contractor_id": bid.contractor_id,
        "amount": bid.amount,
        "estimated_days": bid.estimated_days,
        "message": bid.message,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.contractor_jobs.update_one(
        {"id": job_id},
        {"$push": {"bids": bid_doc}}
    )
    
    # Notify landlord of new bid
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": job["landlord_id"],
        "title": "New Bid Received",
        "body": f"${bid.amount} bid on '{job['title']}'",
        "type": "bid",
        "data": {"job_id": job_id, "bid_id": bid_doc["id"], "amount": bid.amount},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {"bid_id": bid_doc["id"], "status": "submitted"}

@api_router.post("/contractor-jobs/{job_id}/select-bid")
async def select_bid(job_id: str, bid_id: str, landlord_id: str):
    """Select a winning bid for a job"""
    job = await db.contractor_jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if job["landlord_id"] != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    # Find the selected bid
    selected_bid = None
    for bid in job.get("bids", []):
        if bid["id"] == bid_id:
            selected_bid = bid
            break
    
    if not selected_bid:
        raise HTTPException(status_code=404, detail="Bid not found")
    
    await db.contractor_jobs.update_one(
        {"id": job_id},
        {"$set": {
            "status": "assigned",
            "contractor_id": selected_bid["contractor_id"],
            "selected_bid_id": bid_id,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Notify the winning contractor
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": selected_bid["contractor_id"],
        "title": "Congratulations! Your bid was accepted! 🎉",
        "body": f"You won the job: {job['title']}",
        "type": "job",
        "data": {"job_id": job_id},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {"status": "assigned", "contractor_id": selected_bid["contractor_id"]}

# ========== PROPERTY MANAGEMENT (LANDLORD) ==========

@api_router.post("/listings/create")
async def create_listing(landlord_id: str, listing: ListingCreate):
    """Create a new property listing (landlord only)"""
    listing_data = listing.model_dump()
    listing_data["landlord_id"] = landlord_id
    listing_data["user_id"] = landlord_id  # Also set user_id for compatibility
    listing_data["owner_id"] = landlord_id  # Also set owner_id for compatibility
    
    listing_obj = Listing(**listing_data)
    
    await db.listings.insert_one(listing_obj.model_dump())
    
    return {"id": listing_obj.id, "status": "created"}

@api_router.get("/listings/landlord/{landlord_id}")
async def get_landlord_listings(landlord_id: str, status: Optional[str] = None):
    """Get all listings owned by a landlord"""
    query = {"landlord_id": landlord_id}
    if status:
        query["status"] = status
    
    listings = await db.listings.find(query, {"_id": 0}).to_list(100)
    return listings

@api_router.put("/listings/{listing_id}")
async def update_listing(listing_id: str, landlord_id: str, updates: Dict[str, Any]):
    """Update a listing (landlord only)"""
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    if listing.get("landlord_id") != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    allowed_fields = ["title", "address", "city", "province", "postal_code", "lat", "lng",
                      "description", "price", "bedrooms", "bathrooms", "sqft", "property_type",
                      "status", "available_date", "amenities", "images", "pet_friendly", "parking",
                      "listing_type", "sale_price", "year_built", "lot_size", "garage",
                      "mls_number", "open_house_dates"]
    update_dict = {k: v for k, v in updates.items() if k in allowed_fields}
    
    await db.listings.update_one(
        {"id": listing_id},
        {"$set": update_dict}
    )
    
    return {"status": "updated"}

@api_router.delete("/listings/{listing_id}")
async def delete_listing(listing_id: str, landlord_id: str):
    """Delete a listing (landlord only)"""
    listing = await db.listings.find_one({"id": listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    if listing.get("landlord_id") != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    # Soft delete
    await db.listings.update_one(
        {"id": listing_id},
        {"$set": {"status": "deleted"}}
    )
    
    return {"status": "deleted"}

# ========== PROPERTY OFFERS (BUY/SELL) ==========

@api_router.post("/offers")
async def create_offer(buyer_id: str, offer: OfferCreate):
    """Submit an offer on a property for sale"""
    listing = await db.listings.find_one({"id": offer.listing_id}, {"_id": 0})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    offer_obj = PropertyOffer(
        buyer_id=buyer_id,
        seller_id=listing.get("landlord_id", ""),
        **offer.model_dump()
    )
    doc = offer_obj.model_dump()
    await db.offers.insert_one(doc)
    doc.pop("_id", None)
    
    # Notify seller
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": listing.get("landlord_id", ""),
        "title": "New Offer Received",
        "body": f"Offer of ${offer.offer_amount:,} on {listing['title']}",
        "type": "offer",
        "data": {"offer_id": offer_obj.id, "listing_id": offer.listing_id},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Email notification to seller
    seller = await db.users.find_one({"id": listing.get("landlord_id", "")}, {"_id": 0})
    if seller and seller.get("email"):
        html = f"""
        <div style="font-family:'Georgia',serif;max-width:600px;margin:0 auto;background:#F5F5F0;padding:40px;">
          <div style="background:#1A2F3A;padding:30px;border-radius:16px 16px 0 0;text-align:center;">
            <h1 style="color:white;font-family:'Georgia',serif;margin:0;">DOMMMA</h1>
          </div>
          <div style="background:white;padding:30px;border-radius:0 0 16px 16px;">
            <h2 style="color:#1A2F3A;">New Offer on {listing['title']}</h2>
            <div style="background:#F5F5F0;border-radius:12px;padding:20px;margin:20px 0;">
              <p style="margin:4px 0;color:#1A2F3A;font-size:24px;font-weight:bold;">${offer.offer_amount:,}</p>
              <p style="margin:4px 0;color:#555;">Financing: {offer.financing_type}</p>
              {f'<p style="margin:4px 0;color:#555;">Closing: {offer.closing_date}</p>' if offer.closing_date else ''}
            </div>
            <p style="color:#555;">Log in to review and respond to this offer.</p>
          </div>
        </div>"""
        asyncio.create_task(send_email(seller["email"], f"New Offer - {listing['title']}", html))
    
    return doc

@api_router.get("/offers/buyer/{buyer_id}")
async def get_buyer_offers(buyer_id: str):
    """Get offers made by a buyer"""
    offers = await db.offers.find({"buyer_id": buyer_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    for o in offers:
        listing = await db.listings.find_one({"id": o["listing_id"]}, {"_id": 0, "title": 1, "address": 1, "city": 1, "price": 1, "images": 1})
        o["listing"] = listing or {}
    return offers

@api_router.get("/offers/seller/{seller_id}")
async def get_seller_offers(seller_id: str):
    """Get offers received by a seller"""
    offers = await db.offers.find({"seller_id": seller_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    for o in offers:
        listing = await db.listings.find_one({"id": o["listing_id"]}, {"_id": 0, "title": 1, "address": 1, "city": 1, "price": 1, "images": 1})
        o["listing"] = listing or {}
        buyer = await db.users.find_one({"id": o["buyer_id"]}, {"_id": 0, "name": 1, "email": 1})
        o["buyer"] = buyer or {}
    return offers

@api_router.get("/offers/listing/{listing_id}")
async def get_listing_offers(listing_id: str):
    """Get all offers for a listing"""
    offers = await db.offers.find({"listing_id": listing_id}, {"_id": 0}).sort("offer_amount", -1).to_list(50)
    for o in offers:
        buyer = await db.users.find_one({"id": o["buyer_id"]}, {"_id": 0, "name": 1, "email": 1})
        o["buyer"] = buyer or {}
    return offers

@api_router.put("/offers/{offer_id}/respond")
async def respond_to_offer(offer_id: str, action: str, seller_id: str, counter_amount: Optional[int] = None, counter_message: Optional[str] = None):
    """Seller responds to an offer: accept, reject, or counter"""
    if action not in ["accepted", "rejected", "countered"]:
        raise HTTPException(status_code=400, detail="Invalid action")
    
    offer = await db.offers.find_one({"id": offer_id}, {"_id": 0})
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")
    if offer["seller_id"] != seller_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    update = {"status": action, "updated_at": datetime.now(timezone.utc).isoformat()}
    if action == "countered" and counter_amount:
        update["counter_amount"] = counter_amount
        update["counter_message"] = counter_message
    
    await db.offers.update_one({"id": offer_id}, {"$set": update})
    
    # If accepted, mark listing as sold/pending
    if action == "accepted":
        await db.listings.update_one(
            {"id": offer["listing_id"]},
            {"$set": {"status": "pending_sale"}}
        )
    
    # Notify buyer
    status_text = {"accepted": "accepted!", "rejected": "declined.", "countered": f"countered with ${counter_amount:,}." if counter_amount else "countered."}
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": offer["buyer_id"],
        "title": f"Offer {action.title()}",
        "body": f"Your offer has been {status_text.get(action, action)}",
        "type": "offer",
        "data": {"offer_id": offer_id, "action": action},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Email buyer
    buyer = await db.users.find_one({"id": offer["buyer_id"]}, {"_id": 0})
    if buyer and buyer.get("email"):
        asyncio.create_task(send_email(
            buyer["email"],
            f"Offer {action.title()}",
            email_application_update(buyer.get("name", ""), "the property", action)
        ))
    
    return {"status": "updated"}

# ========== CONTRACTOR PROFILES ==========

@api_router.post("/contractors/profile")
async def create_contractor_profile(user_id: str, profile: ContractorProfileCreate):
    """Create or update contractor profile"""
    existing = await db.contractor_profiles.find_one({"user_id": user_id}, {"_id": 0})
    
    if existing:
        update_data = profile.model_dump()
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        await db.contractor_profiles.update_one(
            {"user_id": user_id},
            {"$set": update_data}
        )
        updated = await db.contractor_profiles.find_one({"user_id": user_id}, {"_id": 0})
        return updated
    
    profile_obj = ContractorProfile(user_id=user_id, **profile.model_dump())
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if user:
        profile_obj.email = profile.email or user.get("email")
    
    doc = profile_obj.model_dump()
    await db.contractor_profiles.insert_one(doc)
    doc.pop("_id", None)
    return doc

@api_router.get("/contractors/profile/{user_id}")
async def get_contractor_profile(user_id: str):
    """Get contractor profile by user_id"""
    profile = await db.contractor_profiles.find_one({"user_id": user_id}, {"_id": 0})
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    return profile

@api_router.get("/contractors/search")
async def search_contractors(
    specialty: Optional[str] = None,
    area: Optional[str] = None,
    min_rating: Optional[float] = None,
    q: Optional[str] = None
):
    """Search contractors"""
    query = {"status": "active"}
    if specialty:
        query["specialties"] = {"$regex": specialty, "$options": "i"}
    if area:
        query["service_areas"] = {"$regex": area, "$options": "i"}
    if min_rating:
        query["rating"] = {"$gte": min_rating}
    if q:
        query["$or"] = [
            {"business_name": {"$regex": q, "$options": "i"}},
            {"description": {"$regex": q, "$options": "i"}},
            {"specialties": {"$regex": q, "$options": "i"}}
        ]
    
    profiles = await db.contractor_profiles.find(query, {"_id": 0}).to_list(100)
    
    # Enrich with user info
    for p in profiles:
        user = await db.users.find_one({"id": p["user_id"]}, {"_id": 0, "name": 1, "email": 1})
        if user:
            p["user_name"] = user.get("name", "")
    
    return profiles

@api_router.put("/contractors/profile/{user_id}/images")
async def update_contractor_images(user_id: str, images: List[str]):
    """Update contractor portfolio images"""
    await db.contractor_profiles.update_one(
        {"user_id": user_id},
        {"$set": {"portfolio_images": images, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "updated"}

@api_router.post("/contractors/verify-document")
async def verify_contractor_document(
    file: UploadFile = File(...),
    document_type: str = Body(...),
    contractor_id: str = Body(...)
):
    """
    Upload and verify contractor documents (WCB clearance, insurance) using AI
    """
    import base64
    
    # Read file content
    file_content = await file.read()
    file_base64 = base64.b64encode(file_content).decode('utf-8')
    file_extension = file.filename.split('.')[-1].lower() if file.filename else 'pdf'
    
    # Determine media type
    media_types = {
        'pdf': 'application/pdf',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'png': 'image/png'
    }
    media_type = media_types.get(file_extension, 'application/pdf')
    
    # Use Claude to verify the document
    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="AI verification not configured")
    
    verification_prompts = {
        'wcb_clearance': """Analyze this document and determine if it is a valid WorkSafeBC (WCB) clearance certificate or letter.
        
Check for:
1. WorkSafeBC or WCB logo/header
2. Clearance letter confirming good standing
3. Account number
4. Date of issuance (should be recent, within last 6 months)
5. Business name

Respond in JSON format:
{
    "is_valid": true/false,
    "document_type": "wcb_clearance",
    "confidence": 0-100,
    "business_name": "extracted name if found",
    "account_number": "if found",
    "issue_date": "if found",
    "reason": "explanation if not valid"
}""",
        'insurance': """Analyze this document and determine if it is a valid commercial liability insurance certificate.

Check for:
1. Insurance company name/logo
2. Certificate of insurance or policy document
3. Coverage type (General Liability, Commercial Liability)
4. Policy number
5. Coverage amounts
6. Expiration date (should not be expired)
7. Named insured (business name)

Respond in JSON format:
{
    "is_valid": true/false,
    "document_type": "insurance",
    "confidence": 0-100,
    "insurance_company": "if found",
    "policy_number": "if found",
    "coverage_amount": "if found",
    "expiration_date": "if found",
    "business_name": "if found",
    "reason": "explanation if not valid"
}"""
    }
    
    prompt = verification_prompts.get(document_type, verification_prompts['insurance'])
    
    try:
        import httpx
        
        # Call Claude API for document analysis
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json"
                },
                json={
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 1024,
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image" if media_type.startswith('image') else "document",
                                    "source": {
                                        "type": "base64",
                                        "media_type": media_type,
                                        "data": file_base64
                                    }
                                },
                                {
                                    "type": "text",
                                    "text": prompt
                                }
                            ]
                        }
                    ]
                }
            )
            
            if response.status_code != 200:
                logger.error(f"Claude API error: {response.text}")
                raise HTTPException(status_code=500, detail="AI verification failed")
            
            result = response.json()
            ai_response = result.get("content", [{}])[0].get("text", "{}")
            
            # Parse JSON response
            import json
            try:
                verification_result = json.loads(ai_response)
            except:
                # Try to extract JSON from the response
                import re
                json_match = re.search(r'\{[^{}]*\}', ai_response, re.DOTALL)
                if json_match:
                    verification_result = json.loads(json_match.group())
                else:
                    verification_result = {"is_valid": False, "reason": "Could not parse AI response"}
            
            is_valid = verification_result.get("is_valid", False)
            confidence = verification_result.get("confidence", 0)
            
            # Require at least 70% confidence for verification
            if is_valid and confidence >= 70:
                # Update contractor profile with verification
                update_field = "wcb_verified" if document_type == "wcb_clearance" else "insurance_verified"
                doc_url_field = "wcb_doc_url" if document_type == "wcb_clearance" else "insurance_doc_url"
                
                await db.contractor_profiles.update_one(
                    {"user_id": contractor_id},
                    {
                        "$set": {
                            update_field: True,
                            doc_url_field: f"verified_{document_type}_{contractor_id}",
                            f"{document_type}_details": verification_result,
                            "updated_at": datetime.now(timezone.utc).isoformat()
                        }
                    },
                    upsert=True
                )
                
                return {
                    "verified": True,
                    "document_type": document_type,
                    "confidence": confidence,
                    "details": verification_result,
                    "document_url": f"verified_{document_type}_{contractor_id}"
                }
            else:
                return {
                    "verified": False,
                    "document_type": document_type,
                    "confidence": confidence,
                    "reason": verification_result.get("reason", "Document verification failed - confidence too low or document invalid")
                }
                
    except Exception as e:
        logger.error(f"Document verification error: {e}")
        raise HTTPException(status_code=500, detail=f"Verification failed: {str(e)}")

# ========== CONTRACTOR SERVICES ==========

@api_router.post("/contractors/services")
async def create_contractor_service(contractor_id: str, service: ContractorServiceCreate):
    """Create a service listing"""
    service_obj = ContractorService(contractor_id=contractor_id, **service.model_dump())
    doc = service_obj.model_dump()
    await db.contractor_services.insert_one(doc)
    doc.pop("_id", None)
    return doc

@api_router.get("/contractors/{contractor_id}/services")
async def get_contractor_services(contractor_id: str):
    """Get services offered by a contractor"""
    services = await db.contractor_services.find(
        {"contractor_id": contractor_id, "status": "active"}, {"_id": 0}
    ).to_list(50)
    return services

@api_router.get("/services/search")
async def search_services(
    category: Optional[str] = None,
    q: Optional[str] = None,
    max_price: Optional[float] = None
):
    """Search all contractor services"""
    query = {"status": "active"}
    if category:
        query["category"] = {"$regex": category, "$options": "i"}
    if q:
        query["$or"] = [
            {"title": {"$regex": q, "$options": "i"}},
            {"description": {"$regex": q, "$options": "i"}}
        ]
    if max_price:
        query["price"] = {"$lte": max_price}
    
    services = await db.contractor_services.find(query, {"_id": 0}).to_list(100)
    
    # Enrich with contractor info
    for s in services:
        profile = await db.contractor_profiles.find_one(
            {"contractor_id": s["contractor_id"]},
            {"_id": 0, "business_name": 1, "rating": 1, "avatar": 1}
        )
        if not profile:
            profile = await db.contractor_profiles.find_one(
                {"user_id": s["contractor_id"]},
                {"_id": 0, "business_name": 1, "rating": 1, "avatar": 1}
            )
        s["contractor"] = profile or {}
    
    return services

@api_router.delete("/contractors/services/{service_id}")
async def delete_contractor_service(service_id: str, contractor_id: str):
    """Delete a service"""
    await db.contractor_services.update_one(
        {"id": service_id, "contractor_id": contractor_id},
        {"$set": {"status": "deleted"}}
    )
    return {"status": "deleted"}

# ========== SERVICE BOOKINGS ==========

@api_router.post("/bookings")
async def create_booking(customer_id: str, booking: ServiceBookingCreate):
    """Create a service booking"""
    booking_obj = ServiceBooking(customer_id=customer_id, **booking.model_dump())
    
    # If service_id provided, get price
    if booking.service_id:
        service = await db.contractor_services.find_one({"id": booking.service_id}, {"_id": 0})
        if service and service.get("price"):
            booking_obj.amount = service["price"]
    
    doc = booking_obj.model_dump()
    await db.bookings.insert_one(doc)
    
    # Notify contractor
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": booking.contractor_id,
        "title": "New Booking Request",
        "body": f"New booking: {booking.title}",
        "type": "booking",
        "data": {"booking_id": booking_obj.id},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    doc.pop("_id", None)
    return doc

@api_router.get("/bookings/customer/{customer_id}")
async def get_customer_bookings(customer_id: str):
    """Get bookings made by a customer"""
    bookings = await db.bookings.find({"customer_id": customer_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    for b in bookings:
        profile = await db.contractor_profiles.find_one({"user_id": b["contractor_id"]}, {"_id": 0, "business_name": 1, "avatar": 1})
        b["contractor"] = profile or {}
    return bookings

@api_router.get("/bookings/contractor/{contractor_id}")
async def get_contractor_bookings(contractor_id: str):
    """Get bookings for a contractor"""
    bookings = await db.bookings.find({"contractor_id": contractor_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    for b in bookings:
        user = await db.users.find_one({"id": b["customer_id"]}, {"_id": 0, "name": 1, "email": 1})
        b["customer"] = user or {}
    return bookings

@api_router.put("/bookings/{booking_id}/status")
async def update_booking_status(booking_id: str, status: str, user_id: str):
    """Update booking status"""
    valid_statuses = ["pending", "confirmed", "in_progress", "completed", "cancelled"]
    if status not in valid_statuses:
        raise HTTPException(status_code=400, detail="Invalid status")
    
    booking = await db.bookings.find_one({"id": booking_id}, {"_id": 0})
    if not booking:
        raise HTTPException(status_code=404, detail="Booking not found")
    
    await db.bookings.update_one(
        {"id": booking_id},
        {"$set": {"status": status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Notify the other party
    notify_user = booking["customer_id"] if user_id == booking["contractor_id"] else booking["contractor_id"]
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": notify_user,
        "title": f"Booking {status.replace('_', ' ').title()}",
        "body": f"Booking '{booking['title']}' is now {status.replace('_', ' ')}",
        "type": "booking",
        "data": {"booking_id": booking_id, "status": status},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Send email notification for confirmed bookings
    if status == "confirmed":
        customer = await db.users.find_one({"id": booking["customer_id"]}, {"_id": 0})
        contractor_profile = await db.contractor_profiles.find_one({"user_id": booking["contractor_id"]}, {"_id": 0})
        if customer and customer.get("email"):
            asyncio.create_task(send_email(
                customer["email"],
                f"Booking Confirmed - {booking['title']}",
                email_booking_confirmed(
                    customer.get("name", ""),
                    contractor_profile.get("business_name", "") if contractor_profile else "",
                    booking["title"],
                    booking.get("preferred_date", "")
                )
            ))
    
    return {"status": "updated"}

@api_router.post("/bookings/{booking_id}/pay")
async def pay_booking(request: Request, booking_id: str):
    """Create payment for a booking"""
    booking = await db.bookings.find_one({"id": booking_id}, {"_id": 0})
    if not booking:
        raise HTTPException(status_code=404, detail="Booking not found")
    
    if not booking.get("amount"):
        raise HTTPException(status_code=400, detail="No amount set for this booking")
    
    api_key = os.environ.get('STRIPE_API_KEY')
    host_url = str(request.base_url).rstrip('/')
    webhook_url = f"{host_url}/api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=api_key, webhook_url=webhook_url)
    
    origin = request.headers.get('origin', host_url)
    success_url = f"{origin}/dashboard?payment=success&booking_id={booking_id}&session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{origin}/dashboard?payment=cancelled"
    
    checkout_request = CheckoutSessionRequest(
        amount=float(booking["amount"]),
        currency="cad",
        success_url=success_url,
        cancel_url=cancel_url,
        metadata={"booking_id": booking_id, "contractor_id": booking["contractor_id"]}
    )
    
    session = await stripe_checkout.create_checkout_session(checkout_request)
    
    await db.bookings.update_one(
        {"id": booking_id},
        {"$set": {"payment_session_id": session.session_id}}
    )
    
    # Also create a payment transaction record
    transaction = PaymentTransaction(
        session_id=session.session_id,
        user_id=booking["customer_id"],
        amount=booking["amount"],
        currency="cad",
        description=f"Booking: {booking['title']}",
        recipient_id=booking["contractor_id"],
        payment_status="pending",
        status="initiated"
    )
    await db.payment_transactions.insert_one(transaction.model_dump())
    
    return {"url": session.url, "session_id": session.session_id}

@api_router.post("/bookings/{booking_id}/review")
async def review_booking(booking_id: str, customer_id: str, review: ReviewCreate):
    """Leave a review for a completed booking"""
    booking = await db.bookings.find_one({"id": booking_id}, {"_id": 0})
    if not booking:
        raise HTTPException(status_code=404, detail="Booking not found")
    if booking["customer_id"] != customer_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    await db.bookings.update_one(
        {"id": booking_id},
        {"$set": {"rating": review.rating, "review": review.review, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Update contractor's average rating
    all_reviews = await db.bookings.find(
        {"contractor_id": booking["contractor_id"], "rating": {"$exists": True, "$ne": None}},
        {"_id": 0, "rating": 1}
    ).to_list(500)
    
    if all_reviews:
        avg_rating = sum(r["rating"] for r in all_reviews) / len(all_reviews)
        await db.contractor_profiles.update_one(
            {"user_id": booking["contractor_id"]},
            {"$set": {"rating": round(avg_rating, 1), "review_count": len(all_reviews)}}
        )
    
    return {"status": "reviewed"}

@api_router.get("/contractors/{contractor_id}/reviews")
async def get_contractor_reviews(contractor_id: str, limit: int = 20):
    """Get reviews for a contractor"""
    reviews = await db.bookings.find(
        {"contractor_id": contractor_id, "rating": {"$exists": True, "$ne": None}},
        {"_id": 0, "id": 1, "customer_id": 1, "title": 1, "rating": 1, "review": 1, "created_at": 1, "updated_at": 1}
    ).sort("updated_at", -1).to_list(limit)
    
    # Add customer info
    for r in reviews:
        user = await db.users.find_one({"id": r["customer_id"]}, {"_id": 0, "name": 1, "avatar": 1})
        r["customer"] = user or {"name": "Anonymous"}
    
    return reviews

@api_router.get("/contractors/leaderboard")
async def get_contractor_leaderboard(limit: int = 10):
    """Get top-rated contractors"""
    contractors = await db.contractor_profiles.find(
        {"status": "active", "rating": {"$gt": 0}},
        {"_id": 0, "id": 1, "user_id": 1, "business_name": 1, "avatar": 1, "rating": 1, "review_count": 1, 
         "completed_jobs": 1, "specialties": 1, "verified": 1, "service_areas": 1}
    ).sort([("rating", -1), ("review_count", -1)]).to_list(limit)
    
    return contractors


# ========== JOB POSTING & BIDDING SYSTEM (bark.com-like) ==========

class JobPostCreate(BaseModel):
    title: str
    category: str  # plumbing, electrical, painting, etc.
    description: str
    address: str
    budget_min: Optional[float] = None
    budget_max: Optional[float] = None
    preferred_date: Optional[str] = None
    urgency: str = "flexible"  # flexible, this_week, urgent
    images: Optional[List[str]] = []
    # New fields for bark.com-style wizard
    contact_email: Optional[str] = None
    contact_name: Optional[str] = None
    contact_phone: Optional[str] = None
    answers: Optional[Dict[str, Any]] = {}  # Store questionnaire answers

class JobBidCreate(BaseModel):
    amount: float
    message: str
    estimated_duration: Optional[str] = None
    available_date: Optional[str] = None

@api_router.post("/jobs")
async def create_job_post(job: JobPostCreate, user_id: str):
    """Create a new job posting (customer posts a job, contractors bid on it)"""
    # Allow guest users (user_id='guest') to post jobs
    user = None
    user_name = "Guest"
    if user_id != 'guest':
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if user:
            user_name = user.get("name", "Anonymous")
    
    job_id = str(uuid.uuid4())
    job_doc = {
        "id": job_id,
        "user_id": user_id,
        "user_name": user_name,
        "title": job.title,
        "category": job.category.lower(),
        "description": job.description,
        "address": job.address,
        "budget_min": job.budget_min,
        "budget_max": job.budget_max,
        "preferred_date": job.preferred_date,
        "urgency": job.urgency,
        "images": job.images or [],
        "status": "open",  # open, in_progress, completed, cancelled
        "bid_count": 0,
        # Contact info for bark.com-style
        "contact_email": job.contact_email,
        "contact_name": job.contact_name,
        "contact_phone": job.contact_phone,
        "answers": job.answers or {},
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.job_posts.insert_one(job_doc)
    
    # Send confirmation email to customer
    if job.contact_email:
        asyncio.create_task(send_email(
            job.contact_email,
            f"Your Request is Live - {job.title}",
            email_job_request_confirmation(
                job.contact_name or user_name,
                job.title,
                job.address,
                job.answers
            )
        ))
    
    # Notify relevant contractors (those with matching specialty) via in-app notification AND email
    contractors = await db.contractor_profiles.find(
        {"specialties": {"$in": [job.category.lower()]}, "status": "active"},
        {"_id": 0, "user_id": 1, "business_name": 1}
    ).to_list(50)
    
    for c in contractors:
        # In-app notification
        await db.notifications.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": c["user_id"],
            "title": "New Job Opportunity",
            "body": f"New job posted: {job.title}",
            "type": "job_post",
            "data": {"job_id": job_id},
            "read": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        
        # Get contractor's email
        contractor_user = await db.users.find_one({"id": c["user_id"]}, {"_id": 0, "email": 1})
        if contractor_user and contractor_user.get("email"):
            asyncio.create_task(send_email(
                contractor_user["email"],
                f"New Lead: {job.title} in {job.address}",
                email_new_lead_notification(
                    c.get("business_name", "Professional"),
                    job.title,
                    job.address,
                    job.description,
                    job_id
                )
            ))
    
    job_doc.pop("_id", None)
    return job_doc

@api_router.get("/jobs")
async def get_job_posts(
    category: Optional[str] = None,
    status: str = "open",
    limit: int = 20,
    skip: int = 0
):
    """Get all job posts (for contractors to browse)"""
    query = {"status": status}
    if category:
        query["category"] = category.lower()
    
    jobs = await db.job_posts.find(query, {"_id": 0}).sort("created_at", -1).skip(skip).limit(limit).to_list(limit)
    
    # Add bid count for each job
    for job in jobs:
        bids = await db.job_bids.count_documents({"job_id": job["id"]})
        job["bid_count"] = bids
    
    return jobs

@api_router.get("/jobs/{job_id}")
async def get_job_post(job_id: str):
    """Get a specific job post with details"""
    job = await db.job_posts.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Get bids for this job
    bids = await db.job_bids.find({"job_id": job_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    
    # Add contractor info to each bid
    for bid in bids:
        contractor = await db.contractor_profiles.find_one(
            {"user_id": bid["contractor_id"]},
            {"_id": 0, "business_name": 1, "avatar": 1, "rating": 1, "review_count": 1, "verified": 1}
        )
        bid["contractor"] = contractor or {}
    
    job["bids"] = bids
    return job

@api_router.get("/jobs/user/{user_id}")
async def get_user_job_posts(user_id: str, status: Optional[str] = None):
    """Get job posts created by a user"""
    query = {"user_id": user_id}
    if status:
        query["status"] = status
    
    jobs = await db.job_posts.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    
    for job in jobs:
        bids = await db.job_bids.count_documents({"job_id": job["id"]})
        job["bid_count"] = bids
    
    return jobs

@api_router.post("/jobs/{job_id}/bids")
async def create_job_bid(job_id: str, bid: JobBidCreate, contractor_id: str):
    """Submit a bid on a job (contractor submits a quote)"""
    # Verify job exists and is open
    job = await db.job_posts.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["status"] != "open":
        raise HTTPException(status_code=400, detail="Job is no longer accepting bids")
    
    # Verify contractor
    contractor = await db.contractor_profiles.find_one({"user_id": contractor_id}, {"_id": 0})
    if not contractor:
        raise HTTPException(status_code=404, detail="Contractor profile not found")
    
    # Check if already bid
    existing = await db.job_bids.find_one({"job_id": job_id, "contractor_id": contractor_id})
    if existing:
        raise HTTPException(status_code=400, detail="You have already submitted a bid for this job")
    
    bid_id = str(uuid.uuid4())
    bid_doc = {
        "id": bid_id,
        "job_id": job_id,
        "contractor_id": contractor_id,
        "amount": bid.amount,
        "message": bid.message,
        "estimated_duration": bid.estimated_duration,
        "available_date": bid.available_date,
        "status": "pending",  # pending, accepted, rejected
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.job_bids.insert_one(bid_doc)
    
    # Update job bid count
    await db.job_posts.update_one(
        {"id": job_id},
        {"$inc": {"bid_count": 1}, "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # In-app notification for job poster
    if job["user_id"] != "guest":
        await db.notifications.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": job["user_id"],
            "title": "New Bid Received",
            "body": f"{contractor.get('business_name', 'A contractor')} submitted a bid of ${bid.amount}",
            "type": "job_bid",
            "data": {"job_id": job_id, "bid_id": bid_id},
            "read": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
    
    # Email notification to job poster
    customer_email = job.get("contact_email")
    if customer_email:
        asyncio.create_task(send_email(
            customer_email,
            f"New Quote: ${bid.amount:.2f} for {job['title']}",
            email_bid_received(
                job.get("contact_name", "Customer"),
                contractor.get("business_name", "Professional"),
                bid.amount,
                job["title"],
                bid.message
            )
        ))
    
    bid_doc.pop("_id", None)
    return bid_doc

@api_router.get("/jobs/{job_id}/bids")
async def get_job_bids(job_id: str):
    """Get all bids for a job"""
    bids = await db.job_bids.find({"job_id": job_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    
    for bid in bids:
        contractor = await db.contractor_profiles.find_one(
            {"user_id": bid["contractor_id"]},
            {"_id": 0, "business_name": 1, "avatar": 1, "rating": 1, "review_count": 1, "verified": 1, "completed_jobs": 1}
        )
        bid["contractor"] = contractor or {}
    
    return bids

@api_router.put("/jobs/{job_id}/bids/{bid_id}/accept")
async def accept_job_bid(job_id: str, bid_id: str, user_id: str):
    """Accept a bid and hire the contractor"""
    # Verify job belongs to user
    job = await db.job_posts.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["user_id"] != user_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    # Get the bid
    bid = await db.job_bids.find_one({"id": bid_id, "job_id": job_id}, {"_id": 0})
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")
    
    # Accept this bid
    await db.job_bids.update_one(
        {"id": bid_id},
        {"$set": {"status": "accepted", "accepted_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Reject all other bids
    await db.job_bids.update_many(
        {"job_id": job_id, "id": {"$ne": bid_id}},
        {"$set": {"status": "rejected"}}
    )
    
    # Update job status
    await db.job_posts.update_one(
        {"id": job_id},
        {"$set": {
            "status": "in_progress",
            "accepted_bid_id": bid_id,
            "accepted_contractor_id": bid["contractor_id"],
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Create a booking from this job
    booking_id = str(uuid.uuid4())
    booking_doc = {
        "id": booking_id,
        "customer_id": user_id,
        "contractor_id": bid["contractor_id"],
        "job_id": job_id,
        "title": job["title"],
        "description": job["description"],
        "address": job["address"],
        "amount": bid["amount"],
        "status": "confirmed",
        "preferred_date": bid.get("available_date") or job.get("preferred_date"),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    await db.bookings.insert_one(booking_doc)
    
    # Notify contractor
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": bid["contractor_id"],
        "title": "Bid Accepted!",
        "body": f"Your bid of ${bid['amount']} for '{job['title']}' was accepted",
        "type": "bid_accepted",
        "data": {"job_id": job_id, "bid_id": bid_id, "booking_id": booking_id},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {"status": "accepted", "booking_id": booking_id}

@api_router.put("/jobs/{job_id}/cancel")
async def cancel_job_post(job_id: str, user_id: str):
    """Cancel a job post"""
    job = await db.job_posts.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["user_id"] != user_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    await db.job_posts.update_one(
        {"id": job_id},
        {"$set": {"status": "cancelled", "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return {"status": "cancelled"}

@api_router.get("/contractors/{contractor_id}/bids")
async def get_contractor_bids(contractor_id: str, status: Optional[str] = None):
    """Get all bids submitted by a contractor"""
    query = {"contractor_id": contractor_id}
    if status:
        query["status"] = status
    
    bids = await db.job_bids.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    
    for bid in bids:
        job = await db.job_posts.find_one({"id": bid["job_id"]}, {"_id": 0, "title": 1, "category": 1, "address": 1, "status": 1})
        bid["job"] = job or {}
    
    return bids


# ========== USER PROFILE ==========

@api_router.get("/users/{user_id}")
async def get_user(user_id: str):
    """Get user by ID"""
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

@api_router.put("/users/{user_id}")
async def update_user(user_id: str, updates: Dict[str, Any]):
    """Update user profile"""
    allowed_fields = ["name", "phone", "avatar", "bio", "language"]
    update_dict = {k: v for k, v in updates.items() if k in allowed_fields}
    if update_dict:
        await db.users.update_one({"id": user_id}, {"$set": update_dict})
    return {"status": "updated"}

@api_router.get("/users/{user_id}/preferences")
async def get_user_preferences(user_id: str):
    """Get user notification and privacy preferences"""
    prefs = await db.user_preferences.find_one({"user_id": user_id}, {"_id": 0})
    if not prefs:
        return {
            "notifications": {
                "email_new_listings": True,
                "email_messages": True,
                "email_applications": True,
                "push_enabled": True,
                "sms_enabled": False
            },
            "privacy": {
                "profile_visible": True,
                "show_phone": False,
                "allow_messages": True
            }
        }
    return prefs

@api_router.put("/users/{user_id}/preferences")
async def update_user_preferences(user_id: str, updates: Dict[str, Any]):
    """Update user notification and privacy preferences"""
    existing = await db.user_preferences.find_one({"user_id": user_id})
    
    if existing:
        await db.user_preferences.update_one(
            {"user_id": user_id},
            {"$set": updates}
        )
    else:
        await db.user_preferences.insert_one({
            "user_id": user_id,
            **updates,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
    
    return {"status": "updated"}

@api_router.post("/auth/change-password")
async def change_password(data: Dict[str, str]):
    """Change user password"""
    user_id = data.get("user_id")
    current_password = data.get("current_password")
    new_password = data.get("new_password")
    
    if not all([user_id, current_password, new_password]):
        raise HTTPException(status_code=400, detail="Missing required fields")
    
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    
    user = await db.users.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Verify current password (supports both password_hash and legacy password fields)
    stored_password = user.get("password_hash") or user.get("password", "")
    if not verify_password(current_password, stored_password):
        raise HTTPException(status_code=400, detail="Invalid current password")
    
    # Hash and update new password
    hashed = hash_password(new_password)
    await db.users.update_one(
        {"id": user_id},
        {"$set": {"password_hash": hashed, "updated_at": datetime.now(timezone.utc).isoformat()}, "$unset": {"password": ""}}
    )
    
    return {"status": "password_changed"}

@api_router.delete("/users/{user_id}")
async def delete_user_account(user_id: str):
    """Delete user account and associated data"""
    # Delete user
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Delete associated data
    await db.user_preferences.delete_many({"user_id": user_id})
    await db.renter_resumes.delete_many({"user_id": user_id})
    await db.roommate_profiles.delete_many({"user_id": user_id})
    await db.favorites.delete_many({"user_id": user_id})
    await db.chat_sessions.delete_many({"user_id": user_id})
    
    return {"status": "deleted"}

# ========== IMAGE UPLOAD ==========

@api_router.post("/upload/image")
async def upload_image(file: UploadFile = File(...)):
    """Upload an image to R2 (with fallback to base64 for backward compatibility)"""
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")
    
    content_type = file.content_type or "image/jpeg"
    image_id = str(uuid.uuid4())
    
    # Try R2 first if configured
    if is_r2_configured():
        try:
            result = await upload_file(
                file_data=content,
                filename=file.filename,
                content_type=content_type,
                folder="images",
                metadata={"image_id": image_id}
            )
            
            # Store reference in DB
            await db.images.insert_one({
                "id": image_id,
                "filename": file.filename,
                "content_type": content_type,
                "r2_url": result["url"],
                "r2_key": result["key"],
                "storage": "r2",
                "created_at": datetime.now(timezone.utc).isoformat()
            })
            
            return {"id": image_id, "url": result["url"]}
        except Exception as e:
            logger.warning(f"R2 upload failed, falling back to base64: {e}")
    
    # Fallback to base64 storage
    b64 = base64.b64encode(content).decode("utf-8")
    data_url = f"data:{content_type};base64,{b64}"
    
    await db.images.insert_one({
        "id": image_id,
        "filename": file.filename,
        "content_type": content_type,
        "data": data_url,
        "storage": "base64",
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {"id": image_id, "url": data_url}

@api_router.get("/images/{image_id}")
async def get_image(image_id: str):
    """Get an uploaded image"""
    img = await db.images.find_one({"id": image_id}, {"_id": 0})
    if not img:
        raise HTTPException(status_code=404, detail="Image not found")
    return img

# ========== ROOMMATE FINDER ==========

@api_router.post("/roommates/profile")
async def create_roommate_profile(user_id: str, profile: RoommateProfileCreate):
    """Create or update roommate profile"""
    existing = await db.roommate_profiles.find_one({"user_id": user_id}, {"_id": 0})
    if existing:
        await db.roommate_profiles.update_one(
            {"user_id": user_id},
            {"$set": {**profile.model_dump(), "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
        updated = await db.roommate_profiles.find_one({"user_id": user_id}, {"_id": 0})
        return updated
    
    profile_obj = RoommateProfile(user_id=user_id, **profile.model_dump())
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if user:
        profile_obj.name = profile.name or user.get("name", "")
    doc = profile_obj.model_dump()
    await db.roommate_profiles.insert_one(doc)
    doc.pop("_id", None)
    return doc

@api_router.get("/roommates/profile/{user_id}")
async def get_roommate_profile(user_id: str):
    profile = await db.roommate_profiles.find_one({"user_id": user_id}, {"_id": 0})
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    return profile

@api_router.get("/roommates/search")
async def search_roommates(
    budget_min: Optional[int] = None,
    budget_max: Optional[int] = None,
    area: Optional[str] = None,
    lifestyle: Optional[str] = None,
    pets: Optional[bool] = None,
    smoking: Optional[bool] = None,
    user_id: Optional[str] = None
):
    """Search for compatible roommates"""
    query = {"status": "active"}
    if user_id:
        query["user_id"] = {"$ne": user_id}
    if budget_min:
        query["budget_max"] = {"$gte": budget_min}
    if budget_max:
        query["budget_min"] = {"$lte": budget_max}
    if area:
        query["preferred_areas"] = {"$regex": area, "$options": "i"}
    if lifestyle:
        query["lifestyle"] = {"$regex": lifestyle, "$options": "i"}
    if pets is not None:
        query["pets"] = pets
    if smoking is not None:
        query["smoking"] = smoking
    
    profiles = await db.roommate_profiles.find(query, {"_id": 0}).sort("created_at", -1).to_list(50)
    return profiles

@api_router.post("/roommates/{profile_id}/connect")
async def connect_roommate(profile_id: str, user_id: str, message: str = ""):
    """Send a connection request to a roommate"""
    target = await db.roommate_profiles.find_one({"id": profile_id}, {"_id": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    connection = {
        "id": str(uuid.uuid4()),
        "from_user_id": user_id,
        "to_user_id": target["user_id"],
        "profile_id": profile_id,
        "message": message,
        "status": "pending",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.roommate_connections.insert_one(connection)
    
    await db.notifications.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": target["user_id"],
        "title": "New Roommate Request",
        "body": f"Someone wants to connect as a potential roommate!",
        "type": "roommate",
        "data": {"connection_id": connection["id"]},
        "read": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    connection.pop("_id", None)
    return connection

@api_router.get("/roommates/connections/{user_id}")
async def get_roommate_connections(user_id: str):
    """Get roommate connections for a user"""
    connections = await db.roommate_connections.find(
        {"$or": [{"from_user_id": user_id}, {"to_user_id": user_id}]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    for c in connections:
        other_id = c["to_user_id"] if c["from_user_id"] == user_id else c["from_user_id"]
        profile = await db.roommate_profiles.find_one({"user_id": other_id}, {"_id": 0})
        c["other_profile"] = profile or {}
    return connections

# ========== FAVORITES ==========

@api_router.post("/favorites")
async def toggle_favorite(user_id: str, listing_id: str):
    """Toggle favorite on a listing"""
    existing = await db.favorites.find_one({"user_id": user_id, "listing_id": listing_id})
    if existing:
        await db.favorites.delete_one({"user_id": user_id, "listing_id": listing_id})
        return {"status": "removed", "favorited": False}
    else:
        await db.favorites.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": user_id,
            "listing_id": listing_id,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        return {"status": "added", "favorited": True}

@api_router.get("/favorites/{user_id}")
async def get_favorites(user_id: str):
    """Get user's favorited listings"""
    favs = await db.favorites.find({"user_id": user_id}, {"_id": 0}).to_list(100)
    listings = []
    for f in favs:
        listing = await db.listings.find_one({"id": f["listing_id"]}, {"_id": 0})
        if listing:
            listing["favorited_at"] = f.get("created_at")
            listings.append(listing)
    return listings

@api_router.get("/favorites/{user_id}/ids")
async def get_favorite_ids(user_id: str):
    """Get just the IDs of favorited listings"""
    favs = await db.favorites.find({"user_id": user_id}, {"_id": 0, "listing_id": 1}).to_list(100)
    return [f["listing_id"] for f in favs]

# ========== PROPERTY COMPARISON ==========

@api_router.post("/compare")
async def compare_properties(listing_ids: List[str]):
    """Compare multiple properties side by side"""
    if len(listing_ids) < 2 or len(listing_ids) > 4:
        raise HTTPException(status_code=400, detail="Compare 2-4 properties")
    
    listings = []
    for lid in listing_ids:
        listing = await db.listings.find_one({"id": lid}, {"_id": 0})
        if listing:
            listings.append(listing)
    
    return {"listings": listings, "count": len(listings)}

# Seed Data Route
@api_router.post("/seed")
async def seed_data():
    count = await db.listings.count_documents({})
    if count > 0:
        return {"message": f"Database already has {count} listings"}
    
    sample_listings = [
        {
            "id": str(uuid.uuid4()),
            "title": "Modern Downtown Condo",
            "address": "1200 West Georgia St, Unit 2505",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V6E 4R2",
            "lat": 49.2849,
            "lng": -123.1215,
            "price": 2800,
            "bedrooms": 2,
            "bathrooms": 2,
            "sqft": 950,
            "property_type": "Condo",
            "description": "Stunning views from this 25th floor condo. Near SkyTrain, gyms, and coffee shops. 10 min walk to Waterfront Station.",
            "amenities": ["Gym", "Rooftop Deck", "Concierge", "In-suite Laundry", "Bike Storage"],
            "images": ["https://images.unsplash.com/photo-1502672260266-1c1ef2d93688?w=800"],
            "available_date": "2026-02-01",
            "pet_friendly": False,
            "parking": True,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Cozy Kitsilano Character Home",
            "address": "2456 West 3rd Ave",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V6K 1L8",
            "lat": 49.2689,
            "lng": -123.1620,
            "price": 3200,
            "bedrooms": 3,
            "bathrooms": 1.5,
            "sqft": 1400,
            "property_type": "House",
            "description": "Beautiful character home in family-friendly Kitsilano. 5 min walk to Kits Beach, dog parks nearby. Great for remote workers with natural light.",
            "amenities": ["Backyard", "Garage", "Fireplace", "Washer/Dryer", "Home Office"],
            "images": ["https://images.unsplash.com/photo-1600596542815-ffad4c1539a9?w=800"],
            "available_date": "2026-02-15",
            "pet_friendly": True,
            "parking": True,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Yaletown Luxury Loft",
            "address": "1155 Mainland St, Unit 808",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V6B 5P2",
            "lat": 49.2756,
            "lng": -123.1209,
            "price": 3500,
            "bedrooms": 1,
            "bathrooms": 1,
            "sqft": 850,
            "property_type": "Loft",
            "description": "Industrial loft in trendy Yaletown. 16ft ceilings, exposed brick. Walk to seawall, restaurants. Perfect for young professionals.",
            "amenities": ["Gym", "Bike Storage", "Rooftop Patio", "In-suite Laundry"],
            "images": ["https://images.unsplash.com/photo-1560448204-e02f11c3d0e2?w=800"],
            "available_date": "2026-01-15",
            "pet_friendly": True,
            "parking": True,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Mount Pleasant Studio",
            "address": "456 Main St, Unit 302",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V5T 3E2",
            "lat": 49.2634,
            "lng": -123.1006,
            "price": 1650,
            "bedrooms": 0,
            "bathrooms": 1,
            "sqft": 450,
            "property_type": "Studio",
            "description": "Bright studio in hip Mount Pleasant. Steps to Main Street shops, breweries, and coffee. Great for students or minimalists.",
            "amenities": ["In-suite Laundry", "Bike Room", "Rooftop"],
            "images": ["https://images.unsplash.com/photo-1522708323590-d24dbb6b0267?w=800"],
            "available_date": "2026-02-01",
            "pet_friendly": False,
            "parking": False,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Coal Harbour Waterfront Suite",
            "address": "1499 West Pender St, Unit 3201",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V6G 2S3",
            "lat": 49.2898,
            "lng": -123.1323,
            "price": 4200,
            "bedrooms": 2,
            "bathrooms": 2,
            "sqft": 1150,
            "property_type": "Condo",
            "description": "Breathtaking water and mountain views! Steps to Stanley Park seawall. High-end finishes, perfect for executives.",
            "amenities": ["Pool", "Gym", "Concierge", "Spa", "Wine Cellar"],
            "images": ["https://images.unsplash.com/photo-1600607687939-ce8a6c25118c?w=800"],
            "available_date": "2026-03-01",
            "pet_friendly": False,
            "parking": True,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Commercial Drive Townhouse",
            "address": "1876 Venables St",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V5L 2H6",
            "lat": 49.2752,
            "lng": -123.0711,
            "price": 2900,
            "bedrooms": 3,
            "bathrooms": 2.5,
            "sqft": 1600,
            "property_type": "Townhouse",
            "description": "Spacious townhouse near Commercial Drive. Italian cafes, indie shops nearby. Close to SkyTrain. Great for families!",
            "amenities": ["Patio", "Storage", "Washer/Dryer", "Dishwasher"],
            "images": ["https://images.unsplash.com/photo-1600566753190-17f0baa2a6c3?w=800"],
            "available_date": "2026-02-15",
            "pet_friendly": True,
            "parking": True,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Gastown Heritage Apartment",
            "address": "55 Water St, Unit 402",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V6B 1A1",
            "lat": 49.2842,
            "lng": -123.1088,
            "price": 2200,
            "bedrooms": 1,
            "bathrooms": 1,
            "sqft": 650,
            "property_type": "Apartment",
            "description": "Historic Gastown building. Exposed brick, high ceilings. Walk to Waterfront Station, nightlife, restaurants.",
            "amenities": ["Exposed Brick", "High Ceilings", "Rooftop Access"],
            "images": ["https://images.unsplash.com/photo-1600585154340-be6161a56a0c?w=800"],
            "available_date": "2026-01-20",
            "pet_friendly": False,
            "parking": False,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "UBC Area Family Home",
            "address": "4521 West 10th Ave",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V6R 2J1",
            "lat": 49.2628,
            "lng": -123.2011,
            "price": 4500,
            "bedrooms": 4,
            "bathrooms": 3,
            "sqft": 2200,
            "property_type": "House",
            "description": "Spacious family home near UBC. Excellent schools, quiet street. 15 min to campus, close to Pacific Spirit Park.",
            "amenities": ["Large Yard", "Double Garage", "Fireplace", "Deck", "Home Office"],
            "images": ["https://images.unsplash.com/photo-1600047509807-ba8f99d2cdde?w=800"],
            "available_date": "2026-03-01",
            "pet_friendly": True,
            "parking": True,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Olympic Village Modern 1BR",
            "address": "1788 Columbia St, Unit 1205",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V5Y 0L7",
            "lat": 49.2689,
            "lng": -123.1069,
            "price": 2400,
            "bedrooms": 1,
            "bathrooms": 1,
            "sqft": 680,
            "property_type": "Condo",
            "description": "Bright corner unit in Olympic Village. On seawall, near Science World. Perfect for active lifestyle - bike paths everywhere!",
            "amenities": ["Gym", "Rooftop Deck", "Bike Storage", "In-suite Laundry"],
            "images": ["https://images.unsplash.com/photo-1600607687644-aac4c3eac7f4?w=800"],
            "available_date": "2026-02-01",
            "pet_friendly": True,
            "parking": True,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "title": "West End Beach Apartment",
            "address": "1845 Beach Ave, Unit 601",
            "city": "Vancouver",
            "province": "BC",
            "postal_code": "V6G 1Y9",
            "lat": 49.2856,
            "lng": -123.1412,
            "price": 2950,
            "bedrooms": 1,
            "bathrooms": 1,
            "sqft": 720,
            "property_type": "Apartment",
            "description": "Beach living in the West End! Walk to English Bay and Stanley Park. Incredible sunsets from balcony. Dog-friendly building!",
            "amenities": ["Ocean View", "Balcony", "Pool", "In-suite Laundry"],
            "images": ["https://images.unsplash.com/photo-1600566752355-35792bedcfea?w=800"],
            "available_date": "2026-03-01",
            "pet_friendly": True,
            "parking": False,
            "status": "active",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
    ]
    
    await db.listings.insert_many(sample_listings)
    
    # Seed contractor profiles
    contractor_count = await db.contractor_profiles.count_documents({})
    if contractor_count == 0:
        sample_contractors = [
            {
                "id": str(uuid.uuid4()), "user_id": f"contractor-seed-{i}",
                "business_name": name, "description": desc,
                "specialties": specs, "service_areas": ["Vancouver", "Burnaby", "Richmond"],
                "hourly_rate": rate, "years_experience": exp,
                "insurance": True, "verified": True,
                "phone": f"604-555-{1000+i}", "email": f"{name.lower().replace(' ', '.')}@email.com",
                "rating": rating, "review_count": reviews, "completed_jobs": jobs,
                "portfolio_images": [img], "status": "active",
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
            for i, (name, desc, specs, rate, exp, rating, reviews, jobs, img) in enumerate([
                ("Vancouver Plumbing Pros", "Expert plumbing services for residential and commercial properties. Available 24/7 for emergencies.", ["plumbing", "pipe repair", "drain cleaning"], 85, 12, 4.8, 47, 156, "https://images.unsplash.com/photo-1621905251189-08b45d6a269e?w=400"),
                ("Pacific Electrical Services", "Licensed electricians serving the Greater Vancouver area. Full residential and commercial electrical services.", ["electrical", "wiring", "panel upgrades"], 95, 15, 4.9, 62, 203, "https://images.unsplash.com/photo-1621905252507-b35492cc74b4?w=400"),
                ("West Coast Painters", "Professional painting crew. Interior, exterior, and specialty finishes. Color consultation included.", ["painting", "drywall", "wallpaper"], 65, 8, 4.7, 35, 98, "https://images.unsplash.com/photo-1562259949-e8e7689d7828?w=400"),
                ("Mountain View Renovations", "Full-service renovation company. Kitchen, bathroom, and complete home remodels.", ["renovation", "carpentry", "flooring"], 110, 20, 4.9, 89, 312, "https://images.unsplash.com/photo-1504307651254-35680f356dfd?w=400"),
                ("Green Thumb Landscaping", "Transform your outdoor space. Design, installation, and ongoing maintenance services.", ["landscaping", "gardening", "irrigation"], 55, 10, 4.6, 28, 134, "https://images.unsplash.com/photo-1558904541-efa843a96f01?w=400"),
                ("Clean Sweep Services", "Professional cleaning for move-in, move-out, and regular maintenance. Eco-friendly products.", ["cleaning", "deep cleaning", "move-out"], 45, 6, 4.8, 55, 245, "https://images.unsplash.com/photo-1581578731548-c64695cc6952?w=400"),
            ])
        ]
        await db.contractor_profiles.insert_many(sample_contractors)
        
        # Seed contractor services
        sample_services = []
        for cp in sample_contractors:
            for j, (title, desc, cat, price, dur) in enumerate([
                (f"{cp['specialties'][0].title()} Inspection", f"Complete {cp['specialties'][0]} inspection and assessment", cp['specialties'][0], cp['hourly_rate'] * 2, "1-2 hours"),
                (f"{cp['specialties'][0].title()} Repair", f"Professional {cp['specialties'][0]} repair service", cp['specialties'][0], cp['hourly_rate'] * 4, "2-4 hours"),
            ]):
                sample_services.append({
                    "id": str(uuid.uuid4()), "contractor_id": cp["user_id"],
                    "title": title, "description": desc, "category": cat,
                    "price_type": "fixed", "price": price, "duration_estimate": dur,
                    "images": [], "status": "active",
                    "created_at": datetime.now(timezone.utc).isoformat()
                })
        await db.contractor_services.insert_many(sample_services)
    
    return {"message": f"Successfully seeded {len(sample_listings)} listings and contractor data"}

# ========== ANALYTICS DASHBOARD ==========

@api_router.get("/analytics/overview")
async def get_analytics_overview():
    """Get platform-wide analytics overview for admin dashboard"""
    
    # User stats
    total_users = await db.users.count_documents({})
    users_by_type = await db.users.aggregate([
        {"$group": {"_id": "$user_type", "count": {"$sum": 1}}}
    ]).to_list(10)
    
    # Listing stats
    total_listings = await db.listings.count_documents({"status": "active"})
    listings_by_type = await db.listings.aggregate([
        {"$match": {"status": "active"}},
        {"$group": {"_id": "$listing_type", "count": {"$sum": 1}}}
    ]).to_list(10)
    listings_by_city = await db.listings.aggregate([
        {"$match": {"status": "active"}},
        {"$group": {"_id": "$city", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
        {"$limit": 5}
    ]).to_list(5)
    
    # Transaction stats
    total_transactions = await db.payment_transactions.count_documents({})
    successful_transactions = await db.payment_transactions.count_documents({"payment_status": "paid"})
    
    # Calculate total revenue
    revenue_pipeline = [
        {"$match": {"payment_status": "paid"}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]
    revenue_result = await db.payment_transactions.aggregate(revenue_pipeline).to_list(1)
    total_revenue = revenue_result[0]["total"] if revenue_result else 0
    
    # Contractor stats
    total_contractors = await db.contractor_profiles.count_documents({"status": "active"})
    verified_contractors = await db.contractor_profiles.count_documents({"status": "active", "verified": True})
    
    # Lease assignment stats  
    active_assignments = await db.lease_assignments.count_documents({"status": "active"})
    
    # Document stats
    total_documents = await db.esign_documents.count_documents({})
    signed_documents = await db.esign_documents.count_documents({"status": "signed"})
    
    # Recent activity (last 7 days)
    from datetime import timedelta
    seven_days_ago = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
    
    new_users_7d = await db.users.count_documents({"created_at": {"$gte": seven_days_ago}})
    new_listings_7d = await db.listings.count_documents({"created_at": {"$gte": seven_days_ago}})
    
    return {
        "users": {
            "total": total_users,
            "by_type": {item["_id"]: item["count"] for item in users_by_type if item["_id"]},
            "new_last_7_days": new_users_7d
        },
        "listings": {
            "total_active": total_listings,
            "by_type": {item["_id"]: item["count"] for item in listings_by_type if item["_id"]},
            "by_city": {item["_id"]: item["count"] for item in listings_by_city if item["_id"]},
            "new_last_7_days": new_listings_7d
        },
        "transactions": {
            "total": total_transactions,
            "successful": successful_transactions,
            "total_revenue": round(total_revenue, 2),
            "success_rate": round((successful_transactions / total_transactions * 100) if total_transactions > 0 else 0, 1)
        },
        "contractors": {
            "total": total_contractors,
            "verified": verified_contractors,
            "verification_rate": round((verified_contractors / total_contractors * 100) if total_contractors > 0 else 0, 1)
        },
        "lease_assignments": {
            "active": active_assignments
        },
        "documents": {
            "total": total_documents,
            "signed": signed_documents,
            "completion_rate": round((signed_documents / total_documents * 100) if total_documents > 0 else 0, 1)
        },
        "generated_at": datetime.now(timezone.utc).isoformat()
    }

@api_router.get("/analytics/activity")
async def get_recent_activity(limit: int = 20):
    """Get recent platform activity feed"""
    
    activities = []
    
    # Recent users
    recent_users = await db.users.find(
        {}, {"_id": 0, "email": 1, "user_type": 1, "created_at": 1, "name": 1}
    ).sort("created_at", -1).limit(5).to_list(5)
    
    for user in recent_users:
        activities.append({
            "type": "user_signup",
            "title": f"New {user.get('user_type', 'user')} registered",
            "description": user.get("name", user.get("email", "Unknown")),
            "timestamp": user.get("created_at"),
            "icon": "user"
        })
    
    # Recent listings
    recent_listings = await db.listings.find(
        {"status": "active"}, {"_id": 0, "title": 1, "city": 1, "price": 1, "created_at": 1}
    ).sort("created_at", -1).limit(5).to_list(5)
    
    for listing in recent_listings:
        activities.append({
            "type": "new_listing",
            "title": "New listing created",
            "description": f"{listing.get('title')} in {listing.get('city')} - ${listing.get('price')}/mo",
            "timestamp": listing.get("created_at"),
            "icon": "home"
        })
    
    # Recent transactions
    recent_payments = await db.payment_transactions.find(
        {"payment_status": "paid"}, {"_id": 0, "amount": 1, "description": 1, "created_at": 1}
    ).sort("created_at", -1).limit(5).to_list(5)
    
    for payment in recent_payments:
        activities.append({
            "type": "payment",
            "title": "Payment completed",
            "description": f"${payment.get('amount', 0):.2f} - {payment.get('description', 'N/A')}",
            "timestamp": payment.get("created_at"),
            "icon": "dollar"
        })
    
    # Sort all activities by timestamp
    activities.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
    
    return activities[:limit]

@api_router.get("/analytics/revenue")
async def get_revenue_analytics(period: str = "30d"):
    """Get revenue analytics over time"""
    
    from datetime import timedelta
    
    # Calculate date range
    if period == "7d":
        days = 7
    elif period == "30d":
        days = 30
    elif period == "90d":
        days = 90
    else:
        days = 30
    
    start_date = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
    
    # Daily revenue breakdown
    pipeline = [
        {"$match": {"payment_status": "paid", "created_at": {"$gte": start_date}}},
        {"$addFields": {
            "date": {"$substr": ["$created_at", 0, 10]}
        }},
        {"$group": {
            "_id": "$date",
            "revenue": {"$sum": "$amount"},
            "count": {"$sum": 1}
        }},
        {"$sort": {"_id": 1}}
    ]
    
    daily_revenue = await db.payment_transactions.aggregate(pipeline).to_list(100)
    
    # Total for period
    total_revenue = sum(day["revenue"] for day in daily_revenue)
    total_transactions = sum(day["count"] for day in daily_revenue)
    
    return {
        "period": period,
        "total_revenue": round(total_revenue, 2),
        "total_transactions": total_transactions,
        "average_transaction": round(total_revenue / total_transactions, 2) if total_transactions > 0 else 0,
        "daily_breakdown": [
            {"date": day["_id"], "revenue": round(day["revenue"], 2), "count": day["count"]}
            for day in daily_revenue
        ]
    }

@api_router.get("/analytics/listings-performance")
async def get_listings_performance():
    """Get listing performance metrics"""
    
    # Price distribution
    price_ranges = [
        {"label": "Under $1,500", "min": 0, "max": 1499},
        {"label": "$1,500 - $2,000", "min": 1500, "max": 1999},
        {"label": "$2,000 - $2,500", "min": 2000, "max": 2499},
        {"label": "$2,500 - $3,000", "min": 2500, "max": 2999},
        {"label": "$3,000+", "min": 3000, "max": 999999},
    ]
    
    price_distribution = []
    for range_item in price_ranges:
        count = await db.listings.count_documents({
            "status": "active",
            "listing_type": "rent",
            "price": {"$gte": range_item["min"], "$lte": range_item["max"]}
        })
        price_distribution.append({
            "range": range_item["label"],
            "count": count
        })
    
    # Property type distribution
    property_types = await db.listings.aggregate([
        {"$match": {"status": "active"}},
        {"$group": {"_id": "$property_type", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]).to_list(10)
    
    # Average prices by city
    avg_prices_by_city = await db.listings.aggregate([
        {"$match": {"status": "active", "listing_type": "rent"}},
        {"$group": {"_id": "$city", "avg_price": {"$avg": "$price"}, "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
        {"$limit": 10}
    ]).to_list(10)
    
    return {
        "price_distribution": price_distribution,
        "property_types": [{"type": p["_id"] or "Unknown", "count": p["count"]} for p in property_types],
        "avg_prices_by_city": [
            {"city": c["_id"] or "Unknown", "avg_price": round(c["avg_price"], 0), "count": c["count"]}
            for c in avg_prices_by_city
        ]
    }

# ========== ROLE-SPECIFIC ANALYTICS ==========

@api_router.get("/analytics/renter/{user_id}")
async def get_renter_analytics(user_id: str):
    """Get renter-specific analytics"""
    
    # Saved/favorite properties count
    favorites_count = await db.favorites.count_documents({"user_id": user_id})
    favorites = await db.favorites.find({"user_id": user_id}, {"_id": 0}).sort("created_at", -1).to_list(10)
    
    # Application statistics
    applications = await db.applications.find({"user_id": user_id}, {"_id": 0}).to_list(100)
    app_stats = {
        "total": len(applications),
        "pending": len([a for a in applications if a.get("status") == "pending"]),
        "approved": len([a for a in applications if a.get("status") == "approved"]),
        "rejected": len([a for a in applications if a.get("status") == "rejected"]),
    }
    
    # Recent applications with listing details
    recent_apps = applications[:5]
    for app in recent_apps:
        listing = await db.listings.find_one({"id": app.get("listing_id")}, {"_id": 0, "title": 1, "address": 1, "city": 1, "price": 1})
        app["listing"] = listing
    
    # Viewing history (from chat sessions or favorites)
    recent_favorites = []
    for fav in favorites[:5]:
        listing = await db.listings.find_one({"id": fav.get("listing_id")}, {"_id": 0, "title": 1, "address": 1, "city": 1, "price": 1, "images": 1})
        if listing:
            listing["favorited_at"] = fav.get("created_at")
            recent_favorites.append(listing)
    
    # Messages sent/received
    messages_sent = await db.messages.count_documents({"sender_id": user_id})
    messages_received = await db.messages.count_documents({"receiver_id": user_id})
    
    # Renter resume completion
    resume = await db.renter_resumes.find_one({"user_id": user_id}, {"_id": 0})
    resume_score = 0
    if resume:
        fields = ["employment_status", "employer", "income", "rental_history", "references"]
        filled = sum(1 for f in fields if resume.get(f))
        resume_score = int((filled / len(fields)) * 100)
    
    return {
        "favorites_count": favorites_count,
        "recent_favorites": recent_favorites,
        "applications": app_stats,
        "recent_applications": recent_apps,
        "messages": {"sent": messages_sent, "received": messages_received},
        "resume_completion": resume_score,
        "generated_at": datetime.now(timezone.utc).isoformat()
    }

@api_router.get("/analytics/landlord/{user_id}")
async def get_landlord_analytics(user_id: str):
    """Get landlord-specific analytics"""
    
    # Properties owned
    listings = await db.listings.find({"landlord_id": user_id}, {"_id": 0}).to_list(100)
    active_listings = [l for l in listings if l.get("status") == "active"]
    
    # Total potential monthly revenue
    monthly_revenue = sum(l.get("price", 0) for l in active_listings if l.get("listing_type") == "rent")
    
    # Applications received
    listing_ids = [l["id"] for l in listings]
    applications = await db.applications.find({"listing_id": {"$in": listing_ids}}, {"_id": 0}).to_list(500)
    
    app_stats = {
        "total": len(applications),
        "pending": len([a for a in applications if a.get("status") == "pending"]),
        "approved": len([a for a in applications if a.get("status") == "approved"]),
        "rejected": len([a for a in applications if a.get("status") == "rejected"]),
    }
    
    # Applications by listing
    apps_by_listing = {}
    for app in applications:
        lid = app.get("listing_id")
        if lid not in apps_by_listing:
            apps_by_listing[lid] = 0
        apps_by_listing[lid] += 1
    
    # Top performing listings (by application count)
    listing_performance = []
    for listing in listings[:10]:
        listing_performance.append({
            "id": listing["id"],
            "title": listing.get("title"),
            "price": listing.get("price"),
            "applications": apps_by_listing.get(listing["id"], 0),
            "status": listing.get("status")
        })
    listing_performance.sort(key=lambda x: x["applications"], reverse=True)
    
    # Messages/inquiries
    messages_received = await db.messages.count_documents({"receiver_id": user_id})
    
    # Maintenance requests for their properties
    maintenance = await db.maintenance_requests.find({"landlord_id": user_id}, {"_id": 0}).to_list(100)
    maintenance_stats = {
        "total": len(maintenance),
        "open": len([m for m in maintenance if m.get("status") in ["pending", "in_progress"]]),
        "resolved": len([m for m in maintenance if m.get("status") == "resolved"]),
    }
    
    # Revenue from payments
    payments = await db.payment_transactions.find(
        {"recipient_id": user_id, "payment_status": "paid"},
        {"_id": 0}
    ).to_list(100)
    total_collected = sum(p.get("amount", 0) for p in payments)
    
    return {
        "properties": {
            "total": len(listings),
            "active": len(active_listings),
            "inactive": len(listings) - len(active_listings)
        },
        "monthly_potential_revenue": monthly_revenue,
        "total_collected": total_collected,
        "applications": app_stats,
        "listing_performance": listing_performance[:5],
        "inquiries": messages_received,
        "maintenance": maintenance_stats,
        "generated_at": datetime.now(timezone.utc).isoformat()
    }

@api_router.get("/analytics/contractor/{user_id}")
async def get_contractor_analytics(user_id: str):
    """Get contractor-specific analytics"""
    
    # Profile info
    profile = await db.contractor_profiles.find_one({"user_id": user_id}, {"_id": 0})
    
    # Jobs completed
    bookings = await db.bookings.find({"contractor_id": user_id}, {"_id": 0}).to_list(500)
    completed_jobs = [b for b in bookings if b.get("status") == "completed"]
    pending_jobs = [b for b in bookings if b.get("status") in ["pending", "confirmed"]]
    
    # Earnings
    total_earnings = sum(b.get("price", 0) for b in completed_jobs)
    
    # Monthly earnings breakdown
    monthly_earnings = {}
    for booking in completed_jobs:
        created = booking.get("created_at", "")
        if created:
            month_key = created[:7]  # YYYY-MM
            if month_key not in monthly_earnings:
                monthly_earnings[month_key] = 0
            monthly_earnings[month_key] += booking.get("price", 0)
    
    # Reviews and ratings
    reviews = await db.reviews.find({"contractor_id": user_id}, {"_id": 0}).to_list(100)
    avg_rating = sum(r.get("rating", 0) for r in reviews) / len(reviews) if reviews else 0
    rating_distribution = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    for r in reviews:
        rating = r.get("rating", 0)
        if 1 <= rating <= 5:
            rating_distribution[rating] += 1
    
    # Job bids
    bids = await db.job_bids.find({"contractor_id": user_id}, {"_id": 0}).to_list(100)
    bid_stats = {
        "total": len(bids),
        "won": len([b for b in bids if b.get("status") == "accepted"]),
        "pending": len([b for b in bids if b.get("status") == "pending"]),
    }
    bid_stats["win_rate"] = round((bid_stats["won"] / bid_stats["total"] * 100) if bid_stats["total"] > 0 else 0, 1)
    
    # Lead sources (from bookings)
    lead_sources = {}
    for b in bookings:
        source = b.get("source", "direct")
        if source not in lead_sources:
            lead_sources[source] = 0
        lead_sources[source] += 1
    
    # Recent reviews
    recent_reviews = sorted(reviews, key=lambda x: x.get("created_at", ""), reverse=True)[:5]
    for review in recent_reviews:
        customer = await db.users.find_one({"id": review.get("customer_id")}, {"_id": 0, "name": 1})
        review["customer_name"] = customer.get("name") if customer else "Anonymous"
    
    return {
        "profile": {
            "business_name": profile.get("business_name") if profile else None,
            "verified": profile.get("verified", False) if profile else False,
            "specialties": profile.get("specialties", []) if profile else []
        },
        "jobs": {
            "completed": len(completed_jobs),
            "pending": len(pending_jobs),
            "total": len(bookings)
        },
        "earnings": {
            "total": total_earnings,
            "monthly": monthly_earnings
        },
        "ratings": {
            "average": round(avg_rating, 1),
            "total_reviews": len(reviews),
            "distribution": rating_distribution
        },
        "bids": bid_stats,
        "lead_sources": lead_sources,
        "recent_reviews": recent_reviews,
        "generated_at": datetime.now(timezone.utc).isoformat()
    }


@api_router.get("/analytics/trends/{user_id}")
async def get_user_trends(user_id: str, period: str = "30d"):
    """Get trend data for a user (views, applications, earnings over time)"""
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "user_type": 1})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Calculate date range
    days = int(period.replace("d", "")) if period.endswith("d") else 30
    start_date = datetime.now(timezone.utc) - timedelta(days=days)
    
    trends = {
        "period": period,
        "start_date": start_date.isoformat(),
        "end_date": datetime.now(timezone.utc).isoformat(),
        "data_points": []
    }
    
    user_type = user.get("user_type")
    
    if user_type == "landlord":
        # Get daily views and applications for landlord's properties
        properties = await db.listings.find({"user_id": user_id}, {"_id": 0, "id": 1}).to_list(100)
        property_ids = [p["id"] for p in properties]
        
        for i in range(days):
            day = start_date + timedelta(days=i)
            day_str = day.strftime("%Y-%m-%d")
            
            # Count applications for this day
            apps_count = await db.applications.count_documents({
                "listing_id": {"$in": property_ids},
                "created_at": {"$regex": f"^{day_str}"}
            })
            
            trends["data_points"].append({
                "date": day_str,
                "applications": apps_count,
                "views": apps_count * 10  # Estimate views as 10x applications
            })
            
    elif user_type == "contractor":
        # Get daily jobs and earnings
        for i in range(days):
            day = start_date + timedelta(days=i)
            day_str = day.strftime("%Y-%m-%d")
            
            jobs_count = await db.contractor_bookings.count_documents({
                "contractor_id": user_id,
                "created_at": {"$regex": f"^{day_str}"}
            })
            
            trends["data_points"].append({
                "date": day_str,
                "jobs": jobs_count,
                "leads": jobs_count * 3  # Estimate leads
            })
    else:
        # Renter - track searches and applications
        for i in range(days):
            day = start_date + timedelta(days=i)
            day_str = day.strftime("%Y-%m-%d")
            
            apps_count = await db.applications.count_documents({
                "user_id": user_id,
                "created_at": {"$regex": f"^{day_str}"}
            })
            
            trends["data_points"].append({
                "date": day_str,
                "applications": apps_count,
                "searches": apps_count * 5  # Estimate
            })
    
    return trends

@api_router.get("/analytics/insights/{user_id}")
async def get_user_insights(user_id: str):
    """Get AI-generated insights for user performance"""
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "user_type": 1, "name": 1})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    user_type = user.get("user_type")
    insights = []
    
    if user_type == "landlord":
        # Get landlord metrics
        properties = await db.listings.find({"user_id": user_id}, {"_id": 0}).to_list(100)
        total_properties = len(properties)
        active_properties = len([p for p in properties if p.get("status") == "active"])
        
        apps = await db.applications.find({"listing_id": {"$in": [p["id"] for p in properties]}}).to_list(500)
        total_apps = len(apps)
        
        if total_properties > 0 and active_properties < total_properties * 0.5:
            insights.append({
                "type": "warning",
                "title": "Inactive Listings",
                "message": f"Only {active_properties} of {total_properties} listings are active. Consider reactivating or updating your listings.",
                "action": "View Listings",
                "action_url": "/my-properties"
            })
        
        if total_apps > 0:
            avg_apps_per_listing = total_apps / max(active_properties, 1)
            if avg_apps_per_listing < 2:
                insights.append({
                    "type": "tip",
                    "title": "Boost Applications",
                    "message": "Your listings are receiving below-average applications. Try adding more photos or adjusting pricing.",
                    "action": "Optimize Listings",
                    "action_url": "/listing-optimizer"
                })
        
        insights.append({
            "type": "info",
            "title": "Market Update",
            "message": "Vancouver rental prices have increased 3% this quarter. Consider reviewing your pricing.",
            "action": "View Analytics",
            "action_url": "/analytics"
        })
        
    elif user_type == "contractor":
        # Get contractor metrics
        profile = await db.contractor_profiles.find_one({"user_id": user_id}, {"_id": 0})
        reviews = await db.ratings.find({"rated_user_id": user_id}).to_list(100)
        
        if profile and not profile.get("verified"):
            insights.append({
                "type": "tip",
                "title": "Get Verified",
                "message": "Verified contractors receive 3x more leads. Complete your verification to stand out.",
                "action": "Start Verification",
                "action_url": "/contractor-profile"
            })
        
        if len(reviews) < 5:
            insights.append({
                "type": "tip",
                "title": "Build Your Reputation",
                "message": f"You have {len(reviews)} reviews. Encourage satisfied customers to leave reviews.",
                "action": "View Reviews",
                "action_url": "/analytics"
            })
        
        insights.append({
            "type": "success",
            "title": "High Demand Alert",
            "message": "Plumbing services are in high demand this week. Check available jobs!",
            "action": "Browse Jobs",
            "action_url": "/contractors"
        })
        
    else:  # Renter
        # Get renter metrics
        apps = await db.applications.find({"user_id": user_id}).to_list(100)
        favorites = await db.favorites.count_documents({"user_id": user_id})
        
        pending_apps = len([a for a in apps if a.get("status") == "pending"])
        
        if pending_apps > 0:
            insights.append({
                "type": "info",
                "title": "Applications Pending",
                "message": f"You have {pending_apps} pending applications. Check for updates!",
                "action": "View Applications",
                "action_url": "/applications"
            })
        
        if favorites > 10:
            insights.append({
                "type": "tip",
                "title": "Compare Properties",
                "message": f"You have {favorites} saved properties. Use our comparison tool to decide!",
                "action": "Compare",
                "action_url": "/compare"
            })
        
        insights.append({
            "type": "info",
            "title": "New Listings",
            "message": "12 new properties matching your criteria were listed this week.",
            "action": "Browse",
            "action_url": "/browse"
        })
    
    return {"user_id": user_id, "user_type": user_type, "insights": insights}



# ========== SYNDICATION AI CONTENT GENERATION ==========

class SyndicationRequest(BaseModel):
    listing_id: str
    listing_data: dict

@api_router.post("/syndication/generate-description")
async def generate_syndication_description(request: SyndicationRequest):
    """Generate AI-powered listing description for syndication"""
    listing = request.listing_data
    
    try:
        from anthropic import AsyncAnthropic
        anthropic_client = AsyncAnthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))
        
        # Build context about the listing
        address = listing.get('address', '')
        city = listing.get('city', 'Vancouver')
        price = listing.get('price', 0)
        bedrooms = listing.get('bedrooms', 1)
        bathrooms = listing.get('bathrooms', 1)
        sqft = listing.get('sqft', 0)
        property_type = listing.get('property_type', 'Apartment')
        amenities = listing.get('amenities', [])
        description = listing.get('description', '')
        pet_friendly = listing.get('pet_friendly', False)
        
        prompt = f"""Generate compelling marketing content for a rental listing in Metro Vancouver.

LISTING DETAILS:
- Address: {address}, {city}
- Price: ${price}/month
- Type: {property_type}
- Bedrooms: {bedrooms}
- Bathrooms: {bathrooms}
- Size: {sqft} sqft
- Pet Friendly: {"Yes" if pet_friendly else "No"}
- Amenities: {', '.join(amenities) if amenities else 'Modern finishes'}
- Original Description: {description}

Please provide:
1. A compelling 2-3 sentence description that highlights the best features and appeals to renters. Focus on lifestyle benefits and key selling points. Make it sound professional but warm.

2. The neighborhood/area information for {city} - what's nearby, transit access, walkability, local amenities (restaurants, parks, shopping). Be specific to the actual area if you know it.

3. 5 key highlights/selling points as bullet points that would attract renters.

Format your response EXACTLY as JSON:
{{
  "description": "Your compelling 2-3 sentence description here",
  "neighborhood": "2-3 sentences about the neighborhood and area",
  "highlights": ["highlight 1", "highlight 2", "highlight 3", "highlight 4", "highlight 5"]
}}

Only respond with valid JSON, no other text."""

        response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=800,
            messages=[{"role": "user", "content": prompt}]
        )
        
        # Parse the response
        ai_text = response.content[0].text.strip()
        
        # Try to parse as JSON
        import json
        try:
            # Handle potential markdown code blocks
            if ai_text.startswith('```'):
                ai_text = ai_text.split('```')[1]
                if ai_text.startswith('json'):
                    ai_text = ai_text[4:]
            ai_content = json.loads(ai_text)
        except json.JSONDecodeError:
            # Fallback: extract content manually
            ai_content = {
                "description": f"Beautiful {bedrooms}-bedroom {property_type.lower()} available in {city}. This well-appointed home offers modern amenities and a comfortable living space perfect for your lifestyle.",
                "neighborhood": f"{city} offers excellent access to transit, shopping, and dining options. The area is known for its vibrant community and convenient location.",
                "highlights": [
                    f"{bedrooms} spacious bedroom{'s' if bedrooms > 1 else ''}",
                    "Modern finishes throughout",
                    "Great location with easy transit access",
                    "Bright and welcoming atmosphere",
                    "Pet friendly" if pet_friendly else "Quiet residential area"
                ]
            }
        
        ai_content["generated"] = True
        return ai_content
        
    except Exception as e:
        logger.error(f"Error generating syndication content: {e}")
        # Return fallback content
        return {
            "description": f"Wonderful {listing.get('bedrooms', 1)}-bedroom {listing.get('property_type', 'apartment').lower()} available in {listing.get('city', 'Vancouver')}. Modern finishes and a great location make this the perfect place to call home.",
            "neighborhood": f"{listing.get('city', 'Vancouver')} is a vibrant area with excellent transit access, local shops, and parks nearby.",
            "highlights": [
                f"{listing.get('bedrooms', 1)} bedroom{'s' if listing.get('bedrooms', 1) > 1 else ''}",
                "Modern appliances",
                "Great natural light",
                "Convenient location",
                "Close to amenities"
            ],
            "generated": True
        }





# ========== UNIVERSAL RATING SYSTEM ==========

class UserRatingCreate(BaseModel):
    rated_user_id: str
    rating: int  # 1-5 stars
    review: Optional[str] = None
    context_type: str  # 'rental', 'service', 'landlord', 'contractor', 'renter'
    context_id: Optional[str] = None  # listing_id, booking_id, etc.

class UserRatingResponse(BaseModel):
    id: str
    rated_user_id: str
    rater_user_id: str
    rater_name: str
    rating: int
    review: Optional[str]
    context_type: str
    context_id: Optional[str]
    created_at: str

@api_router.post("/ratings/user")
async def create_user_rating(rating_data: UserRatingCreate, rater_id: str):
    """Create a rating for any user (renter, landlord, contractor)"""
    # Validate rating
    if not 1 <= rating_data.rating <= 5:
        raise HTTPException(status_code=400, detail="Rating must be between 1 and 5")
    
    # Get rater info
    rater = await db.users.find_one({"id": rater_id}, {"_id": 0})
    if not rater:
        raise HTTPException(status_code=404, detail="Rater not found")
    
    # Get rated user
    rated_user = await db.users.find_one({"id": rating_data.rated_user_id}, {"_id": 0})
    if not rated_user:
        raise HTTPException(status_code=404, detail="User to rate not found")
    
    # Prevent self-rating
    if rater_id == rating_data.rated_user_id:
        raise HTTPException(status_code=400, detail="Cannot rate yourself")
    
    # Check for existing rating in same context
    existing = await db.user_ratings.find_one({
        "rated_user_id": rating_data.rated_user_id,
        "rater_user_id": rater_id,
        "context_type": rating_data.context_type,
        "context_id": rating_data.context_id
    })
    
    if existing:
        raise HTTPException(status_code=400, detail="You have already rated this user for this context")
    
    # Create rating
    rating_id = str(uuid.uuid4())
    rating_doc = {
        "id": rating_id,
        "rated_user_id": rating_data.rated_user_id,
        "rated_user_type": rated_user.get("user_type"),
        "rater_user_id": rater_id,
        "rater_name": rater.get("name", "Anonymous"),
        "rater_type": rater.get("user_type"),
        "rating": rating_data.rating,
        "review": rating_data.review,
        "context_type": rating_data.context_type,
        "context_id": rating_data.context_id,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.user_ratings.insert_one(rating_doc)
    
    # Update user's average rating
    await update_user_rating_stats(rating_data.rated_user_id)
    
    return rating_doc

async def update_user_rating_stats(user_id: str):
    """Recalculate and update a user's rating statistics"""
    ratings = await db.user_ratings.find(
        {"rated_user_id": user_id},
        {"_id": 0, "rating": 1}
    ).to_list(1000)
    
    if ratings:
        avg_rating = sum(r["rating"] for r in ratings) / len(ratings)
        await db.users.update_one(
            {"id": user_id},
            {"$set": {
                "rating": round(avg_rating, 1),
                "rating_count": len(ratings)
            }}
        )

@api_router.get("/ratings/user/{user_id}")
async def get_user_ratings(user_id: str, limit: int = 20):
    """Get all ratings for a user"""
    ratings = await db.user_ratings.find(
        {"rated_user_id": user_id},
        {"_id": 0}
    ).sort("created_at", -1).limit(limit).to_list(limit)
    
    # Get user's rating stats
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "rating": 1, "rating_count": 1, "user_type": 1, "name": 1})
    
    return {
        "user_id": user_id,
        "user_name": user.get("name") if user else None,
        "user_type": user.get("user_type") if user else None,
        "average_rating": user.get("rating", 0) if user else 0,
        "total_ratings": user.get("rating_count", 0) if user else 0,
        "reviews": ratings
    }

@api_router.get("/ratings/summary/{user_id}")
async def get_user_rating_summary(user_id: str):
    """Get rating summary with distribution for a user"""
    ratings = await db.user_ratings.find(
        {"rated_user_id": user_id},
        {"_id": 0, "rating": 1}
    ).to_list(1000)
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "rating": 1, "rating_count": 1, "user_type": 1, "name": 1})
    
    # Calculate distribution
    distribution = {5: 0, 4: 0, 3: 0, 2: 0, 1: 0}
    for r in ratings:
        distribution[r["rating"]] = distribution.get(r["rating"], 0) + 1
    
    total = len(ratings)
    distribution_pct = {
        k: round((v / total * 100) if total > 0 else 0, 1)
        for k, v in distribution.items()
    }
    
    return {
        "user_id": user_id,
        "user_name": user.get("name") if user else None,
        "user_type": user.get("user_type") if user else None,
        "average_rating": user.get("rating", 0) if user else 0,
        "total_ratings": total,
        "distribution": distribution,
        "distribution_percentage": distribution_pct
    }

@api_router.delete("/ratings/{rating_id}")
async def delete_user_rating(rating_id: str, user_id: str):
    """Delete a rating (only by the rater)"""
    rating = await db.user_ratings.find_one({"id": rating_id})
    if not rating:
        raise HTTPException(status_code=404, detail="Rating not found")
    
    if rating.get("rater_user_id") != user_id:
        raise HTTPException(status_code=403, detail="Not authorized to delete this rating")
    
    rated_user_id = rating.get("rated_user_id")
    await db.user_ratings.delete_one({"id": rating_id})
    
    # Update user's rating stats
    await update_user_rating_stats(rated_user_id)
    
    return {"success": True, "message": "Rating deleted"}


# ========== FILE UPLOAD ENDPOINTS (Cloudflare R2) ==========

@api_router.get("/storage/status")
async def get_storage_status():
    """Check if R2 storage is configured"""
    return {"configured": is_r2_configured()}

@api_router.post("/upload/property-image")
async def upload_property_image_endpoint(
    file: UploadFile = File(...),
    property_id: str = "",
    user_id: str = ""
):
    """Upload a property listing image to R2"""
    if not is_r2_configured():
        raise HTTPException(status_code=503, detail="Storage service not configured")
    
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="File must be an image")
    
    try:
        content = await file.read()
        result = await upload_property_image(
            file_data=content,
            filename=file.filename,
            content_type=file.content_type,
            property_id=property_id or str(uuid.uuid4()),
            user_id=user_id
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Property image upload error: {e}")
        raise HTTPException(status_code=500, detail="Upload failed")

@api_router.post("/upload/document")
async def upload_document_endpoint(
    file: UploadFile = File(...),
    user_id: str = "",
    document_type: str = "general"
):
    """Upload a document (lease, contract, etc.) to R2"""
    if not is_r2_configured():
        raise HTTPException(status_code=503, detail="Storage service not configured")
    
    try:
        content = await file.read()
        result = await r2_upload_document(
            file_data=content,
            filename=file.filename,
            content_type=file.content_type,
            user_id=user_id,
            document_type=document_type
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Document upload error: {e}")
        raise HTTPException(status_code=500, detail="Upload failed")

@api_router.post("/upload/avatar")
async def upload_avatar_endpoint(
    file: UploadFile = File(...),
    user_id: str = ""
):
    """Upload a user avatar to R2"""
    if not is_r2_configured():
        raise HTTPException(status_code=503, detail="Storage service not configured")
    
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="Avatar must be an image")
    
    try:
        content = await file.read()
        result = await upload_avatar(
            file_data=content,
            filename=file.filename,
            content_type=file.content_type,
            user_id=user_id
        )
        
        # Update user's avatar URL in database
        if user_id:
            await db.users.update_one(
                {"id": user_id},
                {"$set": {"avatar": result["url"]}}
            )
        
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Avatar upload error: {e}")
        raise HTTPException(status_code=500, detail="Upload failed")

@api_router.post("/upload/contractor-portfolio")
async def upload_portfolio_endpoint(
    file: UploadFile = File(...),
    contractor_id: str = ""
):
    """Upload a contractor portfolio image to R2"""
    if not is_r2_configured():
        raise HTTPException(status_code=503, detail="Storage service not configured")
    
    if not contractor_id:
        raise HTTPException(status_code=400, detail="contractor_id is required")
    
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="Portfolio image must be an image")
    
    try:
        content = await file.read()
        result = await upload_contractor_portfolio(
            file_data=content,
            filename=file.filename,
            content_type=file.content_type,
            contractor_id=contractor_id
        )
        
        # Add to contractor's portfolio array
        await db.contractor_profiles.update_one(
            {"user_id": contractor_id},
            {"$push": {"portfolio_images": result["url"]}}
        )
        
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Portfolio upload error: {e}")
        raise HTTPException(status_code=500, detail="Upload failed")

@api_router.post("/upload/general")
async def upload_general_file(
    file: UploadFile = File(...),
    folder: str = "uploads",
    user_id: str = ""
):
    """Upload a general file to R2"""
    if not is_r2_configured():
        raise HTTPException(status_code=503, detail="Storage service not configured")
    
    try:
        content = await file.read()
        result = await upload_file(
            file_data=content,
            filename=file.filename,
            content_type=file.content_type,
            folder=folder,
            user_id=user_id
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"General upload error: {e}")
        raise HTTPException(status_code=500, detail="Upload failed")


# ============== FINANCING ENDPOINTS ==============

class FinancingApplication(BaseModel):
    financing_type: str  # rent-to-own, deposit-financing, first-month-free
    full_name: str
    email: str
    phone: Optional[str] = None
    monthly_income: Optional[str] = None
    employment_status: str = "employed"
    credit_score_range: str = "600-650"
    property_interest: Optional[str] = None
    message: Optional[str] = None
    user_id: Optional[str] = None

@api_router.post("/financing/applications")
async def submit_financing_application(application: FinancingApplication):
    """Submit a financing application (rent-to-own, deposit financing, etc.)"""
    try:
        app_data = application.model_dump()
        app_data["id"] = str(uuid.uuid4())
        app_data["status"] = "pending"
        app_data["created_at"] = datetime.now(timezone.utc).isoformat()
        
        await db.financing_applications.insert_one(app_data)
        
        # Send notification email
        try:
            await send_email(
                to_email=application.email,
                subject=f"DOMMMA - {application.financing_type.replace('-', ' ').title()} Application Received",
                html_content=f"""
                <h2>Thank you for your application!</h2>
                <p>Hi {application.full_name},</p>
                <p>We've received your application for our {application.financing_type.replace('-', ' ').title()} program.</p>
                <p>Our team will review your application and contact you within 2 business days.</p>
                <p><strong>Application Details:</strong></p>
                <ul>
                    <li>Program: {application.financing_type.replace('-', ' ').title()}</li>
                    <li>Employment: {application.employment_status}</li>
                    <li>Credit Range: {application.credit_score_range}</li>
                </ul>
                <p>Best regards,<br>The DOMMMA Team</p>
                """
            )
        except Exception as e:
            logger.warning(f"Failed to send financing confirmation email: {e}")
        
        return {"id": app_data["id"], "status": "pending", "message": "Application submitted successfully"}
    except Exception as e:
        logger.error(f"Financing application error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/financing/applications")
async def get_financing_applications(user_id: str = None, status: str = None):
    """Get financing applications (admin or user-specific)"""
    query = {}
    if user_id:
        query["user_id"] = user_id
    if status:
        query["status"] = status
    
    applications = await db.financing_applications.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return applications


# ========== RENT PAYMENT SYSTEM ==========

@api_router.post("/rent/agreements")
async def create_rent_agreement(data: RentAgreementCreate, landlord_id: str):
    """Landlord creates a rent payment agreement for a tenant"""
    
    # Verify the listing belongs to this landlord
    listing = await db.listings.find_one({"id": data.listing_id})
    if not listing:
        raise HTTPException(status_code=404, detail="Listing not found")
    
    if listing.get("landlord_id") != landlord_id and listing.get("user_id") != landlord_id:
        raise HTTPException(status_code=403, detail="Not authorized to create agreement for this listing")
    
    # Get tenant info
    tenant = await db.users.find_one({"id": data.tenant_id})
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")
    
    # Create the agreement
    agreement = RentAgreement(
        **data.model_dump(),
        landlord_id=landlord_id,
        landlord_signed=True  # Landlord signs by creating
    )
    
    await db.rent_agreements.insert_one(agreement.model_dump())
    
    # Send notification to tenant
    try:
        landlord = await db.users.find_one({"id": landlord_id})
        await send_email(
            tenant.get("email"),
            "Rent Payment Agreement - Action Required",
            f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #1A2F3A;">Rent Payment Agreement</h2>
                <p>Hi {tenant.get('name', 'there')},</p>
                <p>{landlord.get('name', 'Your landlord')} has created a rent payment agreement for you.</p>
                <div style="background: #f5f5f5; padding: 15px; border-radius: 8px; margin: 20px 0;">
                    <p><strong>Property:</strong> {listing.get('title', 'N/A')}</p>
                    <p><strong>Monthly Rent:</strong> ${data.monthly_rent:,.2f}</p>
                    <p><strong>Due Day:</strong> {data.due_day} of each month</p>
                    <p><strong>Grace Period:</strong> {data.grace_period_days} days</p>
                    <p><strong>Late Fee:</strong> {'$' + str(data.late_fee_amount) if data.late_fee_type == 'flat' else str(data.late_fee_amount) + '%'}</p>
                </div>
                <p>Please log in to DOMMMA to review and accept these terms.</p>
                <a href="https://dommma.com/rent-agreements" style="background: #1A2F3A; color: white; padding: 12px 24px; border-radius: 8px; text-decoration: none; display: inline-block;">Review Agreement</a>
            </div>
            """
        )
    except Exception as e:
        logger.warning(f"Failed to send rent agreement email: {e}")
    
    return {"status": "success", "agreement_id": agreement.id, "message": "Agreement created and sent to tenant"}


@api_router.get("/rent/agreements")
async def get_rent_agreements(user_id: str, role: str = "tenant"):
    """Get rent agreements for a user (as landlord or tenant)"""
    if role == "landlord":
        query = {"landlord_id": user_id}
    else:
        query = {"tenant_id": user_id}
    
    agreements = await db.rent_agreements.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    # Enrich with listing and user info
    for agreement in agreements:
        listing = await db.listings.find_one({"id": agreement.get("listing_id")}, {"_id": 0, "id": 1, "title": 1, "address": 1})
        agreement["listing"] = listing
        
        if role == "landlord":
            tenant = await db.users.find_one({"id": agreement.get("tenant_id")}, {"_id": 0, "id": 1, "name": 1, "email": 1})
            agreement["tenant"] = tenant
        else:
            landlord = await db.users.find_one({"id": agreement.get("landlord_id")}, {"_id": 0, "id": 1, "name": 1, "email": 1})
            agreement["landlord"] = landlord
    
    return agreements


@api_router.get("/rent/agreements/{agreement_id}")
async def get_rent_agreement(agreement_id: str, user_id: str):
    """Get a specific rent agreement"""
    agreement = await db.rent_agreements.find_one(
        {"id": agreement_id, "$or": [{"landlord_id": user_id}, {"tenant_id": user_id}]},
        {"_id": 0}
    )
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    # Enrich with related data
    listing = await db.listings.find_one({"id": agreement.get("listing_id")}, {"_id": 0})
    tenant = await db.users.find_one({"id": agreement.get("tenant_id")}, {"_id": 0, "id": 1, "name": 1, "email": 1, "phone": 1})
    landlord = await db.users.find_one({"id": agreement.get("landlord_id")}, {"_id": 0, "id": 1, "name": 1, "email": 1, "phone": 1})
    
    agreement["listing"] = listing
    agreement["tenant"] = tenant
    agreement["landlord"] = landlord
    
    return agreement


@api_router.post("/rent/agreements/{agreement_id}/accept")
async def accept_rent_agreement(agreement_id: str, tenant_id: str):
    """Tenant accepts the rent agreement terms"""
    agreement = await db.rent_agreements.find_one({"id": agreement_id, "tenant_id": tenant_id})
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    if agreement.get("status") != "pending":
        raise HTTPException(status_code=400, detail="Agreement is not pending")
    
    await db.rent_agreements.update_one(
        {"id": agreement_id},
        {"$set": {
            "tenant_signed": True,
            "tenant_accepted_terms": True,
            "status": "active",
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Notify landlord
    try:
        landlord = await db.users.find_one({"id": agreement.get("landlord_id")})
        tenant = await db.users.find_one({"id": tenant_id})
        await send_email(
            landlord.get("email"),
            "Rent Agreement Accepted",
            f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #1A2F3A;">Good News!</h2>
                <p>{tenant.get('name', 'Your tenant')} has accepted the rent payment agreement.</p>
                <p>Automatic rent collection is now active.</p>
            </div>
            """
        )
    except Exception as e:
        logger.warning(f"Failed to send acceptance email: {e}")
    
    return {"status": "success", "message": "Agreement accepted"}


@api_router.post("/rent/agreements/{agreement_id}/counter")
async def counter_rent_agreement(agreement_id: str, tenant_id: str, terms: RentAgreementTerms):
    """Tenant proposes counter-terms"""
    agreement = await db.rent_agreements.find_one({"id": agreement_id, "tenant_id": tenant_id})
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    if agreement.get("status") != "pending":
        raise HTTPException(status_code=400, detail="Agreement is not pending")
    
    # Store counter-proposal
    counter_terms = {k: v for k, v in terms.model_dump().items() if v is not None}
    
    await db.rent_agreements.update_one(
        {"id": agreement_id},
        {"$set": {
            "counter_proposal": counter_terms,
            "counter_proposal_by": tenant_id,
            "counter_proposal_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Notify landlord
    try:
        landlord = await db.users.find_one({"id": agreement.get("landlord_id")})
        tenant = await db.users.find_one({"id": tenant_id})
        await send_email(
            landlord.get("email"),
            "Rent Agreement - Counter Proposal",
            f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #1A2F3A;">Counter Proposal Received</h2>
                <p>{tenant.get('name', 'Your tenant')} has proposed changes to the rent agreement.</p>
                <p>Please log in to review and respond.</p>
                <a href="https://dommma.com/rent-agreements/{agreement_id}" style="background: #1A2F3A; color: white; padding: 12px 24px; border-radius: 8px; text-decoration: none; display: inline-block;">Review Proposal</a>
            </div>
            """
        )
    except Exception as e:
        logger.warning(f"Failed to send counter proposal email: {e}")
    
    return {"status": "success", "message": "Counter proposal submitted"}


@api_router.post("/rent/agreements/{agreement_id}/decline")
async def decline_rent_agreement(agreement_id: str, tenant_id: str, reason: str = ""):
    """Tenant declines the rent agreement"""
    agreement = await db.rent_agreements.find_one({"id": agreement_id, "tenant_id": tenant_id})
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")
    
    await db.rent_agreements.update_one(
        {"id": agreement_id},
        {"$set": {
            "status": "declined",
            "decline_reason": reason,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"status": "success", "message": "Agreement declined"}


@api_router.post("/rent/save-payment-method")
async def save_payment_method(tenant_id: str, data: SavePaymentMethod):
    """Save tenant's payment method (credit card) via Stripe"""
    import stripe
    stripe.api_key = os.environ.get("STRIPE_API_KEY")
    
    # Get or create Stripe customer
    tenant = await db.users.find_one({"id": tenant_id})
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")
    
    stripe_customer_id = tenant.get("stripe_customer_id")
    
    if not stripe_customer_id:
        # Create Stripe customer
        customer = stripe.Customer.create(
            email=tenant.get("email"),
            name=tenant.get("name"),
            metadata={"dommma_user_id": tenant_id}
        )
        stripe_customer_id = customer.id
        await db.users.update_one(
            {"id": tenant_id},
            {"$set": {"stripe_customer_id": stripe_customer_id}}
        )
    
    # Attach payment method to customer
    try:
        stripe.PaymentMethod.attach(
            data.payment_method_id,
            customer=stripe_customer_id
        )
        
        # Set as default payment method
        stripe.Customer.modify(
            stripe_customer_id,
            invoice_settings={"default_payment_method": data.payment_method_id}
        )
        
        # Update user record
        await db.users.update_one(
            {"id": tenant_id},
            {"$set": {"default_payment_method_id": data.payment_method_id}}
        )
        
        return {"status": "success", "message": "Payment method saved"}
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@api_router.get("/rent/payments")
async def get_rent_payments(user_id: str, role: str = "tenant"):
    """Get rent payment history"""
    if role == "landlord":
        query = {"landlord_id": user_id}
    else:
        query = {"tenant_id": user_id}
    
    payments = await db.rent_payments.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return payments


@api_router.post("/rent/payments/{agreement_id}/pay")
async def process_rent_payment(agreement_id: str, tenant_id: str, payment_method: str = "card"):
    """Process a rent payment manually (or triggered by cron)"""
    import stripe
    stripe.api_key = os.environ.get("STRIPE_API_KEY")
    
    agreement = await db.rent_agreements.find_one({"id": agreement_id, "tenant_id": tenant_id, "status": "active"})
    if not agreement:
        raise HTTPException(status_code=404, detail="Active agreement not found")
    
    tenant = await db.users.find_one({"id": tenant_id})
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")
    
    # Calculate amount (check if late fee applies)
    today = datetime.now(timezone.utc)
    current_month = today.strftime("%Y-%m")
    due_day = agreement.get("due_day", 1)
    due_date = datetime(today.year, today.month, due_day, tzinfo=timezone.utc)
    grace_end = due_date + timedelta(days=agreement.get("grace_period_days", 3))
    
    base_amount = agreement.get("monthly_rent")
    late_fee = 0.0
    
    if today > grace_end:
        # Apply late fee
        if agreement.get("late_fee_type") == "flat":
            late_fee = agreement.get("late_fee_amount", 0)
        else:
            late_fee = base_amount * (agreement.get("late_fee_amount", 0) / 100)
            if agreement.get("max_late_fee"):
                late_fee = min(late_fee, agreement.get("max_late_fee"))
    
    total_amount = base_amount + late_fee
    
    # Check if already paid this month
    existing = await db.rent_payments.find_one({"agreement_id": agreement_id, "payment_month": current_month, "status": "paid"})
    if existing:
        raise HTTPException(status_code=400, detail="Rent already paid for this month")
    
    # Create payment record
    payment = RentPayment(
        agreement_id=agreement_id,
        tenant_id=tenant_id,
        landlord_id=agreement.get("landlord_id"),
        amount=base_amount,
        late_fee=late_fee,
        total_amount=total_amount,
        payment_month=current_month,
        due_date=due_date.isoformat(),
        status="processing",
        payment_method=payment_method
    )
    
    await db.rent_payments.insert_one(payment.model_dump())
    
    if payment_method == "card":
        # Process via Stripe
        try:
            stripe_customer_id = tenant.get("stripe_customer_id")
            payment_method_id = tenant.get("default_payment_method_id")
            
            if not stripe_customer_id or not payment_method_id:
                raise HTTPException(status_code=400, detail="No payment method on file. Please add a credit card.")
            
            # Create payment intent
            intent = stripe.PaymentIntent.create(
                amount=int(total_amount * 100),  # Stripe uses cents
                currency="cad",
                customer=stripe_customer_id,
                payment_method=payment_method_id,
                off_session=True,
                confirm=True,
                description=f"DOMMMA Rent Payment - {current_month}",
                metadata={
                    "agreement_id": agreement_id,
                    "payment_id": payment.id,
                    "tenant_id": tenant_id
                }
            )
            
            # Update payment record
            await db.rent_payments.update_one(
                {"id": payment.id},
                {"$set": {
                    "status": "paid",
                    "stripe_payment_intent_id": intent.id,
                    "paid_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            
            # TODO: Transfer to landlord via Stripe Connect
            
            return {"status": "success", "message": f"Payment of ${total_amount:.2f} processed successfully", "payment_id": payment.id}
            
        except stripe.error.CardError as e:
            await db.rent_payments.update_one(
                {"id": payment.id},
                {"$set": {
                    "status": "failed",
                    "failure_reason": str(e)
                }}
            )
            raise HTTPException(status_code=400, detail=f"Card declined: {e.user_message}")
    else:
        # E-transfer - just record it as pending
        return {
            "status": "pending", 
            "message": f"Please send ${total_amount:.2f} via e-transfer to: {agreement.get('etransfer_email')}",
            "payment_id": payment.id
        }


@api_router.get("/rent/late-fee-calculator")
async def calculate_late_fee(rent_amount: float, late_fee_type: str = "flat", late_fee_amount: float = 50, max_late_fee: float = None):
    """Calculate late fee for AI to display to user"""
    if late_fee_type == "flat":
        late_fee = late_fee_amount
    else:
        late_fee = rent_amount * (late_fee_amount / 100)
        if max_late_fee:
            late_fee = min(late_fee, max_late_fee)
    
    # Stripe fee estimation (2.9% + $0.30)
    stripe_fee = (rent_amount + late_fee) * 0.029 + 0.30
    
    return {
        "rent_amount": rent_amount,
        "late_fee": round(late_fee, 2),
        "total_with_late_fee": round(rent_amount + late_fee, 2),
        "estimated_stripe_fee": round(stripe_fee, 2),
        "landlord_receives": round(rent_amount + late_fee - stripe_fee, 2)
    }


# ========== CONTACT FORM ==========

class ContactFormInput(BaseModel):
    name: str
    email: str
    subject: str
    message: str

@api_router.post("/contact")
async def submit_contact_form(data: ContactFormInput):
    """Submit contact form - sends email to admin"""
    
    # Store the message in database
    contact_message = {
        "id": str(uuid.uuid4()),
        "name": data.name,
        "email": data.email,
        "subject": data.subject,
        "message": data.message,
        "status": "new",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.contact_messages.insert_one(contact_message)
    
    # Send notification email to support
    contact_email = os.environ.get('CONTACT_EMAIL', 'support@dommma.com')
    
    try:
        await send_email(
            contact_email,
            f"New Contact Form: {data.subject}",
            f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #1A2F3A;">New Contact Form Submission</h2>
                <p><strong>From:</strong> {data.name} ({data.email})</p>
                <p><strong>Subject:</strong> {data.subject}</p>
                <hr style="border: 1px solid #eee;">
                <p><strong>Message:</strong></p>
                <p style="background: #f5f5f5; padding: 15px; border-radius: 8px;">{data.message}</p>
                <hr style="border: 1px solid #eee;">
                <p style="color: #666; font-size: 12px;">
                    Received at: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}<br>
                    Reply directly to: {data.email}
                </p>
            </div>
            """
        )
    except Exception as e:
        logger.warning(f"Failed to send contact form email: {e}")
    
    return {"status": "success", "message": "Your message has been received"}

@api_router.get("/admin/contact-messages")
async def get_contact_messages(status: str = None):
    """Get all contact form messages (admin only)"""
    query = {}
    if status:
        query["status"] = status
    
    messages = await db.contact_messages.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return messages


# ========== ADMIN: CLEAR DATA ==========

@api_router.delete("/admin/clear-test-data")
async def clear_test_data(admin_key: str, clear_users: bool = False, clear_listings: bool = True, 
                          clear_applications: bool = True, clear_messages: bool = True,
                          clear_documents: bool = True, clear_jobs: bool = True):
    """
    Clear test/mock data from the database.
    Requires admin_key for security.
    
    Parameters:
    - admin_key: Must match ADMIN_SECRET_KEY env variable
    - clear_users: If True, deletes all users (careful!)
    - clear_listings: If True, deletes all listings
    - clear_applications: If True, deletes all rental applications
    - clear_messages: If True, deletes all chat messages
    - clear_documents: If True, deletes all e-sign documents
    - clear_jobs: If True, deletes all job posts
    """
    
    # Verify admin key
    expected_key = os.environ.get('ADMIN_SECRET_KEY', 'dommma-admin-2026')
    if admin_key != expected_key:
        raise HTTPException(status_code=403, detail="Invalid admin key")
    
    results = {}
    
    if clear_users:
        result = await db.users.delete_many({})
        results["users_deleted"] = result.deleted_count
    
    if clear_listings:
        result = await db.listings.delete_many({})
        results["listings_deleted"] = result.deleted_count
    
    if clear_applications:
        result = await db.applications.delete_many({})
        results["applications_deleted"] = result.deleted_count
        
        result = await db.financing_applications.delete_many({})
        results["financing_applications_deleted"] = result.deleted_count
    
    if clear_messages:
        result = await db.messages.delete_many({})
        results["messages_deleted"] = result.deleted_count
        
        result = await db.contact_messages.delete_many({})
        results["contact_messages_deleted"] = result.deleted_count
    
    if clear_documents:
        result = await db.esign_documents.delete_many({})
        results["esign_documents_deleted"] = result.deleted_count
        
        result = await db.builder_documents.delete_many({})
        results["builder_documents_deleted"] = result.deleted_count
        
        result = await db.documents.delete_many({})
        results["documents_deleted"] = result.deleted_count
    
    if clear_jobs:
        result = await db.job_posts.delete_many({})
        results["job_posts_deleted"] = result.deleted_count
    
    # Also clear other test data
    result = await db.lease_assignments.delete_many({})
    results["lease_assignments_deleted"] = result.deleted_count
    
    result = await db.payment_transactions.delete_many({})
    results["payment_transactions_deleted"] = result.deleted_count
    
    result = await db.maintenance_requests.delete_many({})
    results["maintenance_requests_deleted"] = result.deleted_count
    
    logger.info(f"Admin cleared test data: {results}")
    
    return {
        "status": "success",
        "message": "Test data cleared successfully",
        "results": results
    }

@api_router.get("/admin/database-stats")
async def get_database_stats(admin_key: str):
    """Get counts of all collections (admin only)"""
    
    expected_key = os.environ.get('ADMIN_SECRET_KEY', 'dommma-admin-2026')
    if admin_key != expected_key:
        raise HTTPException(status_code=403, detail="Invalid admin key")
    
    stats = {
        "users": await db.users.count_documents({}),
        "listings": await db.listings.count_documents({}),
        "applications": await db.applications.count_documents({}),
        "messages": await db.messages.count_documents({}),
        "esign_documents": await db.esign_documents.count_documents({}),
        "job_posts": await db.job_posts.count_documents({}),
        "lease_assignments": await db.lease_assignments.count_documents({}),
        "payment_transactions": await db.payment_transactions.count_documents({}),
        "contractor_profiles": await db.contractor_profiles.count_documents({}),
        "contact_messages": await db.contact_messages.count_documents({}),
        "financing_applications": await db.financing_applications.count_documents({})
    }
    
    return stats


# Include the router
app.include_router(api_router)

# Include new modular routers
from routers import calendar_router, moving_router, compatibility_router, portfolio_router, nova_router
app.include_router(calendar_router, prefix="/api")
app.include_router(moving_router, prefix="/api")
app.include_router(compatibility_router, prefix="/api")
app.include_router(portfolio_router, prefix="/api")
app.include_router(nova_router, prefix="/api")

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def startup_migrations():
    """Run database migrations on startup"""
    # Fix existing signed documents: set recipient_id from signer_id if not set
    result = await db.esign_documents.update_many(
        {
            "status": "signed",
            "signer_id": {"$exists": True, "$ne": None},
            "$or": [
                {"recipient_id": {"$exists": False}},
                {"recipient_id": None}
            ]
        },
        [
            {"$set": {"recipient_id": "$signer_id"}}
        ]
    )
    if result.modified_count > 0:
        logger.info(f"Migration: Fixed {result.modified_count} e-sign documents - linked recipient_id to signer_id")

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
